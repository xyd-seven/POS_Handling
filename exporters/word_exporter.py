# -*- coding: utf-8 -*-
import io
import os
from datetime import datetime
from PySide6.QtCore import Qt
from PySide6.QtWidgets import QFileDialog, QMessageBox, QProgressDialog, QApplication

def export_word_report(parent_window, segments, truth, table_metrics, canvases_dict, config=None):
    try:
        import docx
        from docx.shared import Inches
    except ImportError:
        QMessageBox.warning(parent_window, "提示", "未检测到 python-docx 模块。请在终端运行: pip install python-docx")
        return

    if not segments:
        QMessageBox.warning(parent_window, "提示", "没有测试数据可供导出。")
        return

    save_path, _ = QFileDialog.getSaveFileName(parent_window, "导出 Word 报告", "测试报告.docx", "Word Documents (*.docx)")
    if not save_path:
        return

    try:
        progress = QProgressDialog("正在生成 Word 报告...", "取消", 0, 11, parent_window)
        progress.setWindowTitle("导出报告")
        progress.setWindowModality(Qt.WindowModal)
        progress.setMinimumDuration(0)
        progress.setValue(0)

        doc = docx.Document()
        doc.add_heading('GNSS 定位精度测试报告', 0)

        doc.add_heading('1. 测试概览', level=1)
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        progress.setValue(1)
        QApplication.processEvents()
        if progress.wasCanceled():
            return

        doc.add_heading('2. 精度指标对比', level=1)
        table = doc.add_table(rows=1, cols=table_metrics.columnCount())
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for j in range(table_metrics.columnCount()):
            hdr_cells[j].text = table_metrics.horizontalHeaderItem(j).text()

        for i in range(table_metrics.rowCount()):
            row_cells = table.add_row().cells
            for j in range(table_metrics.columnCount()):
                item = table_metrics.item(i, j)
                row_cells[j].text = item.text() if item else ""
        progress.setValue(2)
        QApplication.processEvents()
        if progress.wasCanceled():
            return

        doc.add_heading('3. 图表分析', level=1)

        def add_plot_to_doc(canvas, title, progress_value):
            doc.add_heading(title, level=2)
            buf = io.BytesIO()
            canvas.figure.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            doc.add_picture(buf, width=Inches(6.0))
            buf.close()
            progress.setValue(progress_value)
            QApplication.processEvents()

        progress.setLabelText("正在渲染图表...")
        time_zone = getattr(parent_window, 'time_zone', 'UTC')
        show_extrema = getattr(parent_window, 'show_extrema', True)
        x_axis_mode = getattr(parent_window, 'x_axis_mode', '历元数')
        show_sats = getattr(parent_window, 'cb_show_sats', None)
        is_sats_checked = show_sats.isChecked() if show_sats else False
        show_absolute_alt = getattr(parent_window, 'show_absolute_alt', False)
        show_raw_alt = getattr(parent_window, 'show_raw_alt', False)
        speed_unit = getattr(parent_window, 'speed_unit', 'm/s')
        cdf_mode = getattr(parent_window, 'cdf_mode', 'horizontal')
        show_cdf_quantiles = getattr(parent_window, 'show_cdf_quantiles', True)

        c_scatter = canvases_dict['scatter']
        c_trajectory = canvases_dict['trajectory']
        c_status = canvases_dict['status']
        c_epoch_h = canvases_dict['epoch_h']
        c_epoch_v = canvases_dict['epoch_v']
        c_epoch_enu = canvases_dict['epoch_enu']
        c_speed = canvases_dict['speed']
        c_cdf = canvases_dict['cdf']

        c_scatter.render_data('scatter', segments, truth, time_zone)
        c_trajectory.render_data('trajectory', segments, truth)
        c_status.render_data('status', segments, truth, time_zone)
        c_epoch_h.render_data('epoch_h', segments, truth, time_zone, show_extrema=show_extrema, x_axis_mode=x_axis_mode, show_sats=is_sats_checked)
        c_epoch_v.render_data('epoch_v', segments, truth, time_zone, show_absolute_alt, show_extrema=show_extrema, x_axis_mode=x_axis_mode, show_sats=is_sats_checked, show_raw_alt=show_raw_alt)
        c_epoch_enu.render_data('epoch_enu', segments, truth, time_zone, x_axis_mode=x_axis_mode, show_stats=True)
        c_speed.render_data('speed', segments, truth, time_zone, x_axis_mode=x_axis_mode, show_stats=True, speed_unit=speed_unit)
        c_cdf.render_data('cdf', segments, truth, cdf_mode=cdf_mode, speed_unit=speed_unit, show_quantiles=show_cdf_quantiles)
        progress.setValue(3)
        QApplication.processEvents()
        if progress.wasCanceled():
            return

        progress.setLabelText("正在写入图表...")
        add_plot_to_doc(c_trajectory, '3.1 绝对二维轨迹投影图', 4)
        if progress.wasCanceled():
            return
        add_plot_to_doc(c_scatter, '3.2 定位偏差分布图 (靶心图)', 5)
        if progress.wasCanceled():
            return
        add_plot_to_doc(c_status, '3.3 定位解状态分布', 6)
        if progress.wasCanceled():
            return
        add_plot_to_doc(c_epoch_h, '3.4 水平位置误差分布', 7)
        if progress.wasCanceled():
            return
        add_plot_to_doc(c_epoch_v, '3.5 高程位置误差分布', 8)
        if progress.wasCanceled():
            return
        add_plot_to_doc(c_epoch_enu, '3.6 ENU三向误差时域分布', 9)
        if progress.wasCanceled():
            return
        add_plot_to_doc(c_speed, '3.7 动态速度跟踪与误差分布', 10)
        if progress.wasCanceled():
            return
        add_plot_to_doc(c_cdf, '3.8 定位误差累积分布曲线 (CDF)', 11)
        if progress.wasCanceled():
            return

        progress.setLabelText("正在保存报告...")
        doc.save(save_path)
        progress.close()
        QMessageBox.information(parent_window, "导出成功", "Word 报告已成功导出至:\n" + str(save_path))
    except PermissionError:
        QMessageBox.critical(parent_window, "错误", "目标 Word 文件正被另一程序打开，请关闭后重试。")
    except Exception as e:
        QMessageBox.critical(parent_window, "错误", f"报告生成失败: {str(e)}")
