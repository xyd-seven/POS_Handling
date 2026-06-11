# -*- coding: utf-8 -*-
"""
VCOM (精度分析与转换工具 Qt版) - 主运行模块 (双向网格与双误差曲线切换版)
"""
import sys
import os
import json
import math
import copy
from collections import deque
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QHBoxLayout,
                             QVBoxLayout, QPushButton, QLabel, QLineEdit,
                             QListWidget, QListWidgetItem, QTableWidget,
                             QTableWidgetItem, QTabWidget, QGroupBox, QSplitter,
                             QHeaderView, QFileDialog, QMessageBox, QMenuBar,
                             QComboBox, QCheckBox, QProgressDialog, QGridLayout,
                             QTextEdit, QSlider)
from PySide6.QtCore import Qt, Signal, QThread, QPointF, QRectF, QTimer
from PySide6.QtGui import QColor, QIcon, QPixmap, QPainter, QFont, QPen, QBrush, QTextCursor
from PySide6.QtSerialPort import QSerialPort, QSerialPortInfo

# 引入 Matplotlib 导航工具栏以支持缩放和拖拽
from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure

from gnss_parser import (parse_log_line, convert_pogos_to_gga, calculate_metrics,
                    time_str_to_seconds, seconds_to_time_str, gps_tow_to_utc_time,
                    interpolate_dynamic_truth, parse_bk_frame, BKStreamParser,
                    crc16_ccitt)
from plot_widget import PlotWidget
from ui_main import QSS_STYLE, SegmentListItemWidget
from settings_dialog import SettingsDialog
from PIL import Image, ImageDraw

def ensure_flag_icons():
    # Target directory inside workspace
    workspace_dir = os.path.dirname(os.path.abspath(__file__))
    flags_dir = os.path.join(workspace_dir, "resources", "flags")
    os.makedirs(flags_dir, exist_ok=True)

    w, h = 20, 13

    flag_paths = {
        'GPS': os.path.join(flags_dir, 'gps.png'),
        'BD': os.path.join(flags_dir, 'bd.png'),
        'GL': os.path.join(flags_dir, 'gl.png'),
        'GA': os.path.join(flags_dir, 'ga.png'),
        'QZSS': os.path.join(flags_dir, 'qzss.png'),
        'IRNSS': os.path.join(flags_dir, 'irnss.png'),
        'SBAS': os.path.join(flags_dir, 'sbas.png'),
    }

    # Check if all files exist
    all_exist = True
    for path in flag_paths.values():
        if not os.path.exists(path):
            all_exist = False
            break

    if all_exist:
        return flag_paths

    # 1. US Flag (GPS)
    img = Image.new('RGB', (w, h), '#FFFFFF')
    draw = ImageDraw.Draw(img)
    for i in range(13):
        color = '#EF4444' if i % 2 == 0 else '#FFFFFF'
        draw.rectangle([0, i, w - 1, i], fill=color)
    draw.rectangle([0, 0, 8, 6], fill='#1E3A8A')
    stars = [(1, 1), (3, 1), (5, 1), (7, 1),
             (2, 2), (4, 2), (6, 2),
             (1, 3), (3, 3), (5, 3), (7, 3),
             (2, 4), (4, 4), (6, 4),
             (1, 5), (3, 5), (5, 5), (7, 5)]
    for sx, sy in stars:
        draw.point((sx, sy), fill='#FFFFFF')
    img.save(flag_paths['GPS'])

    # 2. China Flag (BD)
    img = Image.new('RGB', (w, h), '#EF4444')
    draw = ImageDraw.Draw(img)
    draw.point((2, 2), fill='#F59E0B')
    draw.point((3, 2), fill='#F59E0B')
    draw.point((2, 3), fill='#F59E0B')
    draw.point((3, 3), fill='#F59E0B')
    draw.point((5, 1), fill='#F59E0B')
    draw.point((6, 2), fill='#F59E0B')
    draw.point((6, 4), fill='#F59E0B')
    draw.point((5, 5), fill='#F59E0B')
    img.save(flag_paths['BD'])

    # 3. Russia Flag (GL)
    img = Image.new('RGB', (w, h), '#FFFFFF')
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w - 1, 3], fill='#FFFFFF')
    draw.rectangle([0, 4, w - 1, 7], fill='#1D4ED8')
    draw.rectangle([0, 8, w - 1, 12], fill='#EF4444')
    img.save(flag_paths['GL'])

    # 4. EU Flag (GA)
    img = Image.new('RGB', (w, h), '#1E40AF')
    draw = ImageDraw.Draw(img)
    cx, cy = 10, 6
    r = 3
    for angle in range(0, 360, 30):
        rad = math.radians(angle)
        sx = int(cx + r * math.cos(rad) + 0.5)
        sy = int(cy + r * math.sin(rad) + 0.5)
        draw.point((sx, sy), fill='#F59E0B')
    img.save(flag_paths['GA'])

    # 5. Japan Flag (QZSS)
    img = Image.new('RGB', (w, h), '#FFFFFF')
    draw = ImageDraw.Draw(img)
    draw.ellipse([7, 3, 13, 9], fill='#EF4444')
    img.save(flag_paths['QZSS'])

    # 6. India Flag (IRNSS)
    img = Image.new('RGB', (w, h), '#F97316')
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w - 1, 3], fill='#F97316')
    draw.rectangle([0, 4, w - 1, 7], fill='#FFFFFF')
    draw.rectangle([0, 8, w - 1, 12], fill='#15803D')
    draw.point((10, 5), fill='#1D4ED8')
    draw.point((10, 6), fill='#1D4ED8')
    draw.point((9, 5), fill='#1D4ED8')
    draw.point((11, 5), fill='#1D4ED8')
    img.save(flag_paths['IRNSS'])

    # 7. SBAS Flag (Satellite Symbol)
    img = Image.new('RGB', (w, h), '#0369A1')
    draw = ImageDraw.Draw(img)
    draw.rectangle([8, 4, 11, 8], fill='#94A3B8')
    draw.rectangle([3, 5, 7, 7], fill='#38BDF8')
    draw.rectangle([12, 5, 16, 7], fill='#38BDF8')
    draw.rectangle([9, 9, 10, 10], fill='#94A3B8')
    img.save(flag_paths['SBAS'])

    return flag_paths

def get_sat_info(prefix_or_talker, prn):
    p = prefix_or_talker.upper() if prefix_or_talker else ''

    # 优先支持已完全确定的系统前缀，避免与 talker ID 重合冲突
    if p == 'SBAS':
        return 'SBAS', prn, 'S'
    elif p == 'QZSS':
        return 'QZSS', prn, 'Q'
    elif p == 'IRNSS':
        return 'IRNSS', prn, 'I'

    # 严格按照博通 BK166X 协议说明书 表 2-7 定义的系统与 PRN 范围进行映射
    if 'GP' in p or 'GPS' in p:
        if 1 <= prn <= 32:
            return 'GPS', prn, 'G'
        elif 33 <= prn <= 64:
            return 'SBAS', prn, 'S'
        elif 193 <= prn <= 202:
            return 'QZSS', prn, 'Q'
    elif 'GL' in p or 'GLONASS' in p:
        if 33 <= prn <= 64:
            return 'SBAS', prn, 'S'
        elif 65 <= prn <= 99:
            return 'GL', prn, 'R'
    elif 'GA' in p or 'GALILEO' in p or 'GAL' in p:
        if 37 <= prn <= 64:
            return 'SBAS', prn, 'S'
        elif 1 <= prn <= 36:
            return 'GA', prn, 'E'
    elif 'BD' in p or 'GB' in p or 'BDS' in p:
        if 1 <= prn <= 63:
            return 'BD', prn, 'B'
    elif 'IR' in p or 'GI' in p or 'IRNSS' in p:
        if 1 <= prn <= 14:
            return 'IRNSS', prn, 'I'
        elif 15 <= prn <= 64:
            return 'SBAS', prn, 'S'

    # 兜底默认分类
    if 'GL' in p:
        return 'GL', prn, 'R'
    elif 'GA' in p:
        return 'GA', prn, 'E'
    elif 'BD' in p or 'GB' in p:
        return 'BD', prn, 'B'
    elif 'IR' in p or 'GI' in p:
        return 'IRNSS', prn, 'I'
    return 'GPS', prn, 'G'

class CNoPlotCanvas(FigureCanvas):
    def __init__(self, parent=None, width=5, height=3, dpi=100):
        # 创建符合主界面 Dark 风格背景色（#172033）的 figure
        self.fig = Figure(figsize=(width, height), dpi=dpi, facecolor='#172033')
        self.ax = self.fig.add_subplot(111)
        self.fig.tight_layout()
        super().__init__(self.fig)
        self.setParent(parent)
        self._apply_base_style()
        self._last_layout_channel_count = None
        self._is_resizing = False
        self._last_hover_time = 0

        # 自动生成并加载国旗图片路径
        self.flag_paths = ensure_flag_icons()
        self.flag_images = {}
        try:
            import matplotlib.image as mpimg
            for prefix, path in self.flag_paths.items():
                if os.path.exists(path):
                    self.flag_images[prefix] = mpimg.imread(path)
        except Exception as e:
            print(f"Error pre-loading flag icons: {e}")

        # 初始化渲染星体位置列表，并绑定鼠标移动事件以展示悬浮气泡
        self.rendered_sats = []
        self._motion_cid = self.fig.canvas.mpl_connect('motion_notify_event', self.on_mouse_move)

    def _apply_base_style(self):
        # 预设图表坐标轴背景色为暗黑色 #0B1120
        self.ax.set_facecolor('#0B1120')
        self.ax.spines['bottom'].set_color('#1E293B')
        self.ax.spines['top'].set_color('#1E293B')
        self.ax.spines['right'].set_color('#1E293B')
        self.ax.spines['left'].set_color('#1E293B')
        self.ax.tick_params(colors='#94A3B8', labelsize=8)
        self.ax.yaxis.grid(True, linestyle='--', color='#1E293B', alpha=0.5)
        self.ax.set_axisbelow(True)

        # 固定 Y 轴在 0~55 dB-Hz
        self.ax.set_ylim(0, 55)
        self.ax.set_ylabel("载噪比 C/No (dB-Hz)", color='#94A3B8', fontsize=9)

    def resizeEvent(self, event):
        from PySide6.QtWidgets import QWidget
        from PySide6.QtCore import QTimer
        QWidget.resizeEvent(self, event)
        self._is_resizing = True
        self._pending_width = event.size().width()
        self._pending_height = event.size().height()
        if not hasattr(self, '_resize_timer'):
            self._resize_timer = QTimer(self)
            self._resize_timer.setSingleShot(True)
            self._resize_timer.timeout.connect(self._handle_delayed_resize)
        self._resize_timer.start(250)

    def _handle_delayed_resize(self):
        if hasattr(self, '_pending_width') and hasattr(self, '_pending_height'):
            dpival = self.figure.dpi
            winch = self._pending_width / dpival
            hinch = self._pending_height / dpival
            self.figure.set_size_inches(winch, hinch, forward=False)
            self._is_resizing = False
            self.draw_idle()

    def get_signal_name(self, pref, sig_id):
        if pref == 'GPS':
            if sig_id == '1': return 'L1CA'
            elif sig_id == '2': return 'L1P'
            elif sig_id == '3': return 'L1M'
            elif sig_id == '4': return 'L2P'
            elif sig_id == '5': return 'L2CM'
            elif sig_id == '6': return 'L2CL'
            elif sig_id == '7': return 'L5I'
            elif sig_id == '8': return 'L5Q'
            return f"L{sig_id}"
        elif pref == 'BD':
            if sig_id == '1': return 'B1I'
            elif sig_id == '3': return 'B1C'
            elif sig_id == '5': return 'B2a'
            elif sig_id == 'B': return 'B2I'
            elif sig_id == '6': return 'B2b'
            elif sig_id == '7': return 'B2AB'
            elif sig_id == '8': return 'B3I'
            return f"B{sig_id}"
        elif pref == 'GL':
            if sig_id == '1': return 'G1CA'
            elif sig_id == '2': return 'G1P'
            elif sig_id == '3': return 'G2CA'
            elif sig_id == '4': return 'G2P'
            return f"G{sig_id}"
        elif pref == 'GA':
            if sig_id == '1': return 'E5a'
            elif sig_id == '2': return 'E5b'
            elif sig_id == '3': return 'E5ab'
            elif sig_id == '4': return 'E6A'
            elif sig_id == '5': return 'E6B'
            elif sig_id == '6': return 'L1A'
            elif sig_id == '7': return 'E1'
            return f"E{sig_id}"
        elif pref == 'QZSS':
            if sig_id == '1': return 'L1CA'
            elif sig_id == '7': return 'L5I'
            return f"Q{sig_id}"
        elif pref == 'IRNSS':
            if sig_id == '1': return 'L5I'
            return f"I{sig_id}"
        elif pref == 'SBAS':
            if sig_id == '1': return 'L1CA'
            elif sig_id == '7': return 'L5I'
            elif sig_id == '8': return 'L5Q'
            return f"S{sig_id}"
        return sig_id

    def render_cno(self, gsv_satellites, used_satellites=None, has_gsa=False, sat_metadata=None):
        """
        绘制可见卫星载噪比柱状图，区分在用和可见卫星。
        gsv_satellites: 字典，格式为 (prefix, prn) -> {signal_id: snr}
        used_satellites: 集合，包含当前在用卫星 (prefix, prn) 元组
        has_gsa: 布尔值，当前会话是否已接收到过 GSA 语句
        sat_metadata: 字典，包含卫星仰角方位角元数据 (prefix, prn) -> {'elevation': elev, 'azimuth': azim}
        """
        self.rendered_sats = []
        self.ax.clear()

        # 恢复符合暗色调主题的画板细节
        self._apply_base_style()

        # 如果没有数据，直接重绘空图表
        if not gsv_satellites:
            self.ax.set_xticks([])
            self.ax.set_xticklabels([])
            self.fig.canvas.draw_idle()
            return

        # 1. 过滤掉所有 snr 均为 0 的卫星
        active_sats = {}
        for (prefix, prn), sig_dict in gsv_satellites.items():
            valid_sig = {sid: snr for sid, snr in sig_dict.items() if snr > 0}
            if valid_sig:
                active_sats[(prefix, prn)] = valid_sig

        if not active_sats:
            self.ax.set_xticks([])
            self.ax.set_xticklabels([])
            self.fig.canvas.draw_idle()
            return

        # 2. 扁平化提取所有活跃的信号通道 [(prefix, prn, sid, val), ...]
        system_order = {
            'GPS': 0, 'QZSS': 1, 'BD': 2, 'GL': 3, 'GA': 4, 'IRNSS': 5, 'SBAS': 6
        }
        def channel_sort_key(item):
            prefix, prn, sid, val = item
            sys_idx = system_order.get(prefix, 99)
            return (sys_idx, prn, sid)

        channels = []
        for (prefix, prn), sig_dict in active_sats.items():
            for sid, val in sig_dict.items():
                channels.append((prefix, prn, sid, val))
        channels.sort(key=channel_sort_key)

        if not channels:
            self.ax.set_xticks([])
            self.ax.set_xticklabels([])
            self.fig.canvas.draw_idle()
            return

        # 3. 计算横坐标
        x_positions = []
        x_labels = []

        # 配色风格 (GPS=绿, 北斗=红, GLONASS=蓝, Galileo=青, QZSS/SBAS/IRNSS=绿/薄荷绿等类似配色)
        colors_map = {
            'GPS': {
                '1': '#22C55E',  # L1 C/A (亮绿)
                '2': '#16A34A',  # L1 P(Y) (中绿)
                '3': '#15803D',  # L1 M (深绿)
                '4': '#14532D',  # L2 P(Y) (极深绿)
                '5': '#4ADE80',  # L2C M (明绿)
                '6': '#86EFAC',  # L2C L (淡绿)
                '7': '#A7F3D0',  # L5 I (薄荷绿)
                '8': '#CCFBF1',  # L5 Q (青绿)
                'default': '#22C55E'
            },
            'QZSS': {
                '1': '#22C55E',  # L1 C/A
                '7': '#A7F3D0',  # L5 I
                'default': '#22C55E'
            },
            'SBAS': {
                '1': '#22C55E',  # L1 C/A
                '7': '#A7F3D0',  # L5 I
                '8': '#CCFBF1',  # L5 Q
                'default': '#22C55E'
            },
            'IRNSS': {
                '1': '#A7F3D0',  # L5 I
                'default': '#A7F3D0'
            },
            'BD': {
                '1': '#EF4444',  # B1I (亮红)
                '3': '#FB7185',  # B1C (玫瑰粉红)
                '5': '#D946EF',  # B2a (极光洋红)
                'B': '#DC2626',  # B2I (中红)
                '6': '#F87171',  # B2b (粉红)
                '7': '#FEE2E2',  # B2 A+B (极淡红)
                '8': '#991B1B',  # B3I (深红)
                'default': '#EF4444'
            },
            'GL': {
                '1': '#FBBF24',  # G1 C/A (亮金黄)
                '2': '#F59E0B',  # G1 P (琥珀黄)
                '3': '#D97706',  # G2 C/A (金黄棕)
                '4': '#92400E',  # G2 P (深古铜)
                'default': '#F59E0B'
            },
            'GA': {
                '1': '#22D3EE',  # E5a (浅青)
                '2': '#0EA5E9',  # E5b (中青/蓝)
                '3': '#06B6D4',  # E5 a+b (亮青)
                '4': '#0891B2',  # E6 A (中深青)
                '5': '#0E7490',  # E6 B (深青)
                '6': '#155E75',  # L1 A (极深青)
                '7': '#2563EB',  # E1 (深宝蓝)
                'default': '#06B6D4'
            }
        }

        from matplotlib.offsetbox import OffsetImage, AnnotationBbox
        import matplotlib.image as mpimg

        N = len(channels)
        draw_all_flags = N <= 60
        flagged_prefixes = set()

        for idx, (prefix, prn, sid, val) in enumerate(channels):
            x_c = idx + 0.5
            x_positions.append(x_c)

            sig_name = self.get_signal_name(prefix, sid)
            _, _, lbl_char = get_sat_info(prefix, prn)
            if lbl_char == 'S':
                prn_str = f"S{prn:03d}"
            else:
                prn_str = f"{lbl_char}{prn:02d}"

            x_labels.append(f"{sig_name}\n{prn_str}")

            c_dict = colors_map.get(prefix, {})
            color = c_dict.get(sid, c_dict.get('default', '#64748B'))

            # 默认若未收到过任何 GSA 语句，则全部视为在用以保持兼容；若已接收到 GSA，则以当前在用集合为准
            is_used = True
            if has_gsa and used_satellites is not None:
                is_used = (prefix, prn) in used_satellites

            bar_w = 0.98
            if is_used:
                self.ax.bar(x_c, val, width=bar_w, color=color, edgecolor='#0B1120', linewidth=0.5, align='center', alpha=1.0)
                self.ax.text(x_c, val + 0.8, f"{int(val)}", ha='center', va='bottom', color='#FFFFFF', fontsize=9, fontweight='semibold')
            else:
                self.ax.bar(x_c, val, width=bar_w, facecolor=color, edgecolor=color, linewidth=1.2, align='center', alpha=0.25)
                self.ax.text(x_c, val + 0.8, f"{int(val)}", ha='center', va='bottom', color='#94A3B8', fontsize=9, fontweight='semibold')

            # 绘制国旗图标
            should_draw_flag = draw_all_flags or prefix not in flagged_prefixes
            img = self.flag_images.get(prefix) if should_draw_flag and hasattr(self, 'flag_images') else None
            if img is not None:
                try:
                    imagebox = OffsetImage(img, zoom=1.0)
                    ab = AnnotationBbox(imagebox, (x_c, 0),
                                        xybox=(0, -10),
                                        xycoords=('data', 'axes fraction'),
                                        boxcoords="offset points",
                                        frameon=False)
                    self.ax.add_artist(ab)
                    flagged_prefixes.add(prefix)
                except Exception as e:
                    print(f"Error drawing flag: {e}")

            # 记录渲染星体用于鼠标悬停计算
            elev = None
            azim = None
            if sat_metadata is not None:
                meta = sat_metadata.get((prefix, prn), {})
                elev = meta.get('elevation')
                azim = meta.get('azimuth')
            self.rendered_sats.append({
                'center_x': x_c,
                'prefix': prefix,
                'prn': prn,
                'sig_id': sid,
                'snr': val,
                'elevation': elev,
                'azimuth': azim,
                'sig_dict': active_sats.get((prefix, prn), {sid: val})
            })

        # 4. 绘制分界虚线网格
        for x_grid in range(1, N):
            self.ax.axvline(x_grid, color='#1E293B', linestyle='--', linewidth=0.8, alpha=0.5)

        # 5. X轴属性与 Padding 调整以避让国旗图标
        self.ax.set_xlim(0, N)
        self.ax.set_xticks(x_positions)
        self.ax.tick_params(axis='x', colors='#94A3B8', labelsize=8, pad=18)
        self.ax.set_xticklabels(x_labels, ha='center', color='#F8FAFC', fontsize=8, fontweight='bold')

        if self._last_layout_channel_count != N:
            import warnings
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", category=UserWarning)
                self.fig.tight_layout()
            self._last_layout_channel_count = N
        self.fig.canvas.draw_idle()

    def on_mouse_move(self, event):
        from PySide6.QtGui import QCursor
        from PySide6.QtWidgets import QToolTip
        import time

        now = time.monotonic()
        if now - self._last_hover_time < 0.05:
            return
        self._last_hover_time = now

        if event.inaxes != self.ax or not hasattr(self, 'rendered_sats') or not self.rendered_sats:
            QToolTip.hideText()
            return

        x_mouse = event.xdata
        y_mouse = event.ydata

        if x_mouse is None or y_mouse is None:
            QToolTip.hideText()
            return

        idx = int(x_mouse)
        if idx < 0 or idx >= len(self.rendered_sats) or not (0 <= y_mouse <= 55):
            QToolTip.hideText()
            return

        hovered_sat = self.rendered_sats[idx]

        # 提取当前悬停卫星的数据
        prefix = hovered_sat['prefix']
        prn = hovered_sat['prn']
        elev = hovered_sat['elevation']
        azim = hovered_sat['azimuth']
        sig_dict = hovered_sat.get('sig_dict', {hovered_sat['sig_id']: hovered_sat['snr']})

        # 卫星标识前缀
        _, _, lbl_char = get_sat_info(prefix, prn)
        if lbl_char == 'S':
            sat_id_str = f"S{prn:03d}"
        else:
            sat_id_str = f"{lbl_char}{prn:02d}"

        # 仰角与方位角显示字符串
        elev_str = f"{elev}°" if elev is not None else "--"
        azim_str = f"{azim}°" if azim is not None else "--"

        # u-center 颜色映射配置，使气泡内频段字体颜色和图表中一致
        colors_map = {
            'GPS': {
                '1': '#22C55E', '2': '#16A34A', '3': '#15803D', '4': '#14532D', '5': '#4ADE80', '6': '#86EFAC', '7': '#A7F3D0', '8': '#CCFBF1', 'default': '#22C55E'
            },
            'QZSS': {
                '1': '#22C55E', '7': '#A7F3D0', 'default': '#22C55E'
            },
            'SBAS': {
                '1': '#22C55E', '7': '#A7F3D0', '8': '#CCFBF1', 'default': '#22C55E'
            },
            'IRNSS': {
                '1': '#A7F3D0', 'default': '#A7F3D0'
            },
            'BD': {
                '1': '#EF4444', '3': '#FCA5A5', '5': '#FECACA', 'B': '#DC2626', '6': '#F87171', '7': '#FEE2E2', '8': '#991B1B', 'default': '#EF4444'
            },
            'GL': {
                '1': '#38BDF8', '2': '#0EA5E9', '3': '#1D4ED8', '4': '#1E3A8A', 'default': '#38BDF8'
            },
            'GA': {
                '1': '#22D3EE', '2': '#0EA5E9', '3': '#06B6D4', '4': '#0891B2', '5': '#0E7490', '6': '#155E75', '7': '#06B6D4', 'default': '#06B6D4'
            }
        }

        def get_signal_name(pref, sig_id):
            if pref == 'GPS':
                if sig_id == '1': return 'L1 C/A'
                elif sig_id == '2': return 'L1 P(Y)'
                elif sig_id == '3': return 'L1 M'
                elif sig_id == '4': return 'L2 P(Y)'
                elif sig_id == '5': return 'L2C M'
                elif sig_id == '6': return 'L2C L'
                elif sig_id == '7': return 'L5 I'
                elif sig_id == '8': return 'L5 Q'
                return f"L{sig_id}"
            elif pref == 'BD':
                if sig_id == '1': return 'B1I'
                elif sig_id == '3': return 'B1C'
                elif sig_id == '5': return 'B2a'
                elif sig_id == 'B': return 'B2I'
                elif sig_id == '6': return 'B2b'
                elif sig_id == '7': return 'B2 A+B'
                elif sig_id == '8': return 'B3I'
                return f"B{sig_id}"
            elif pref == 'GL':
                if sig_id == '1': return 'G1 C/A'
                elif sig_id == '2': return 'G1 P'
                elif sig_id == '3': return 'G2 C/A'
                elif sig_id == '4': return 'G2 P'
                return f"G{sig_id}"
            elif pref == 'GA':
                if sig_id == '1': return 'E5a'
                elif sig_id == '2': return 'E5b'
                elif sig_id == '3': return 'E5 a+b'
                elif sig_id == '4': return 'E6 A'
                elif sig_id == '5': return 'E6 B'
                elif sig_id == '6': return 'L1 A'
                elif sig_id == '7': return 'E1'
                return f"E{sig_id}"
            elif pref == 'QZSS':
                if sig_id == '1': return 'L1 C/A'
                elif sig_id == '7': return 'L5 I'
                return f"Q{sig_id}"
            elif pref == 'IRNSS':
                if sig_id == '1': return 'L5 I'
                return f"I{sig_id}"
            elif pref == 'SBAS':
                if sig_id == '1': return 'L1 C/A'
                elif sig_id == '7': return 'L5 I'
                elif sig_id == '8': return 'L5 Q'
                return f"S{sig_id}"
            return sig_id

        # 构造精美的 HTML 表格气泡，完全匹配模组厂商参考图设计
        html = f"""
        <div style="background-color: #172033; color: #F8FAFC; border: 1px solid #334155; border-radius: 6px; padding: 8px; font-family: 'Consolas', 'Segoe UI', 'Microsoft YaHei', sans-serif;">
            <div style="font-weight: bold; font-size: 13px; color: #38BDF8; margin-bottom: 6px; border-bottom: 1px solid #1E293B; padding-bottom: 4px;">
                {sat_id_str} &nbsp; E:{elev_str} &nbsp; A:{azim_str}
            </div>
            <table style="border-collapse: collapse; font-size: 11px; width: 100%;">
                <tr style="border-bottom: 1px solid #1E293B; color: #94A3B8; font-weight: bold;">
                    <th style="text-align: left; padding: 2px 12px 2px 0;">signal</th>
                    <th style="text-align: right; padding: 2px 0 2px 12px;">CN0</th>
                </tr>
        """

        sids = sorted(list(sig_dict.keys()))
        for sid in sids:
            val = sig_dict[sid]
            sig_name = get_signal_name(prefix, sid)

            c_dict = colors_map.get(prefix, {})
            color = c_dict.get(sid, c_dict.get('default', '#64748B'))

            html += f"""
                <tr style="color: {color};">
                    <td style="text-align: left; padding: 2px 12px 2px 0; font-weight: bold;">{sig_name}</td>
                    <td style="text-align: right; padding: 2px 0 2px 12px; font-weight: bold;">{int(val)}</td>
                </tr>
            """

        html += """
            </table>
        </div>
        """

        QToolTip.showText(QCursor.pos(), html, self)

class LogParserThread(QThread):
    progress_updated = Signal(int)
    finished_parsing = Signal(object)
    error_occurred = Signal(str)

    def __init__(self, filepath, leap_secs, strict_checksum=False, parent=None):
        super().__init__(parent)
        self.filepath = filepath
        self.leap_secs = leap_secs
        self.strict_checksum = strict_checksum
        self.is_cancelled = False

    def run(self):
        try:
            file_size = os.path.getsize(self.filepath)
            processed_size = 0
            file_epochs = []
            sentence_types = {}
            first_time_sec = None
            last_time_sec = None
            first_time_str = ''
            last_time_str = ''
            gga_status_events = []
            gsa_status_events = []
            last_time_sec_for_gsa = None
            last_emit_progress = -1

            with open(self.filepath, 'r', encoding='utf-8', errors='replace') as f:
                for line in f:
                    if self.is_cancelled:
                        break

                    line_len = len(line.encode('utf-8', errors='replace'))
                    processed_size += line_len

                    if not line.startswith('$'):
                        continue

                    comma_idx = line.find(',')
                    if comma_idx != -1:
                        stype = line[:comma_idx]
                        sentence_types[stype] = sentence_types.get(stype, 0) + 1

                    epoch = parse_log_line(line, self.leap_secs, self.strict_checksum)
                    if epoch:
                        if epoch['type'] in ['GGA', 'POGOS', 'PODRS', 'RMC']:
                            epoch['file_id'] = self.filepath
                            file_epochs.append(epoch)
                            last_time_sec_for_gsa = epoch['utc_time_sec']
                            if first_time_sec is None:
                                first_time_sec = epoch['utc_time_sec']
                                first_time_str = epoch['time_str']
                            last_time_sec = epoch['utc_time_sec']
                            last_time_str = epoch['time_str']

                            if epoch['type'] == 'GGA':
                                gga_status_events.append((epoch['utc_time_sec'], {
                                    'quality': epoch['quality'],
                                    'num_sats': epoch['num_sats'],
                                    'hdop': epoch['hdop']
                                }))
                        elif epoch['type'] == 'GSA' and last_time_sec_for_gsa is not None:
                            gsa_status_events.append((last_time_sec_for_gsa, {
                                'vdop': epoch['vdop'],
                                'pdop': epoch['pdop']
                            }))

                    if processed_size % (1024 * 100) < line_len:  # Update roughly every 100KB
                        progress = int((processed_size / file_size) * 100) if file_size > 0 else 0
                        if progress > last_emit_progress:
                            self.progress_updated.emit(progress)
                            last_emit_progress = progress

            if file_epochs:
                from gnss_parser import unwrap_times
                unwrapped = unwrap_times([ep['utc_time_sec'] for ep in file_epochs])
                for i, ep in enumerate(file_epochs):
                    ep['utc_time_sec'] = unwrapped[i]
                first_time_sec = file_epochs[0]['utc_time_sec']
                last_time_sec = file_epochs[-1]['utc_time_sec']

                def unwrap_status_events(events):
                    if not events:
                        return []
                    times = unwrap_times([item[0] for item in events])
                    return [(times[i], events[i][1]) for i in range(len(events))]

                gga_status_events = unwrap_status_events(gga_status_events)
                gsa_status_events = unwrap_status_events(gsa_status_events)

                import bisect
                gga_status_events.sort(key=lambda item: item[0])
                gsa_status_events.sort(key=lambda item: item[0])
                gga_status_times = [item[0] for item in gga_status_events]
                gsa_status_times = [item[0] for item in gsa_status_events]

                def nearest_status(events, event_times, target_time, max_delta=1.0):
                    if not events:
                        return None
                    best_fields = None
                    best_delta = None
                    pos = bisect.bisect_left(event_times, target_time)
                    for candidate in (pos - 1, pos):
                        if candidate < 0 or candidate >= len(events):
                            continue
                        event_time, fields = events[candidate]
                        delta = abs(event_time - target_time)
                        if delta <= max_delta and (best_delta is None or delta < best_delta):
                            best_delta = delta
                            best_fields = fields
                    return best_fields

                for ep in file_epochs:
                    if ep['type'] in ['POGOS', 'PODRS']:
                        gga_fields = nearest_status(gga_status_events, gga_status_times, ep['utc_time_sec'])
                        if gga_fields:
                            ep.update(gga_fields)

                    gsa_fields = nearest_status(gsa_status_events, gsa_status_times, ep['utc_time_sec'])
                    if gsa_fields:
                        ep.update(gsa_fields)

            self.progress_updated.emit(100)
            result = {
                'file_epochs': file_epochs,
                'first_time_sec': first_time_sec,
                'last_time_sec': last_time_sec,
                'first_time_str': first_time_str,
                'last_time_str': last_time_str,
                'sentence_types': sentence_types
            }
            self.finished_parsing.emit(result)
        except Exception as e:
            self.error_occurred.emit(str(e))

def create_switch_icons():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    res_dir = os.path.join(base_dir, "resources")
    # Make sure resources directory exists
    os.makedirs(res_dir, exist_ok=True)
    # switch_off.png (36x20)
    pix_off = QPixmap(36, 20)
    pix_off.fill(QColor(0, 0, 0, 0))
    painter = QPainter(pix_off)
    painter.setRenderHint(QPainter.Antialiasing)
    # Draw capsule background
    painter.setPen(QPen(QColor("#475569"), 1.5))
    painter.setBrush(QBrush(QColor("#1E293B")))
    painter.drawRoundedRect(QRectF(1, 1, 34, 18), 9, 9)
    # Draw knob
    painter.setPen(Qt.NoPen)
    painter.setBrush(QBrush(QColor("#64748B")))
    painter.drawEllipse(QRectF(3, 3, 14, 14))
    painter.end()
    try:
        pix_off.save(os.path.join(res_dir, "switch_off.png"), "PNG")
    except Exception:
        pass

    # switch_on.png (36x20)
    pix_on = QPixmap(36, 20)
    pix_on.fill(QColor(0, 0, 0, 0))
    painter = QPainter(pix_on)
    painter.setRenderHint(QPainter.Antialiasing)
    # Draw capsule background
    painter.setPen(QPen(QColor("#0EA5E9"), 1.5))
    painter.setBrush(QBrush(QColor("#0F172A")))
    painter.drawRoundedRect(QRectF(1, 1, 34, 18), 9, 9)
    # Draw knob (glowing neon blue)
    painter.setPen(Qt.NoPen)
    painter.setBrush(QBrush(QColor("#38BDF8")))
    painter.drawEllipse(QRectF(19, 3, 14, 14))
    painter.end()
    try:
        pix_on.save(os.path.join(res_dir, "switch_on.png"), "PNG")
    except Exception:
        pass

def create_app_icon():
    # Create a 256x256 high-resolution icon
    pixmap = QPixmap(256, 256)
    pixmap.fill(QColor("#0B1120")) # Dark background matching QSS QMainWindow

    painter = QPainter(pixmap)
    painter.setRenderHint(QPainter.Antialiasing)

    # Draw a stylish target/orbit (satellite positioning theme)
    # Circle 1 (Outer orbit)
    pen = QPen(QColor("#1E293B"), 4)
    painter.setPen(pen)
    painter.drawEllipse(QRectF(28, 28, 200, 200))

    # Circle 2 (Inner orbit)
    pen.setColor(QColor("#334155"))
    pen.setWidth(2)
    painter.setPen(pen)
    painter.drawEllipse(QRectF(56, 56, 144, 144))

    # Crosshair lines
    painter.drawLine(QPointF(128, 10), QPointF(128, 246))
    painter.drawLine(QPointF(10, 128), QPointF(246, 128))

    # Draw a stylish "V" (for VCOM) in the center
    font = QFont("Outfit", 96, QFont.Bold)
    painter.setFont(font)
    painter.setPen(QColor("#60A5FA")) # Electric blue
    painter.drawText(QRectF(0, 0, 256, 240), Qt.AlignCenter, "V")

    # Draw a glowing center dot
    painter.setPen(Qt.NoPen)
    painter.setBrush(QBrush(QColor("#EF4444"))) # Red target dot
    painter.drawEllipse(QPointF(128, 128), 12, 12)

    painter.end()

    # Save physical icon files to the resources directory for PyInstaller
    base_dir = os.path.dirname(os.path.abspath(__file__))
    res_dir = os.path.join(base_dir, "resources")
    os.makedirs(res_dir, exist_ok=True)
    if not os.path.exists(os.path.join(res_dir, "icon.png")):
        try:
            pixmap.save(os.path.join(res_dir, "icon.png"), "PNG")
        except Exception:
            pass
    if not os.path.exists(os.path.join(res_dir, "icon.ico")):
        try:
            pixmap.save(os.path.join(res_dir, "icon.ico"), "ICO")
        except Exception:
            pass

    # Also generate modern switch toggle PNGs for checkboxes
    create_switch_icons()

    return QIcon(pixmap)

def create_refresh_icon(color_hex):
    from PySide6.QtGui import QPixmap, QPainter, QPen, QBrush, QColor, QIcon, QPolygon
    from PySide6.QtCore import Qt, QPoint, QRectF
    pixmap = QPixmap(32, 32)
    pixmap.fill(Qt.transparent)
    painter = QPainter(pixmap)
    painter.setRenderHint(QPainter.Antialiasing)

    pen = QPen(QColor(color_hex))
    pen.setWidth(3)
    pen.setCapStyle(Qt.RoundCap)
    painter.setPen(pen)

    # Draw arc
    painter.drawArc(QRectF(6, 6, 20, 20), 960, -4800)

    # Draw arrowhead
    brush = QBrush(QColor(color_hex))
    painter.setBrush(brush)
    painter.setPen(Qt.NoPen)
    points = QPolygon([QPoint(23, 6), QPoint(23, 16), QPoint(14, 11)])
    painter.drawPolygon(points)

    painter.end()
    return QIcon(pixmap)

def is_system_dark_mode():
    """读取 Windows 注册表以精准获取操作系统真实的深色/浅色模式"""
    import platform
    if platform.system() == "Windows":
        try:
            import winreg
            registry = winreg.ConnectRegistry(None, winreg.HKEY_CURRENT_USER)
            key = winreg.OpenKey(registry, r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize")
            value, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
            return value == 0  # 0 表示深色模式 (AppsUseLightTheme = 0)
        except Exception:
            pass
    # 针对非 Windows 系统或读取失败的备用方案
    try:
        from PySide6.QtWidgets import QApplication
        from PySide6.QtGui import QPalette
        app = QApplication.instance()
        if app:
            bg_color = app.palette().color(QPalette.ColorRole.Window)
            return bg_color.lightness() < 128
    except Exception:
        pass
    return False

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("VCOM定位精度分析工具")
        self.setWindowIcon(create_app_icon())
        self.resize(1180, 720)
        self.setStyleSheet(QSS_STYLE)

        self.app_config = {}
        self.parser_thread = None
        self.dynamic_parser_thread = None

        # 数据状态管理
        self.raw_lines = []
        self.parsed_epochs = []
        self.file_epochs_map = {}  # 字典映射 file_id -> epochs list 以加速筛选
        self.sentence_types = {}
        self.time_range = {'start': 0, 'end': 0}

        self.truth_mode = 'auto' # 'auto' 或 'manual'
        self.truth = {'lat': 0.0, 'lon': 0.0, 'alt': 0.0}

        self.segments = []
        self.segment_counter = 0
        self.default_colors = [
            '#2563EB', '#EF4444', '#16A34A', '#D97706', '#9333EA', '#0D9488',
            '#F97316', '#EC4899', '#78350F', '#84CC16', '#4338CA', '#4B5563'
        ]
        self.first_time_str = ''
        self.last_time_str = ''
        self.time_zone = 'UTC'
        self.show_absolute_alt = False  # 是否显示高程绝对值误差
        self.show_raw_alt = False  # 是否显示高程绝对物理值 (而不是误差绝对值)
        self.show_extrema = True  # 是否显示最值标注
        self.x_axis_mode = '历元数'  # 是否使用时间轴对齐 X 轴

        # 初始化串口组件与录制状态
        self.serial_port = QSerialPort(self)
        self.serial_port.readyRead.connect(self.handle_serial_read)
        self.serial_buffer = BKStreamParser()
        self.record_file = None
        self.record_filepath = None
        self.record_error_reported = False
        self._is_programmatic_scroll = False
        self.realtime_raw_epochs = deque(maxlen=6000)
        self.latest_quality = 0
        self.latest_num_sats = 0
        self.latest_hdop = 1.0
        self.latest_pdop = 1.0

        # 版本查询状态
        self.version_lines = []
        self.version_timer = None
        self.waiting_for_version = False

        # 实时数据超时清空定时器
        self.realtime_timeout_timer = QTimer(self)
        self.realtime_timeout_timer.setSingleShot(True)
        self.realtime_timeout_timer.timeout.connect(self.reset_live_status_ui)

        # 卫星载噪比数据缓存 (prefix, prn) -> {signal_id: snr}
        self.gsv_satellites = {}
        self.used_satellites = set()  # 在用(定位解算)卫星集合，存储 (prefix, prn) 元组
        self.has_received_gsa = False  # 标记是否接收到过 GSA 语句
        self.sat_metadata = {}  # 卫星仰角方位角元数据缓存 (prefix, prn) -> {'elevation': elev, 'azimuth': azim}
        self._last_cno_snapshot = None

        # 载噪比柱状图限频定时器
        self.cno_refresh_timer = QTimer(self)
        self.cno_refresh_timer.timeout.connect(self.update_cno_chart)

        # 串口回放相关状态
        self.replay_blocks = []
        self.replay_filepath = None
        self.replay_index = 0
        self.replay_snapshot_interval = 500
        self.replay_snapshots = {}
        self.is_replaying = False
        self.replay_timer = QTimer(self)
        self.replay_timer.timeout.connect(self.replay_tick)
        self.is_slider_dragging = False
        self.last_recompute_time = 0.0
        self.last_live_recompute_time = 0.0
        self.replay_start_time = 0.0
        self.replay_start_index = 0

        # 初始化 UI
        self.init_ui()
        self.setAcceptDrops(True)
        self.load_config()
        self.refresh_serial_ports()
        self.pending_parse_queue = []
        self.total_queue_count = 0

    def get_leap_seconds(self):
        try:
            return int(self.txt_leap.text())
        except ValueError:
            return 18

    def init_ui(self):
        # 1. 顶部菜单栏
        self.menu_bar = QMenuBar(self)
        self.setMenuBar(self.menu_bar)
        file_menu = self.menu_bar.addMenu("文件")
        file_menu.addAction("打开日志", self.on_import_clicked)
        file_menu.addAction("退出", self.close)
        settings_menu = self.menu_bar.addMenu("设置")
        settings_menu.addAction("首选项...", self.on_settings_clicked)
        self.menu_bar.addMenu("帮助")

        # 主中央分割窗
        self.main_splitter = QSplitter(Qt.Horizontal)
        self.main_splitter.setHandleWidth(7)
        self.setCentralWidget(self.main_splitter)

        # 2. 左侧工作区 (包含 Tab 嵌套画布及工具栏)
        self.left_widget = QWidget()
        left_layout = QVBoxLayout(self.left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(0)

        self.tab_widget = QTabWidget()

        # 定义白色图表卡片内嵌工具栏样式，采用浅色背景以契合白色图表卡片，确保深色(黑色)图标在任何系统模式下都清晰可见
        TOOLBAR_STYLE = """
        QToolBar {
            background-color: #F8FAFC;
            border: none;
            border-bottom: 1px solid #E2E8F0;
            border-top-left-radius: 8px;
            border-top-right-radius: 8px;
        }
        QToolBar QToolButton {
            background-color: transparent;
            border: none;
            border-radius: 4px;
            padding: 4px;
            margin: 2px;
        }
        QToolBar QToolButton:hover {
            background-color: #E2E8F0;
        }
        QToolBar QToolButton:pressed {
            background-color: #CBD5E1;
        }
        QToolBar QToolButton:disabled {
            background-color: transparent;
            opacity: 0.5;
        }
        QToolBar QLabel {
            color: #475569;
            font-size: 11px;
            font-weight: bold;
            padding-left: 10px;
            background-color: transparent;
        }
        QToolBar::separator {
            width: 1px;
            background-color: #E2E8F0;
            margin-top: 4px;
            margin-bottom: 4px;
            margin-left: 6px;
            margin-right: 6px;
        }
        """


        # A. 靶心图容器页
        self.tab_scatter = QWidget()
        layout_scatter = QVBoxLayout(self.tab_scatter)
        layout_scatter.setContentsMargins(12, 12, 12, 12)
        layout_scatter.setSpacing(0)

        # 创建白卡容器
        self.card_scatter = QWidget()
        self.card_scatter.setStyleSheet("background-color: #FFFFFF; border: 1px solid #334155; border-radius: 8px;")
        card_layout_scatter = QVBoxLayout(self.card_scatter)
        card_layout_scatter.setContentsMargins(0, 0, 0, 0)
        card_layout_scatter.setSpacing(0)

        self.canvas_scatter = PlotWidget(self.card_scatter)
        self.canvas_scatter.setStyleSheet("background-color: #FFFFFF; border-bottom-left-radius: 8px; border-bottom-right-radius: 8px;")
        self.toolbar_scatter = NavigationToolbar(self.canvas_scatter, self.card_scatter)
        self.toolbar_scatter.setStyleSheet(TOOLBAR_STYLE)

        card_layout_scatter.addWidget(self.toolbar_scatter)
        card_layout_scatter.addWidget(self.canvas_scatter)
        layout_scatter.addWidget(self.card_scatter)

        # B. 定位质量图容器页
        self.tab_status = QWidget()
        layout_status = QVBoxLayout(self.tab_status)
        layout_status.setContentsMargins(12, 12, 12, 12)
        layout_status.setSpacing(0)

        # 创建白卡容器
        self.card_status = QWidget()
        self.card_status.setStyleSheet("background-color: #FFFFFF; border: 1px solid #334155; border-radius: 8px;")
        card_layout_status = QVBoxLayout(self.card_status)
        card_layout_status.setContentsMargins(0, 0, 0, 0)
        card_layout_status.setSpacing(0)

        self.canvas_status = PlotWidget(self.card_status)
        self.canvas_status.setStyleSheet("background-color: #FFFFFF; border-bottom-left-radius: 8px; border-bottom-right-radius: 8px;")
        self.toolbar_status = NavigationToolbar(self.canvas_status, self.card_status)
        self.toolbar_status.setStyleSheet(TOOLBAR_STYLE)

        card_layout_status.addWidget(self.toolbar_status)
        card_layout_status.addWidget(self.canvas_status)
        layout_status.addWidget(self.card_status)

        # C. 水平误差历元图容器页
        self.tab_epoch_h = QWidget()
        layout_epoch_h = QVBoxLayout(self.tab_epoch_h)
        layout_epoch_h.setContentsMargins(12, 12, 12, 12)
        layout_epoch_h.setSpacing(0)

        # 创建白卡容器
        self.card_epoch_h = QWidget()
        self.card_epoch_h.setStyleSheet("background-color: #FFFFFF; border: 1px solid #334155; border-radius: 8px;")
        card_layout_epoch_h = QVBoxLayout(self.card_epoch_h)
        card_layout_epoch_h.setContentsMargins(0, 0, 0, 0)
        card_layout_epoch_h.setSpacing(0)

        self.canvas_epoch_h = PlotWidget(self.card_epoch_h)
        self.canvas_epoch_h.setStyleSheet("background-color: #FFFFFF; border-bottom-left-radius: 8px; border-bottom-right-radius: 8px;")
        self.toolbar_epoch_h = NavigationToolbar(self.canvas_epoch_h, self.card_epoch_h)
        self.toolbar_epoch_h.setStyleSheet(TOOLBAR_STYLE)

        card_layout_epoch_h.addWidget(self.toolbar_epoch_h)
        card_layout_epoch_h.addWidget(self.canvas_epoch_h)
        layout_epoch_h.addWidget(self.card_epoch_h)

        # D. 高程误差历元图容器页
        self.tab_epoch_v = QWidget()
        layout_epoch_v = QVBoxLayout(self.tab_epoch_v)
        layout_epoch_v.setContentsMargins(12, 12, 12, 12)
        layout_epoch_v.setSpacing(0)

        # 创建白卡容器
        self.card_epoch_v = QWidget()
        self.card_epoch_v.setStyleSheet("background-color: #FFFFFF; border: 1px solid #334155; border-radius: 8px;")
        card_layout_epoch_v = QVBoxLayout(self.card_epoch_v)
        card_layout_epoch_v.setContentsMargins(0, 0, 0, 0)
        card_layout_epoch_v.setSpacing(0)

        self.canvas_epoch_v = PlotWidget(self.card_epoch_v)
        self.canvas_epoch_v.setStyleSheet("background-color: #FFFFFF; border-bottom-left-radius: 8px; border-bottom-right-radius: 8px;")
        self.toolbar_epoch_v = NavigationToolbar(self.canvas_epoch_v, self.card_epoch_v)
        self.toolbar_epoch_v.setStyleSheet(TOOLBAR_STYLE)

        card_layout_epoch_v.addWidget(self.toolbar_epoch_v)
        card_layout_epoch_v.addWidget(self.canvas_epoch_v)
        layout_epoch_v.addWidget(self.card_epoch_v)

        # E. 绝对轨迹投影页
        self.tab_trajectory = QWidget()
        layout_trajectory = QVBoxLayout(self.tab_trajectory)
        layout_trajectory.setContentsMargins(12, 12, 12, 12)
        layout_trajectory.setSpacing(0)

        self.card_trajectory = QWidget()
        self.card_trajectory.setStyleSheet("background-color: #FFFFFF; border: 1px solid #334155; border-radius: 8px;")
        card_layout_trajectory = QVBoxLayout(self.card_trajectory)
        card_layout_trajectory.setContentsMargins(0, 0, 0, 0)
        card_layout_trajectory.setSpacing(0)

        self.canvas_trajectory = PlotWidget(self.card_trajectory)
        self.canvas_trajectory.setStyleSheet("background-color: #FFFFFF; border-bottom-left-radius: 8px; border-bottom-right-radius: 8px;")
        self.toolbar_trajectory = NavigationToolbar(self.canvas_trajectory, self.card_trajectory)
        self.toolbar_trajectory.setStyleSheet(TOOLBAR_STYLE)

        card_layout_trajectory.addWidget(self.toolbar_trajectory)
        card_layout_trajectory.addWidget(self.canvas_trajectory)
        layout_trajectory.addWidget(self.card_trajectory)

        # F. 实时串口页
        self.tab_serial = QWidget()
        layout_serial = QHBoxLayout(self.tab_serial)
        layout_serial.setContentsMargins(8, 8, 8, 8)
        layout_serial.setSpacing(0)

        # 预先定义数据解析卡片 QTabWidget，并在之后做上下垂直堆叠
        self.dashboard_tab = QTabWidget()
        self.dashboard_tab.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #1E293B;
                background-color: #172033;
                border-radius: 8px;
            }
            QTabBar::tab {
                background-color: #0B1120;
                color: #94A3B8;
                padding: 6px 12px;
                font-size: 11px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                border: 1px solid #1E293B;
                margin-right: 4px;
            }
            QTabBar::tab:selected {
                background-color: #172033;
                color: #38BDF8;
                border-bottom: 2px solid #38BDF8;
            }
        """)

        # 主垂直分割条 (上层为水平切分，下层为载噪比监视器)
        self.serial_vertical_splitter = QSplitter(Qt.Vertical)
        self.serial_vertical_splitter.setHandleWidth(7)
        layout_serial.addWidget(self.serial_vertical_splitter)

        # 上层水平分割条 (左侧堆叠栏 + 右侧终端)
        self.serial_upper_splitter = QSplitter(Qt.Horizontal)
        self.serial_upper_splitter.setHandleWidth(7)
        self.serial_vertical_splitter.addWidget(self.serial_upper_splitter)

        # F.1. 左侧控制和指令面板
        self.serial_left_panel = QWidget()
        serial_left_layout = QVBoxLayout(self.serial_left_panel)
        serial_left_layout.setContentsMargins(4, 4, 4, 4)
        serial_left_layout.setSpacing(10)

        # 串口与回放配置 GroupBox
        self.group_serial_ctrl = QGroupBox("串口与回放配置")
        outer_layout = QVBoxLayout(self.group_serial_ctrl)
        outer_layout.setContentsMargins(4, 12, 4, 4)
        outer_layout.setSpacing(0)

        self.mode_tab = QTabWidget()
        self.mode_tab.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #1E293B;
                background-color: #172033;
                border-radius: 6px;
            }
            QTabBar::tab {
                background-color: #0B1120;
                color: #94A3B8;
                padding: 6px 12px;
                font-size: 11px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                border: 1px solid #1E293B;
                margin-right: 4px;
            }
            QTabBar::tab:selected {
                background-color: #172033;
                color: #38BDF8;
                border-bottom: 1px solid #172033;
            }
        """)

        # --- Tab 1: 串口连接 ---
        self.tab_serial_conn = QWidget()
        ctrl_layout = QGridLayout(self.tab_serial_conn)
        ctrl_layout.setContentsMargins(6, 10, 6, 6)
        ctrl_layout.setSpacing(6)

        ctrl_layout.addWidget(QLabel("串口选择:"), 0, 0)
        self.cmb_port = QComboBox()
        self.cmb_port.setFixedHeight(28)
        combobox_font = self.cmb_port.font()
        combobox_font.setPointSize(10)
        self.cmb_port.setFont(combobox_font)
        ctrl_layout.addWidget(self.cmb_port, 0, 1)

        self.btn_port_refresh = QPushButton()
        from PySide6.QtCore import QSize
        self.btn_port_refresh.setIcon(create_refresh_icon("#38BDF8"))
        self.btn_port_refresh.setIconSize(QSize(14, 14))
        self.btn_port_refresh.setFixedSize(26, 28)
        self.btn_port_refresh.setToolTip("刷新可用串口")
        self.btn_port_refresh.setStyleSheet("""
            QPushButton {
                background-color: rgba(56, 189, 248, 0.15);
                border: 1px solid #334155;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: rgba(56, 189, 248, 0.25);
                border-color: #38BDF8;
                color: #FFFFFF;
                font-size: 16px;
            }
            QPushButton:pressed {
                background-color: rgba(56, 189, 248, 0.35);
                border-color: #0EA5E9;
                font-size: 16px;
            }
        """)
        self.btn_port_refresh.clicked.connect(self.refresh_serial_ports)
        ctrl_layout.addWidget(self.btn_port_refresh, 0, 2)

        ctrl_layout.addWidget(QLabel("波特率:"), 1, 0)
        self.cmb_baud = QComboBox()
        self.cmb_baud.setFixedHeight(28)
        combobox_font = self.cmb_baud.font()
        combobox_font.setPointSize(10)
        self.cmb_baud.setFont(combobox_font)
        self.cmb_baud.addItems(["9600", "115200", "230400", "460800", "921600", "2000000"])
        self.cmb_baud.setCurrentText("115200")
        ctrl_layout.addWidget(self.cmb_baud, 1, 1, 1, 2)

        self.btn_serial_connect = QPushButton("打开串口")
        self.btn_serial_connect.setFixedHeight(32)
        self.btn_serial_connect.setStyleSheet("""
            QPushButton {
                background-color: rgba(16, 185, 129, 0.15);
                border: 1px solid #10B981;
                color: #10B981;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: rgba(16, 185, 129, 0.25);
                color: #FFFFFF;
                font-size: 12px;
            }
            QPushButton:pressed {
                background-color: rgba(16, 185, 129, 0.35);
                font-size: 12px;
            }
        """)
        self.btn_serial_connect.clicked.connect(self.toggle_serial_connection)
        ctrl_layout.addWidget(self.btn_serial_connect, 2, 0, 1, 3)

        row_display_ctrl = QHBoxLayout()
        row_display_ctrl.setSpacing(6)
        self.cb_hex = QCheckBox("Hex显示")
        self.cb_hex.setStyleSheet("color:#94A3B8; font-size:11px;")
        self.cb_scroll = QCheckBox("自动滚动")
        self.cb_scroll.setStyleSheet("color:#94A3B8; font-size:11px;")
        self.cb_scroll.setChecked(True)
        self.cb_scroll.stateChanged.connect(self.on_scroll_checkbox_changed)
        row_display_ctrl.addWidget(self.cb_hex)
        row_display_ctrl.addWidget(self.cb_scroll)
        ctrl_layout.addLayout(row_display_ctrl, 3, 0, 1, 3)

        row_record = QHBoxLayout()
        row_record.setSpacing(6)
        self.cb_record = QCheckBox("录制原始数据")
        self.cb_record.setStyleSheet("color:#38BDF8; font-size:11px;")
        self.cb_record.stateChanged.connect(self.on_record_state_changed)
        row_record.addWidget(self.cb_record)
        ctrl_layout.addLayout(row_record, 4, 0, 1, 3)

        self.btn_clear_console = QPushButton("清空接收区")
        self.btn_clear_console.setFixedHeight(26)
        self.btn_clear_console.clicked.connect(self.clear_serial_console)
        ctrl_layout.addWidget(self.btn_clear_console, 5, 0, 1, 3)

        self.mode_tab.addTab(self.tab_serial_conn, "串口连接")

        # --- Tab 2: 日志回放 ---
        self.tab_replay = QWidget()
        replay_layout = QVBoxLayout(self.tab_replay)
        replay_layout.setContentsMargins(6, 10, 6, 6)
        replay_layout.setSpacing(8)

        row_file = QHBoxLayout()
        row_file.setSpacing(4)
        self.txt_replay_file = QLineEdit()
        self.txt_replay_file.setReadOnly(True)
        self.txt_replay_file.setFixedHeight(26)
        self.txt_replay_file.setPlaceholderText("请选择日志文件...")
        self.txt_replay_file.setStyleSheet("""
            QLineEdit {
                background-color: #0B1120;
                color: #94A3B8;
                border: 1px solid #1E293B;
                border-radius: 4px;
                padding-left: 4px;
                font-size: 11px;
            }
        """)
        self.btn_replay_browse = QPushButton("浏览")
        self.btn_replay_browse.setFixedSize(46, 26)
        self.btn_replay_browse.setStyleSheet("""
            QPushButton {
                background-color: #1E293B;
                color: #F8FAFC;
                border: 1px solid #334155;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #334155;
            }
        """)
        self.btn_replay_browse.clicked.connect(self.on_replay_browse)

        self.btn_replay_clear = QPushButton("清除")
        self.btn_replay_clear.setFixedSize(46, 26)
        self.btn_replay_clear.setEnabled(False)
        self.btn_replay_clear.setStyleSheet("""
            QPushButton {
                background-color: #1E293B;
                color: #EF4444;
                border: 1px solid #EF4444;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: rgba(239, 68, 68, 0.15);
            }
            QPushButton:disabled {
                border-color: #334155;
                color: #475569;
            }
        """)
        self.btn_replay_clear.clicked.connect(self.clear_replay_data)

        row_file.addWidget(self.txt_replay_file)
        row_file.addWidget(self.btn_replay_browse)
        row_file.addWidget(self.btn_replay_clear)
        replay_layout.addLayout(row_file)

        row_ctrl = QHBoxLayout()
        row_ctrl.setSpacing(6)

        self.btn_replay_play = QPushButton("播放")
        self.btn_replay_play.setFixedSize(56, 28)
        self.btn_replay_play.setEnabled(False)
        self.btn_replay_play.setStyleSheet("""
            QPushButton {
                background-color: rgba(56, 189, 248, 0.15);
                border: 1px solid #38BDF8;
                color: #38BDF8;
                font-weight: bold;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: rgba(56, 189, 248, 0.25);
                color: #FFFFFF;
            }
            QPushButton:disabled {
                background-color: transparent;
                border-color: #334155;
                color: #475569;
            }
        """)
        self.btn_replay_play.clicked.connect(self.toggle_replay_playback)

        self.btn_replay_stop = QPushButton("停止")
        self.btn_replay_stop.setFixedSize(56, 28)
        self.btn_replay_stop.setEnabled(False)
        self.btn_replay_stop.setStyleSheet("""
            QPushButton {
                background-color: rgba(248, 113, 113, 0.15);
                border: 1px solid #F87171;
                color: #F87171;
                font-weight: bold;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: rgba(248, 113, 113, 0.25);
                color: #FFFFFF;
            }
            QPushButton:disabled {
                background-color: transparent;
                border-color: #334155;
                color: #475569;
            }
        """)
        self.btn_replay_stop.clicked.connect(self.stop_replay)

        lbl_speed = QLabel("倍速:")
        lbl_speed.setStyleSheet("color: #94A3B8; font-size: 11px;")

        self.cmb_replay_speed = QComboBox()
        self.cmb_replay_speed.setFixedHeight(28)
        self.cmb_replay_speed.setFixedWidth(64)
        self.cmb_replay_speed.addItems(["0.5x", "1.0x", "2.0x", "5.0x", "10.0x"])
        self.cmb_replay_speed.setCurrentText("1.0x")
        self.cmb_replay_speed.currentTextChanged.connect(self.on_replay_speed_changed)

        row_ctrl.addWidget(self.btn_replay_play)
        row_ctrl.addWidget(self.btn_replay_stop)
        row_ctrl.addWidget(lbl_speed)
        row_ctrl.addWidget(self.cmb_replay_speed)
        row_ctrl.addStretch()
        replay_layout.addLayout(row_ctrl)

        self.slider_replay = QSlider(Qt.Horizontal)
        self.slider_replay.setEnabled(False)
        self.slider_replay.setStyleSheet("""
            QSlider::groove:horizontal {
                border: 1px solid #1E293B;
                height: 4px;
                background: #0B1120;
                border-radius: 2px;
            }
            QSlider::sub-page:horizontal {
                background: #38BDF8;
                border-radius: 2px;
            }
            QSlider::handle:horizontal {
                background: #38BDF8;
                border: 1px solid #38BDF8;
                width: 12px;
                height: 12px;
                margin: -4px 0;
                border-radius: 6px;
            }
            QSlider::handle:horizontal:hover {
                background: #F8FAFC;
                border-color: #F8FAFC;
            }
        """)
        self.slider_replay.sliderPressed.connect(self.on_slider_pressed)
        self.slider_replay.sliderReleased.connect(self.on_slider_released)
        self.slider_replay.valueChanged.connect(self.on_slider_value_changed)

        self.lbl_replay_time = QLabel("00:00:00 / 00:00:00")
        self.lbl_replay_time.setStyleSheet("color: #94A3B8; font-size: 11px; font-family: Consolas;")
        self.lbl_replay_time.setAlignment(Qt.AlignCenter)

        replay_layout.addWidget(self.slider_replay)
        replay_layout.addWidget(self.lbl_replay_time)

        self.mode_tab.addTab(self.tab_replay, "数据回放")
        self.mode_tab.currentChanged.connect(self.on_mode_tab_changed)
        outer_layout.addWidget(self.mode_tab)

        serial_left_layout.addWidget(self.group_serial_ctrl)

        # 将数据解析卡片直接垂直叠放在左侧串口配置下方
        serial_left_layout.addWidget(self.dashboard_tab)
        self.serial_upper_splitter.addWidget(self.serial_left_panel)

        # 串口文本终端卡片包裹
        self.group_console = QGroupBox("串口文本终端")
        console_layout = QVBoxLayout(self.group_console)
        console_layout.setContentsMargins(8, 16, 8, 8)
        console_layout.setSpacing(0)

        # F.2.1 定位解析数据及快捷指令 Tab 构建

        # F.2.1 定位解析数据面板
        self.pane_pnt = QWidget()
        self.pane_pnt.setStyleSheet("QLabel { font-size: 13px; }")
        pnt_grid = QGridLayout(self.pane_pnt)
        pnt_grid.setContentsMargins(16, 20, 16, 20)
        pnt_grid.setSpacing(12)

        pnt_grid.addWidget(QLabel("UTC 时间:"), 0, 0)
        self.lbl_pnt_utc = QLabel("--:--:--.--")
        self.lbl_pnt_utc.setStyleSheet("color: #F8FAFC; font-weight: bold; font-family: Consolas; font-size: 15px;")
        pnt_grid.addWidget(self.lbl_pnt_utc, 0, 1)

        pnt_grid.addWidget(QLabel("定位质量:"), 0, 2)
        self.lbl_pnt_quality = QLabel("未定位")
        self.lbl_pnt_quality.setStyleSheet("background-color: #334155; color: #94A3B8; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;")
        self.lbl_pnt_quality.setAlignment(Qt.AlignCenter)
        pnt_grid.addWidget(self.lbl_pnt_quality, 0, 3)

        pnt_grid.addWidget(QLabel("纬度 (Lat):"), 1, 0)
        self.lbl_pnt_lat = QLabel("---.--------")
        self.lbl_pnt_lat.setStyleSheet("color: #E2E8F0; font-weight: bold; font-family: Consolas; font-size: 15px;")
        pnt_grid.addWidget(self.lbl_pnt_lat, 1, 1)

        pnt_grid.addWidget(QLabel("经度 (Lon):"), 1, 2)
        self.lbl_pnt_lon = QLabel("---.--------")
        self.lbl_pnt_lon.setStyleSheet("color: #E2E8F0; font-weight: bold; font-family: Consolas; font-size: 15px;")
        pnt_grid.addWidget(self.lbl_pnt_lon, 1, 3)

        pnt_grid.addWidget(QLabel("椭球高 (HAE):"), 2, 0)
        self.lbl_pnt_alt = QLabel("---.--- 米")
        self.lbl_pnt_alt.setStyleSheet("color: #E2E8F0; font-weight: bold; font-size: 14px;")
        pnt_grid.addWidget(self.lbl_pnt_alt, 2, 1)

        pnt_grid.addWidget(QLabel("解算星数:"), 2, 2)
        self.lbl_pnt_num = QLabel("0 颗")
        self.lbl_pnt_num.setStyleSheet("color: #E2E8F0; font-weight: bold; font-size: 14px;")
        pnt_grid.addWidget(self.lbl_pnt_num, 2, 3)

        pnt_grid.addWidget(QLabel("PDOP:"), 3, 0)
        self.lbl_pnt_pdop = QLabel("---")
        self.lbl_pnt_pdop.setStyleSheet("color: #E2E8F0; font-weight: bold; font-size: 14px;")
        pnt_grid.addWidget(self.lbl_pnt_pdop, 3, 1)

        pnt_grid.addWidget(QLabel("HDOP:"), 3, 2)
        self.lbl_pnt_hdop = QLabel("---")
        self.lbl_pnt_hdop.setStyleSheet("color: #E2E8F0; font-weight: bold; font-size: 14px;")
        pnt_grid.addWidget(self.lbl_pnt_hdop, 3, 3)

        # 均匀分配行拉伸以合理利用垂直空白
        pnt_grid.setRowStretch(0, 1)
        pnt_grid.setRowStretch(1, 1)
        pnt_grid.setRowStretch(2, 1)
        pnt_grid.setRowStretch(3, 1)

        self.dashboard_tab.addTab(self.pane_pnt, "定位基本状态")

        # F.2.2 惯导解析数据面板
        self.pane_ins = QWidget()
        self.pane_ins.setStyleSheet("QLabel { font-size: 13px; }")
        ins_grid = QGridLayout(self.pane_ins)
        ins_grid.setContentsMargins(16, 20, 16, 20)
        ins_grid.setSpacing(12)

        ins_grid.addWidget(QLabel("惯导状态:"), 0, 0)
        self.lbl_ins_status = QLabel("未激活")
        self.lbl_ins_status.setStyleSheet("background-color: #334155; color: #94A3B8; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;")
        self.lbl_ins_status.setAlignment(Qt.AlignCenter)
        ins_grid.addWidget(self.lbl_ins_status, 0, 1)

        ins_grid.addWidget(QLabel("载体运动:"), 0, 2)
        self.lbl_ins_motion = QLabel("未知")
        self.lbl_ins_motion.setStyleSheet("color: #F8FAFC; font-weight: bold; font-size: 14px;")
        ins_grid.addWidget(self.lbl_ins_motion, 0, 3)

        ins_grid.addWidget(QLabel("滚转 (Roll):"), 1, 0)
        self.lbl_ins_roll = QLabel("---.--- 度")
        self.lbl_ins_roll.setStyleSheet("color: #E2E8F0; font-family: Consolas; font-size: 14px;")
        ins_grid.addWidget(self.lbl_ins_roll, 1, 1)

        ins_grid.addWidget(QLabel("俯仰 (Pitch):"), 1, 2)
        self.lbl_ins_pitch = QLabel("---.--- 度")
        self.lbl_ins_pitch.setStyleSheet("color: #E2E8F0; font-family: Consolas; font-size: 14px;")
        ins_grid.addWidget(self.lbl_ins_pitch, 1, 3)

        ins_grid.addWidget(QLabel("航向 (Yaw):"), 2, 0)
        self.lbl_ins_yaw = QLabel("---.--- 度")
        self.lbl_ins_yaw.setStyleSheet("color: #E2E8F0; font-family: Consolas; font-weight: bold; font-size: 15px;")
        ins_grid.addWidget(self.lbl_ins_yaw, 2, 1)

        ins_grid.addWidget(QLabel("前向速度:"), 2, 2)
        self.lbl_ins_speed = QLabel("---.- m/s")
        self.lbl_ins_speed.setStyleSheet("color: #E2E8F0; font-size: 14px; font-weight: bold;")
        ins_grid.addWidget(self.lbl_ins_speed, 2, 3)

        ins_grid.addWidget(QLabel("累计里程:"), 3, 0)
        self.lbl_ins_mileage = QLabel("---.- 米")
        self.lbl_ins_mileage.setStyleSheet("color: #E2E8F0; font-size: 14px; font-weight: bold;")
        ins_grid.addWidget(self.lbl_ins_mileage, 3, 1)

        ins_grid.addWidget(QLabel("周内秒 (TOW):"), 3, 2)
        self.lbl_ins_tow = QLabel("------.---")
        self.lbl_ins_tow.setStyleSheet("color: #94A3B8; font-family: Consolas; font-size: 14px;")
        ins_grid.addWidget(self.lbl_ins_tow, 3, 3)

        # 均匀分配行拉伸以合理利用垂直空白
        ins_grid.setRowStretch(0, 1)
        ins_grid.setRowStretch(1, 1)
        ins_grid.setRowStretch(2, 1)
        ins_grid.setRowStretch(3, 1)

        self.dashboard_tab.addTab(self.pane_ins, "组合惯导参数")

        # F.2.3 快捷指令面板
        self.pane_cmd = QWidget()
        cmd_grid = QGridLayout(self.pane_cmd)
        cmd_grid.setContentsMargins(16, 20, 16, 20)
        cmd_grid.setSpacing(12)

        self.btn_cmd_cold = QPushButton("冷启动")
        self.btn_cmd_cold.setFixedHeight(28)
        self.btn_cmd_cold.setToolTip("发送冷启动复位指令")
        self.btn_cmd_cold.clicked.connect(lambda: self.send_serial_command('cold'))
        cmd_grid.addWidget(self.btn_cmd_cold, 0, 0)

        self.btn_cmd_hot = QPushButton("热启动")
        self.btn_cmd_hot.setFixedHeight(28)
        self.btn_cmd_hot.setToolTip("发送热启动复位指令")
        self.btn_cmd_hot.clicked.connect(lambda: self.send_serial_command('hot'))
        cmd_grid.addWidget(self.btn_cmd_hot, 0, 1)

        self.btn_cmd_ver = QPushButton("查询版本")
        self.btn_cmd_ver.setFixedHeight(28)
        self.btn_cmd_ver.setToolTip("发送查询固件版本指令")
        self.btn_cmd_ver.clicked.connect(lambda: self.send_serial_command('version'))
        cmd_grid.addWidget(self.btn_cmd_ver, 1, 0)

        self.btn_cmd_save = QPushButton("保存配置")
        self.btn_cmd_save.setFixedHeight(28)
        self.btn_cmd_save.setToolTip("发送保存当前配置到 Flash 指令")
        self.btn_cmd_save.clicked.connect(lambda: self.send_serial_command('save'))
        cmd_grid.addWidget(self.btn_cmd_save, 1, 1)

        # 均匀分配行拉伸以合理利用垂直空白
        cmd_grid.setRowStretch(0, 1)
        cmd_grid.setRowStretch(1, 1)

        self.dashboard_tab.addTab(self.pane_cmd, "快捷指令")

        # 将控制台和解析卡片加入上层水平分割条 (控制台在右侧)
        self.txt_console = QTextEdit()
        self.txt_console.setReadOnly(True)
        self.txt_console.document().setMaximumBlockCount(500)
        self.console_scroll_timer = QTimer(self)
        self.console_scroll_timer.setSingleShot(True)
        self.console_scroll_timer.timeout.connect(self.scroll_console_to_bottom)
        self.txt_console.setStyleSheet("""
            background-color: #0B1120;
            color: #10B981;
            font-family: 'Consolas', 'Courier New', monospace;
            font-size: 11px;
            border: 1px solid #1E293B;
            border-radius: 6px;
            padding: 4px;
        """)
        self.txt_console.verticalScrollBar().valueChanged.connect(self.on_console_scrollbar_value_changed)
        console_layout.addWidget(self.txt_console)

        # 新增串口发送控制行
        send_layout = QHBoxLayout()
        send_layout.setSpacing(6)

        self.txt_send_input = QLineEdit()
        self.txt_send_input.setFixedHeight(28)
        self.txt_send_input.setPlaceholderText("输入要发送的命令或数据...")
        self.txt_send_input.setStyleSheet("""
            QLineEdit {
                background-color: #0B1120;
                color: #F8FAFC;
                border: 1px solid #334155;
                border-radius: 4px;
                padding-left: 6px;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 11px;
            }
            QLineEdit:focus {
                border-color: #38BDF8;
            }
        """)
        self.txt_send_input.returnPressed.connect(self.send_custom_data)
        send_layout.addWidget(self.txt_send_input, 1)

        self.cb_send_ln = QCheckBox("加换行")
        self.cb_send_ln.setChecked(True)
        self.cb_send_ln.setStyleSheet("""
            QCheckBox { color: #94A3B8; font-size: 11px; }
            QCheckBox:disabled { color: #475569; }
        """)
        send_layout.addWidget(self.cb_send_ln)

        self.cb_send_hex = QCheckBox("Hex发送")
        self.cb_send_hex.setStyleSheet("color: #94A3B8; font-size: 11px;")
        self.cb_send_hex.stateChanged.connect(self.update_send_ln_state)
        send_layout.addWidget(self.cb_send_hex)

        self.btn_send = QPushButton("发送")
        self.btn_send.setFixedSize(60, 28)
        self.btn_send.setStyleSheet("""
            QPushButton {
                background-color: rgba(56, 189, 248, 0.15);
                border: 1px solid #38BDF8;
                color: #38BDF8;
                font-weight: bold;
                border-radius: 4px;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: rgba(56, 189, 248, 0.25);
                color: #FFFFFF;
            }
            QPushButton:pressed {
                background-color: rgba(56, 189, 248, 0.35);
            }
        """)
        self.btn_send.clicked.connect(self.send_custom_data)
        send_layout.addWidget(self.btn_send)

        console_layout.addLayout(send_layout)

        self.serial_upper_splitter.addWidget(self.group_console)

        # F.2.2 下层：可见卫星载噪比监视器卡片包裹
        self.group_cno = QGroupBox("可见卫星载噪比监视器")
        cno_layout = QVBoxLayout(self.group_cno)
        cno_layout.setContentsMargins(8, 16, 8, 8)
        cno_layout.setSpacing(0)

        self.canvas_cno = CNoPlotCanvas(self.group_cno)
        cno_layout.addWidget(self.canvas_cno)
        self.serial_vertical_splitter.addWidget(self.group_cno)

        # 设置初始分割比例 (左侧堆叠栏宽度 320px，右侧串口终端宽度 920px)
        self.serial_upper_splitter.setSizes([320, 920])
        self.serial_upper_splitter.setCollapsible(0, False)
        self.serial_upper_splitter.setCollapsible(1, False)
        self.serial_vertical_splitter.setSizes([600, 300])
        self.serial_vertical_splitter.setCollapsible(0, False)
        self.serial_vertical_splitter.setCollapsible(1, False)

        # 添加选项卡
        self.tab_widget.addTab(self.tab_scatter, "靶心图")
        self.tab_widget.addTab(self.tab_epoch_h, "水平位置误差历元分布图")
        self.tab_widget.addTab(self.tab_epoch_v, "高程误差历元分布图")
        self.tab_widget.addTab(self.tab_status, "定位质量图")
        self.tab_widget.addTab(self.tab_trajectory, "绝对轨迹图")
        self.tab_widget.addTab(self.tab_serial, "实时串口")

        # E. 精度统计对比页
        self.tab_metrics = QWidget()
        layout_metrics = QVBoxLayout(self.tab_metrics)
        layout_metrics.setContentsMargins(15, 15, 15, 15)

        self.table_metrics = QTableWidget(0, 11)
        self.table_metrics.setHorizontalHeaderLabels([
            "分段名称", "有效对齐率 (%)", "总历元数", "固定率", "CEP50 (m)", "CEP68 (m)", "CEP95 (m)",
            "RMS(水平) (m)", "RMS(高程) (m)", "最大水平偏差 (m)", "最大高程偏差 (m)"
        ])
        self.table_metrics.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table_metrics.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.table_metrics.verticalHeader().setVisible(False)
        self.table_metrics.setStyleSheet("""
            QTableWidget {
                background-color: #FFFFFF;
                color: #0F172A;
                gridline-color: #E2E8F0;
                border: 1px solid #CBD5E1;
                border-radius: 4px;
                font-size: 13px;
            }
            QTableWidget::item {
                color: #0F172A;
            }
            QHeaderView::section {
                background-color: #F1F5F9;
                color: #0F172A;
                font-weight: bold;
                padding: 6px;
                border: 1px solid #E2E8F0;
            }
        """)
        layout_metrics.addWidget(self.table_metrics)
        self.tab_widget.addTab(self.tab_metrics, "精度统计对比")

        self.tab_widget.currentChanged.connect(self.on_tab_changed)
        left_layout.addWidget(self.tab_widget)

        self.main_splitter.addWidget(self.left_widget)

        # 3. 右侧属性侧边栏 (深色)
        self.sidebar_widget = QWidget()
        self.sidebar_widget.setObjectName("sidebar_container")
        sidebar_layout = QVBoxLayout(self.sidebar_widget)
        sidebar_layout.setContentsMargins(8, 8, 8, 8)
        sidebar_layout.setSpacing(10)

        # 3.1 参考位置卡片
        self.group_ref = QGroupBox("参考位置")
        ref_layout = QVBoxLayout(self.group_ref)
        ref_layout.setContentsMargins(8, 12, 8, 8)
        ref_layout.setSpacing(6)

        # Auto / Manual 切换按钮
        switch_layout = QHBoxLayout()
        switch_layout.setSpacing(2)

        self.btn_ref_auto = QPushButton("自动(均值)")
        self.btn_ref_auto.setObjectName("btn_ref_auto")
        self.btn_ref_auto.setProperty("active", True)
        self.btn_ref_auto.clicked.connect(self.set_ref_mode_auto)

        self.btn_ref_manual = QPushButton("自定义")
        self.btn_ref_manual.setObjectName("btn_ref_manual")
        self.btn_ref_manual.setProperty("active", False)
        self.btn_ref_manual.clicked.connect(self.set_ref_mode_manual)

        self.btn_ref_dynamic = QPushButton("动态文件")
        self.btn_ref_dynamic.setObjectName("btn_ref_dynamic")
        self.btn_ref_dynamic.setProperty("active", False)
        self.btn_ref_dynamic.clicked.connect(self.set_ref_mode_dynamic)

        switch_layout.addWidget(self.btn_ref_auto)
        switch_layout.addWidget(self.btn_ref_manual)
        switch_layout.addWidget(self.btn_ref_dynamic)
        ref_layout.addLayout(switch_layout)

        # 历史输入坐标下拉框
        row_history = QHBoxLayout()
        row_history.setSpacing(6)
        lbl_history = QLabel("历史参考坐标：")
        lbl_history.setStyleSheet("color:#94A3B8; font-size:12px; font-weight:bold;")
        self.cmb_history = QComboBox()
        self.cmb_history.setFixedHeight(28)
        self.cmb_history.setDisabled(True)  # 默认Auto模式下禁用
        self.cmb_history.currentIndexChanged.connect(self.on_history_coordinate_selected)
        row_history.addWidget(lbl_history)
        row_history.addWidget(self.cmb_history)
        ref_layout.addLayout(row_history)

        # 经纬高输入框
        self.inputs_layout = QVBoxLayout()
        self.inputs_layout.setSpacing(4)

        # 备注名称
        row_name = QHBoxLayout()
        row_name.addWidget(QLabel("备注 (可选)："))
        self.txt_ref_name = QLineEdit("")
        self.txt_ref_name.setAlignment(Qt.AlignLeft)
        self.txt_ref_name.setPlaceholderText("例如: 静态测试点 A")
        self.txt_ref_name.setDisabled(True)
        self.txt_ref_name.editingFinished.connect(self.on_manual_truth_changed)
        row_name.addWidget(self.txt_ref_name)
        self.inputs_layout.addLayout(row_name)

        # 纬度
        row_lat = QHBoxLayout()
        row_lat.addWidget(QLabel("纬度 (度)："))
        self.txt_lat = QLineEdit("0.00000000")
        self.txt_lat.setAlignment(Qt.AlignLeft)
        self.txt_lat.setDisabled(True)
        self.txt_lat.editingFinished.connect(self.on_manual_truth_changed)
        row_lat.addWidget(self.txt_lat)
        self.inputs_layout.addLayout(row_lat)

        # 经度
        row_lon = QHBoxLayout()
        row_lon.addWidget(QLabel("经度 (度)："))
        self.txt_lon = QLineEdit("0.00000000")
        self.txt_lon.setAlignment(Qt.AlignLeft)
        self.txt_lon.setDisabled(True)
        self.txt_lon.editingFinished.connect(self.on_manual_truth_changed)
        row_lon.addWidget(self.txt_lon)
        self.inputs_layout.addLayout(row_lon)

        # 高程
        row_alt = QHBoxLayout()
        row_alt.addWidget(QLabel("高程 (米)："))
        self.txt_alt = QLineEdit("0.000")
        self.txt_alt.setAlignment(Qt.AlignLeft)
        self.txt_alt.setDisabled(True)
        self.txt_alt.editingFinished.connect(self.on_manual_truth_changed)
        row_alt.addWidget(self.txt_alt)
        self.inputs_layout.addLayout(row_alt)
        self.inputs_widget = QWidget()
        self.inputs_widget.setLayout(self.inputs_layout)
        ref_layout.addWidget(self.inputs_widget)

        # 动态真值输入区
        self.dynamic_layout = QVBoxLayout()
        self.dynamic_layout.setSpacing(4)
        self.lbl_dynamic_file = QLabel("未加载动态真值文件")
        self.lbl_dynamic_file.setWordWrap(True)
        self.lbl_dynamic_file.setStyleSheet("color:#38BDF8;")

        dyn_btn_layout = QHBoxLayout()
        dyn_btn_layout.setSpacing(4)
        self.btn_load_dynamic = QPushButton("导入真值")
        self.btn_load_dynamic.clicked.connect(self.on_load_dynamic_clicked)
        self.btn_clear_dynamic = QPushButton("清除真值")
        self.btn_clear_dynamic.setEnabled(False)
        self.btn_clear_dynamic.clicked.connect(self.on_clear_dynamic_clicked)

        dyn_btn_layout.addWidget(self.btn_load_dynamic)
        dyn_btn_layout.addWidget(self.btn_clear_dynamic)

        self.dynamic_layout.addWidget(self.lbl_dynamic_file)
        self.dynamic_layout.addLayout(dyn_btn_layout)

        self.dynamic_widget = QWidget()
        self.dynamic_widget.setLayout(self.dynamic_layout)
        self.dynamic_widget.hide()
        ref_layout.addWidget(self.dynamic_widget)

        # 闰秒
        row_leap = QHBoxLayout()
        row_leap.addWidget(QLabel("GPS-UTC 闰秒："))
        self.txt_leap = QLineEdit("18")
        self.txt_leap.setFixedWidth(50)
        self.txt_leap.setAlignment(Qt.AlignCenter)
        self.txt_leap.editingFinished.connect(self.recompute_all)
        row_leap.addWidget(self.txt_leap)
        self.inputs_layout.addLayout(row_leap)

        sidebar_layout.addWidget(self.group_ref)

        # 3.2 文件/分析时段列表卡片
        self.group_file = QGroupBox("文件 / 分析分段")
        file_layout = QVBoxLayout(self.group_file)
        file_layout.setContentsMargins(8, 12, 8, 8)
        file_layout.setSpacing(6)

        # 动作区分割与小图标化：采用网格布局，将按钮压缩成 3 列 2 行，提高横向空间利用率
        actions_grid = QGridLayout()
        actions_grid.setSpacing(4)
        actions_grid.setContentsMargins(0, 4, 0, 4)

        self.btn_import = QPushButton("➕ 导入日志")
        self.btn_import.setToolTip("导入 GNSS 定位日志文件，支持 $GNGGA、$POGOS、$PODRS 等格式。")
        self.btn_import.clicked.connect(self.on_import_clicked)
        actions_grid.addWidget(self.btn_import, 0, 0)

        self.btn_add_segment = QPushButton("➕ 新增分段")
        self.btn_add_segment.setToolTip("在下方列表中新增一个自定义的分析时段。")
        self.btn_add_segment.clicked.connect(self.on_add_segment_clicked)
        actions_grid.addWidget(self.btn_add_segment, 0, 1)

        self.btn_export_raw = QPushButton("✂️ 数据截取")
        self.btn_export_raw.setToolTip("截取当前分析时段内的定位日志，并另存为新文件。")
        self.btn_export_raw.clicked.connect(self.on_export_raw_clicked)
        actions_grid.addWidget(self.btn_export_raw, 0, 2)

        self.btn_export_gga = QPushButton("🌐 格式转换")
        self.btn_export_gga.setToolTip("将当前时段内的 POGOS 或 PODRS 格式数据转换并导出为标准的 GGA 语句文件。")
        self.btn_export_gga.clicked.connect(self.on_export_gga_clicked)
        actions_grid.addWidget(self.btn_export_gga, 1, 0)

        self.btn_export_kml = QPushButton("🧭 导出 KML")
        self.btn_export_kml.setToolTip("根据当前时段的数据生成 KML 轨迹文件，可在 Google Earth 等地图中查看。")
        self.btn_export_kml.clicked.connect(self.on_export_kml_clicked)
        actions_grid.addWidget(self.btn_export_kml, 1, 1)

        self.btn_export_report = QPushButton("📄 导出 Word")
        self.btn_export_report.setToolTip("将当前时段的定位精度分析结果与图表自动生成为 Word 格式报告。")
        self.btn_export_report.clicked.connect(self.on_export_report_clicked)
        actions_grid.addWidget(self.btn_export_report, 1, 2)

        file_layout.addLayout(actions_grid)

        # 设置区网格对齐 (2-Column Key-Value Grid)：左右两列排布下拉框与复选框切换开关
        options_layout = QGridLayout()
        options_layout.setSpacing(6)
        options_layout.setContentsMargins(0, 4, 0, 4)

        # 左列：时间下拉框
        row_tz = QHBoxLayout()
        row_tz.setSpacing(2)
        lbl_tz = QLabel("时间:")
        lbl_tz.setStyleSheet("color:#94A3B8; font-size:12px; font-weight:bold;")
        self.cmb_timezone = QComboBox()
        self.cmb_timezone.addItems(["UTC 时间", "北京时间 (UTC+8)"])
        self.cmb_timezone.setFixedWidth(110)
        self.cmb_timezone.setFixedHeight(28)
        self.cmb_timezone.currentTextChanged.connect(self.on_timezone_changed)
        row_tz.addWidget(lbl_tz)
        row_tz.addWidget(self.cmb_timezone)
        row_tz.addStretch()
        options_layout.addLayout(row_tz, 0, 0)

        # 右列：X轴下拉框
        row_xaxis = QHBoxLayout()
        row_xaxis.setSpacing(2)
        lbl_xaxis = QLabel("X轴:")
        lbl_xaxis.setStyleSheet("color:#94A3B8; font-size:12px; font-weight:bold;")
        self.cmb_xaxis = QComboBox()
        self.cmb_xaxis.addItems(["历元数", "时间轴"])
        self.cmb_xaxis.setFixedWidth(90)
        self.cmb_xaxis.setFixedHeight(28)
        self.cmb_xaxis.currentTextChanged.connect(self.on_xaxis_changed)
        row_xaxis.addWidget(lbl_xaxis)
        row_xaxis.addWidget(self.cmb_xaxis)
        row_xaxis.addStretch()
        options_layout.addLayout(row_xaxis, 0, 1)

        # 复选开关：左列 - 高程误差绝对值，右列 - 显示高程值
        self.cb_abs_alt = QCheckBox("高程误差绝对值")
        self.cb_abs_alt.setToolTip("开启后，高程误差将以绝对值（去除正负号）进行计算与展示。")
        self.cb_abs_alt.setStyleSheet("color: #94A3B8; font-size: 11px; font-weight: bold;")
        self.cb_abs_alt.setChecked(False)
        self.cb_abs_alt.stateChanged.connect(self.on_alt_mode_changed)
        options_layout.addWidget(self.cb_abs_alt, 1, 0)

        self.cb_raw_alt = QCheckBox("显示高程值")
        self.cb_raw_alt.setToolTip("开启后，图表将直接展示高程的绝对物理值，而非相对于真值的误差值。")
        self.cb_raw_alt.setStyleSheet("color: #94A3B8; font-size: 11px; font-weight: bold;")
        self.cb_raw_alt.setChecked(False)
        self.cb_raw_alt.stateChanged.connect(self.on_raw_alt_changed)
        options_layout.addWidget(self.cb_raw_alt, 1, 1)

        self.cb_show_extrema = QCheckBox("显示极值")
        self.cb_show_extrema.setToolTip("在误差分布图上自动标注最大误差值与最小误差值点。")
        self.cb_show_extrema.setStyleSheet("color: #94A3B8; font-size: 11px; font-weight: bold;")
        self.cb_show_extrema.setChecked(True)
        self.cb_show_extrema.stateChanged.connect(self.on_extrema_mode_changed)
        options_layout.addWidget(self.cb_show_extrema, 2, 0)

        self.cb_show_sats = QCheckBox("显示卫星与DOP")
        self.cb_show_sats.setToolTip("在下方开启双联屏，联动分析在用卫星数与 DOP 变化趋势。")
        self.cb_show_sats.setStyleSheet("color: #94A3B8; font-size: 11px; font-weight: bold;")
        self.cb_show_sats.setChecked(False)
        self.cb_show_sats.stateChanged.connect(self.on_show_sats_changed)
        options_layout.addWidget(self.cb_show_sats, 2, 1)

        file_layout.addLayout(options_layout)

        # 文件列表
        self.list_segments = QListWidget()
        file_layout.addWidget(self.list_segments)

        sidebar_layout.addWidget(self.group_file)

        # 状态指示
        self.lbl_status = QLabel("等待导入 GNSS 日志...")
        self.lbl_status.setStyleSheet("color:#64748B; font-size:10px;")
        sidebar_layout.addWidget(self.lbl_status)

        self.sidebar_widget.setMinimumWidth(320)
        self.sidebar_widget.setMaximumWidth(450)
        self.main_splitter.addWidget(self.sidebar_widget)
        self.main_splitter.setSizes([860, 320])
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 0)

        # 初始化工具栏样式 (适配系统浅色/深色主题)
        self.update_toolbar_styles()

        # 3. 初始化视图菜单并绑定显示隐藏各个模块的动作
        view_menu = self.menu_bar.addMenu("视图")

        self.action_toggle_serial_ctrl = view_menu.addAction("显示串口配置")
        self.action_toggle_serial_ctrl.setCheckable(True)
        self.action_toggle_serial_ctrl.setChecked(True)
        self.action_toggle_serial_ctrl.triggered.connect(self.toggle_serial_ctrl_visibility)

        self.action_toggle_dashboard = view_menu.addAction("显示定位状态解析")
        self.action_toggle_dashboard.setCheckable(True)
        self.action_toggle_dashboard.setChecked(True)
        self.action_toggle_dashboard.triggered.connect(self.toggle_dashboard_visibility)

        self.action_toggle_console = view_menu.addAction("显示串口打印终端")
        self.action_toggle_console.setCheckable(True)
        self.action_toggle_console.setChecked(True)
        self.action_toggle_console.triggered.connect(self.toggle_console_visibility)

        self.action_toggle_ref = view_menu.addAction("显示参考位置")
        self.action_toggle_ref.setCheckable(True)
        self.action_toggle_ref.setChecked(True)
        self.action_toggle_ref.triggered.connect(self.toggle_ref_visibility)

        self.action_toggle_file = view_menu.addAction("显示文件与分析分段")
        self.action_toggle_file.setCheckable(True)
        self.action_toggle_file.setChecked(True)
        self.action_toggle_file.triggered.connect(self.toggle_file_visibility)

        view_menu.addSeparator()

        self.action_toggle_cno = view_menu.addAction("显示载噪比监视器")
        self.action_toggle_cno.setCheckable(True)
        self.action_toggle_cno.setChecked(True)
        self.action_toggle_cno.triggered.connect(self.toggle_cno_visibility)

    def showEvent(self, event):
        super().showEvent(event)
        # Ensure right sidebar defaults to the minimum functional width (320px) on startup or resize
        sidebar_w = 320
        left_w = self.width() - sidebar_w - self.main_splitter.handleWidth()
        self.main_splitter.setSizes([left_w, sidebar_w])

    def update_toolbar_styles(self):
        """根据系统深浅色主题动态更新绘图工具栏样式"""
        is_dark = is_system_dark_mode()

        if is_dark:
            # 系统深色模式下，Matplotlib 加载的 symbolic 图标为白色，需要深色底工具栏背景
            style = """
            QToolBar {
                background-color: #1E293B;
                border: none;
                border-bottom: 1px solid #334155;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }
            QToolBar QToolButton {
                background-color: transparent;
                border: none;
                border-radius: 4px;
                padding: 4px;
                margin: 2px;
            }
            QToolBar QToolButton:hover {
                background-color: #334155;
            }
            QToolBar QToolButton:pressed {
                background-color: #0F172A;
            }
            QToolBar QToolButton:disabled {
                background-color: transparent;
                opacity: 0.5;
            }
            QToolBar QLabel {
                color: #94A3B8;
                font-size: 11px;
                font-weight: bold;
                padding-left: 10px;
                background-color: transparent;
            }
            QToolBar::separator {
                width: 1px;
                background-color: #334155;
                margin-top: 4px;
                margin-bottom: 4px;
                margin-left: 6px;
                margin-right: 6px;
            }
            """
        else:
            # 系统浅色模式下，Matplotlib 加载的 symbolic 图标为深色，需要浅色底工具栏背景
            style = """
            QToolBar {
                background-color: #F8FAFC;
                border: none;
                border-bottom: 1px solid #E2E8F0;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }
            QToolBar QToolButton {
                background-color: transparent;
                border: none;
                border-radius: 4px;
                padding: 4px;
                margin: 2px;
            }
            QToolBar QToolButton:hover {
                background-color: #E2E8F0;
            }
            QToolBar QToolButton:pressed {
                background-color: #CBD5E1;
            }
            QToolBar QToolButton:disabled {
                background-color: transparent;
                opacity: 0.5;
            }
            QToolBar QLabel {
                color: #475569;
                font-size: 11px;
                font-weight: bold;
                padding-left: 10px;
                background-color: transparent;
            }
            QToolBar::separator {
                width: 1px;
                background-color: #E2E8F0;
                margin-top: 4px;
                margin-bottom: 4px;
                margin-left: 6px;
                margin-right: 6px;
            }
            """

        for attr in ['toolbar_scatter', 'toolbar_status', 'toolbar_epoch_h', 'toolbar_epoch_v', 'toolbar_trajectory']:
            if hasattr(self, attr):
                getattr(self, attr).setStyleSheet(style)

    def changeEvent(self, event):
        from PySide6.QtCore import QEvent
        if event.type() in (QEvent.Type.PaletteChange, QEvent.Type.ThemeChange):
            self.update_toolbar_styles()
        super().changeEvent(event)


    # 4. 参考位置模式切换控制
    def set_ref_mode_auto(self):
        self.truth_mode = 'auto'
        self.btn_ref_auto.setProperty("active", True)
        self.btn_ref_manual.setProperty("active", False)
        self.btn_ref_dynamic.setProperty("active", False)
        self.btn_ref_auto.style().unpolish(self.btn_ref_auto)
        self.btn_ref_auto.style().polish(self.btn_ref_auto)
        self.btn_ref_manual.style().unpolish(self.btn_ref_manual)
        self.btn_ref_manual.style().polish(self.btn_ref_manual)
        self.btn_ref_dynamic.style().unpolish(self.btn_ref_dynamic)
        self.btn_ref_dynamic.style().polish(self.btn_ref_dynamic)

        self.inputs_widget.show()
        self.dynamic_widget.hide()

        self.txt_lat.setDisabled(True)
        self.txt_lon.setDisabled(True)
        self.txt_alt.setDisabled(True)
        self.txt_ref_name.setDisabled(True)
        self.cmb_history.setDisabled(True)
        self.set_avg_point_as_truth()
        self.save_config()

    def set_ref_mode_manual(self):
        self.truth_mode = 'manual'
        self.btn_ref_auto.setProperty("active", False)
        self.btn_ref_manual.setProperty("active", True)
        self.btn_ref_dynamic.setProperty("active", False)
        self.btn_ref_auto.style().unpolish(self.btn_ref_auto)
        self.btn_ref_auto.style().polish(self.btn_ref_auto)
        self.btn_ref_manual.style().unpolish(self.btn_ref_manual)
        self.btn_ref_manual.style().polish(self.btn_ref_manual)
        self.btn_ref_dynamic.style().unpolish(self.btn_ref_dynamic)
        self.btn_ref_dynamic.style().polish(self.btn_ref_dynamic)

        self.inputs_widget.show()
        self.dynamic_widget.hide()

        self.txt_lat.setEnabled(True)
        self.txt_lon.setEnabled(True)
        self.txt_alt.setEnabled(True)
        self.txt_ref_name.setEnabled(True)
        self.cmb_history.setEnabled(True)
        self.on_manual_truth_changed()

    def set_ref_mode_dynamic(self):
        self.truth_mode = 'dynamic'
        self.truth['mode'] = 'dynamic'
        self.btn_ref_auto.setProperty("active", False)
        self.btn_ref_manual.setProperty("active", False)
        self.btn_ref_dynamic.setProperty("active", True)
        self.btn_ref_auto.style().unpolish(self.btn_ref_auto)
        self.btn_ref_auto.style().polish(self.btn_ref_auto)
        self.btn_ref_manual.style().unpolish(self.btn_ref_manual)
        self.btn_ref_manual.style().polish(self.btn_ref_manual)
        self.btn_ref_dynamic.style().unpolish(self.btn_ref_dynamic)
        self.btn_ref_dynamic.style().polish(self.btn_ref_dynamic)

        self.inputs_widget.hide()
        self.dynamic_widget.show()

        self.recompute_all()
        self.save_config()

    def on_load_dynamic_clicked(self):
        filepath, _ = QFileDialog.getOpenFileName(self, "选择真值文件", "", "All Files (*)")
        if not filepath:
            return

        self.lbl_dynamic_file.setText(f"正在加载:\n{os.path.basename(filepath)}")

        leap_secs = self.get_leap_seconds()
        strict = self.app_config.get('strict_nmea_checksum', False)

        self.dynamic_parser_thread = LogParserThread(filepath, leap_secs, strict, self)
        self.dynamic_parser_thread.finished_parsing.connect(lambda res: self.on_dynamic_parse_finished(res, filepath))
        self.dynamic_parser_thread.error_occurred.connect(self.on_dynamic_parse_error)
        self.dynamic_parser_thread.start()

    def on_dynamic_parse_error(self, err_msg):
        QMessageBox.critical(self, "解析错误", f"动态真值文件解析失败: {err_msg}")
        self.lbl_dynamic_file.setText("加载失败")
        self.truth['epochs'] = None
        self.recompute_all()

    def on_dynamic_parse_finished(self, result, filepath):
        file_epochs = result['file_epochs']
        if not file_epochs:
            QMessageBox.warning(self, "警告", "未能在文件中提取到有效的 GGA/RMC 轨迹数据！")
            self.lbl_dynamic_file.setText("未加载有效的真值文件")
            self.truth['epochs'] = None
            self.recompute_all()
            return

        # 去重：优先保留 GGA 语句，因为 GGA 包含高度和更详细的质量状态信息
        best_epochs = {}
        for ep in file_epochs:
            t = ep['utc_time_sec']
            if t not in best_epochs:
                best_epochs[t] = ep
            else:
                # 若已存在且为 RMC，新来的是 GGA 时，进行替换
                if best_epochs[t]['type'] == 'RMC' and ep['type'] == 'GGA':
                    best_epochs[t] = ep

        # 按时间排序转回列表
        dedup_epochs = [best_epochs[t] for t in sorted(best_epochs.keys())]

        self.truth['mode'] = 'dynamic'
        self.truth['file_id'] = filepath
        self.truth['epochs'] = dedup_epochs

        self.lbl_dynamic_file.setText(f"真值加载成功 ({len(dedup_epochs)}个点):\n{os.path.basename(filepath)}")
        self.btn_clear_dynamic.setEnabled(True)
        self.recompute_all()
        self.save_config()

    def on_clear_dynamic_clicked(self):
        self.truth['epochs'] = None
        self.truth['file_id'] = None
        self.lbl_dynamic_file.setText("未加载动态真值文件")
        self.btn_clear_dynamic.setEnabled(False)
        self.recompute_all()
        self.save_config()

    def on_manual_truth_changed(self):
        if self.truth_mode == 'manual':
            try:
                lat = float(self.txt_lat.text())
                lon = float(self.txt_lon.text())
                alt = float(self.txt_alt.text())
                name = self.txt_ref_name.text().strip()

                self.truth['lat'] = lat
                self.truth['lon'] = lon
                self.truth['alt'] = alt

                # 记录有效坐标到历史中，会自动更新下拉框并保存配置
                self.add_coordinate_to_history(lat, lon, alt, name)
                self.recompute_all()
            except ValueError:
                QMessageBox.warning(self, "输入错误", "请输入合法的经度、纬度或高程数值。")
    def cancel_parsing(self):
        self.pending_parse_queue.clear()
        if hasattr(self, 'parser_thread') and self.parser_thread:
            self.parser_thread.is_cancelled = True
        if hasattr(self, 'progress_dialog') and self.progress_dialog:
            self.progress_dialog.close()

    def on_import_clicked(self):
        filepaths, _ = QFileDialog.getOpenFileNames(self, "导入定位原始日志", "", "All Files (*.*);;GNSS Logs (*.log *.txt *.nmea *.dat)")
        if not filepaths:
            return

        # Check Busy Guard
        if (hasattr(self, 'parser_thread') and self.parser_thread and self.parser_thread.isRunning()) or \
           (hasattr(self, 'dynamic_parser_thread') and self.dynamic_parser_thread and self.dynamic_parser_thread.isRunning()):
            QMessageBox.warning(self, "系统繁忙", "当前后台正有数据解析线程运行，请等待当前文件导入完成后再试。")
            return

        self.import_multiple_files(filepaths)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        filepaths = []
        for url in urls:
            local_path = url.toLocalFile()
            if local_path and os.path.isfile(local_path):
                filepaths.append(local_path)

        if not filepaths:
            return

        # Check Busy Guard
        if (hasattr(self, 'parser_thread') and self.parser_thread and self.parser_thread.isRunning()) or \
           (hasattr(self, 'dynamic_parser_thread') and self.dynamic_parser_thread and self.dynamic_parser_thread.isRunning()):
            QMessageBox.warning(self, "系统繁忙", "当前后台正有数据解析线程运行，请等待当前文件导入完成后再试。")
            return

        self.import_multiple_files(filepaths)

    def import_multiple_files(self, filepaths):
        self.pending_parse_queue = list(filepaths)
        self.total_queue_count = len(filepaths)

        if not self.pending_parse_queue:
            return

        self.progress_dialog = QProgressDialog("正在准备解析文件...", "取消", 0, 100, self)
        self.progress_dialog.setWindowTitle("导入日志")
        self.progress_dialog.setWindowModality(Qt.WindowModal)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.setValue(0)
        self.progress_dialog.canceled.connect(self.cancel_parsing)

        self.start_next_queued_parse()

    def start_next_queued_parse(self):
        if not self.pending_parse_queue:
            if hasattr(self, 'progress_dialog') and self.progress_dialog:
                self.progress_dialog.close()
            return

        filepath = self.pending_parse_queue.pop(0)
        current_idx = self.total_queue_count - len(self.pending_parse_queue)

        self.lbl_status.setText(f"正在读取 ({current_idx}/{self.total_queue_count}): {os.path.basename(filepath)}")

        if hasattr(self, 'progress_dialog') and self.progress_dialog:
            self.progress_dialog.setLabelText(f"正在读取并解析 ({current_idx}/{self.total_queue_count}):\n{os.path.basename(filepath)}")
            self.progress_dialog.setValue(0)

        leap_secs = self.get_leap_seconds()
        strict = self.app_config.get('strict_nmea_checksum', False)

        self.parser_thread = LogParserThread(filepath, leap_secs, strict, self)
        self.parser_thread.progress_updated.connect(self.progress_dialog.setValue)
        self.parser_thread.finished_parsing.connect(lambda res: self.on_parse_finished(res, filepath))
        self.parser_thread.error_occurred.connect(self.on_parse_error)

        self.parser_thread.start()

    def on_parse_error(self, err_msg):
        self.pending_parse_queue.clear()
        if hasattr(self, 'progress_dialog') and self.progress_dialog:
            self.progress_dialog.close()
        QMessageBox.critical(self, "读取错误", f"解析失败: {err_msg}")
        self.lbl_status.setText("导入失败")

    def on_parse_finished(self, result, filepath):
        if hasattr(self, 'parser_thread') and self.parser_thread and self.parser_thread.is_cancelled:
            return

        file_epochs = result['file_epochs']
        first_time_sec = result['first_time_sec']
        last_time_sec = result['last_time_sec']
        first_time_str = result['first_time_str']
        last_time_str = result['last_time_str']

        for stype, count in result['sentence_types'].items():
            self.sentence_types[stype] = self.sentence_types.get(stype, 0) + count

        if not file_epochs:
            QMessageBox.warning(self, "警告", f"文件 {os.path.basename(filepath)} 中未检测到任何有效的定位语句（$GNGGA、$POGOS 或 $PODRS）！")
            self.lbl_status.setText("导入失败：无有效定位数据")
            if self.pending_parse_queue:
                self.start_next_queued_parse()
            else:
                if hasattr(self, 'progress_dialog') and self.progress_dialog:
                    self.progress_dialog.close()
            return

        # Remove old epochs of this file from the global list to prevent duplicates
        self.parsed_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') != filepath]

        # Remove old segments associated with this file from both the list and the UI list widget safely
        seg_ids_to_remove = [s['id'] for s in self.segments if s.get('file_id') == filepath]
        for seg_id in seg_ids_to_remove:
            self.segments = [s for s in self.segments if s['id'] != seg_id]
            i = 0
            while i < self.list_segments.count():
                item = self.list_segments.item(i)
                widget = self.list_segments.itemWidget(item)
                if isinstance(widget, SegmentListItemWidget) and widget.seg_id == seg_id:
                    self.list_segments.takeItem(i)
                    continue
                i += 1

        self.parsed_epochs.extend(file_epochs)
        self.parsed_epochs.sort(key=lambda x: x['utc_time_sec'])
        self.file_epochs_map[filepath] = []
        self.file_epochs_map[filepath].extend(file_epochs)

        # 为了后面的 bisect 二分查找，必须保证 file_epochs 是按时间严格递增的
        self.file_epochs_map[filepath].sort(key=lambda x: x['utc_time_sec'])

        # 更新默认起止范围 (合并多文件的范围)
        if self.time_range['start'] == 0 or first_time_sec < self.time_range['start']:
            self.time_range['start'] = first_time_sec
            self.first_time_str = first_time_str
        if self.time_range['end'] == 0 or last_time_sec > self.time_range['end']:
            self.time_range['end'] = last_time_sec
            self.last_time_str = last_time_str
        self.lbl_status.setText(f"已导入: {os.path.basename(filepath)}")

        # 为导入的新文件自动添加一个分析时段，名称设为文件名
        filename_clean = os.path.splitext(os.path.basename(filepath))[0]
        self.add_segment_item(filename_clean, first_time_str, last_time_str, file_id=filepath)

        if self.truth_mode == 'auto':
            self.set_avg_point_as_truth()
        else:
            self.recompute_all()

        if self.pending_parse_queue:
            self.start_next_queued_parse()
        else:
            if hasattr(self, 'progress_dialog') and self.progress_dialog:
                self.progress_dialog.close()

    def on_add_segment_clicked(self):
        if not self.parsed_epochs:
            QMessageBox.warning(self, "提示", "请先导入原始日志数据！")
            return

        # 默认使用最后一个时段的 file_id，如果没有，使用最后一个解析历元的 file_id
        last_file_id = self.segments[-1]['file_id'] if self.segments else self.parsed_epochs[-1]['file_id']
        file_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') == last_file_id]

        start_str = file_epochs[0]['time_str'] if file_epochs else self.first_time_str
        end_str = file_epochs[-1]['time_str'] if file_epochs else self.last_time_str

        self.add_segment_item(f"ID_{self.segment_counter + 1}", start_str, end_str, file_id=last_file_id)

    def set_avg_point_as_truth(self):
        if not self.parsed_epochs:
            return

        # 智能过滤：如果文件中存在标准 GGA 定位数据，优先仅用 GGA 计算参考平均值
        # 只有在完全没有标准 GGA 时，才平均私有 POGOS 数据，防止坐标系偏差导致真值漂移
        gga_epochs = [p for p in self.parsed_epochs if p['type'] == 'GGA']
        target_epochs = gga_epochs if gga_epochs else self.parsed_epochs
        # 防御性过滤：确保参与平均值计算的历元都包含必要的经纬度和高度字段
        target_epochs = [p for p in target_epochs if isinstance(p, dict) and 'lat' in p and 'lon' in p and 'alt' in p]

        n = len(target_epochs)
        if n == 0:
            self.truth['lat'] = 0.0
            self.truth['lon'] = 0.0
            self.truth['alt'] = 0.0
        else:
            sum_lat = sum(p['lat'] for p in target_epochs)
            sum_lon = sum(p['lon'] for p in target_epochs)
            sum_alt = sum(p['alt'] for p in target_epochs)
            self.truth['lat'] = sum_lat / n
            self.truth['lon'] = sum_lon / n
            self.truth['alt'] = sum_alt / n

        self.txt_lat.setText(f"{self.truth['lat']:.8f}")
        self.txt_lon.setText(f"{self.truth['lon']:.8f}")
        self.txt_alt.setText(f"{self.truth['alt']:.3f}")
        self.recompute_all()

    # 6. 分段时段管理与交互
    def add_segment_item(self, name, start_time, end_time, file_id=None):
        if not self.parsed_epochs:
            return

        seg_id = self.segment_counter
        self.segment_counter += 1

        if not file_id and self.segments:
            file_id = self.segments[-1]['file_id']
        elif not file_id:
            file_id = self.parsed_epochs[-1]['file_id']

        # 根据分段关联的具体 file_id 的历元类型判断其支持的协议数据源
        file_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') == file_id]
        has_gga = any(ep['type'] == 'GGA' for ep in file_epochs)
        has_pogos = any(ep['type'] == 'POGOS' for ep in file_epochs)
        has_podrs = any(ep['type'] == 'PODRS' for ep in file_epochs)
        if has_gga:
            default_src = 'GGA'
        elif has_pogos:
            default_src = 'POGOS'
        else:
            default_src = 'PODRS'

        # 寻找目前未被使用的颜色，避免颜色重复
        color = None
        used_colors = {s['color'].upper() for s in self.segments}
        for candidate in self.default_colors:
            if candidate.upper() not in used_colors:
                color = candidate
                break
        if not color:
            color = self.default_colors[seg_id % len(self.default_colors)]

        seg = {
            'id': seg_id,
            'name': name,
            'start_time': start_time,
            'end_time': end_time,
            'source_type': default_src,
            'color': color,
            'active': True,
            'metrics': None,
            'epochs': [],
            'file_id': file_id  # 关联的具体文件标志
        }
        self.segments.append(seg)

        display_start = self.get_display_time(start_time)
        display_end = self.get_display_time(end_time)
        item_widget = SegmentListItemWidget(seg_id, name, display_start, display_end, default_src, color, has_gga, has_pogos, has_podrs)

        item_widget.active_toggled.connect(self.on_seg_active_toggled)
        item_widget.name_changed.connect(self.on_seg_name_changed)
        item_widget.color_changed.connect(self.on_seg_color_changed)
        item_widget.time_changed.connect(self.on_seg_time_changed)
        item_widget.source_changed.connect(self.on_seg_source_changed)
        item_widget.delete_clicked.connect(self.on_seg_delete_clicked)

        list_item = QListWidgetItem(self.list_segments)
        list_item.setSizeHint(item_widget.sizeHint())
        self.list_segments.addItem(list_item)
        self.list_segments.setItemWidget(list_item, item_widget)

        self.recompute_all()

    def on_seg_active_toggled(self, seg_id, is_active):
        for s in self.segments:
            if s['id'] == seg_id:
                s['active'] = is_active
                break
        self.recompute_all()

    def on_seg_name_changed(self, seg_id, new_name):
        for s in self.segments:
            if s['id'] == seg_id:
                s['name'] = new_name
                break
        self.update_metrics_table()
        self.refresh_chart()

    def on_seg_color_changed(self, seg_id, color_hex):
        for s in self.segments:
            if s['id'] == seg_id:
                s['color'] = color_hex
                break
        self.recompute_all()

    def on_seg_time_changed(self, seg_id, start_str, end_str):
        for s in self.segments:
            if s['id'] == seg_id:
                s['start_time'] = self.get_utc_time(start_str)
                s['end_time'] = self.get_utc_time(end_str)
                break
        self.recompute_all()

    def on_timezone_changed(self, text):
        if "北京时间" in text:
            self.time_zone = 'Beijing'
        else:
            self.time_zone = 'UTC'
        self.save_config()

        for i in range(self.list_segments.count()):
            item = self.list_segments.item(i)
            widget = self.list_segments.itemWidget(item)
            if isinstance(widget, SegmentListItemWidget):
                seg_id = widget.seg_id
                seg = next((s for s in self.segments if s['id'] == seg_id), None)
                if seg:
                    display_start = self.get_display_time(seg['start_time'])
                    display_end = self.get_display_time(seg['end_time'])
                    widget.set_times(display_start, display_end)
        self.refresh_chart()

    def on_alt_mode_changed(self, state):
        if not hasattr(self, 'canvas_epoch_v') or self.canvas_epoch_v is None:
            return
        self.show_absolute_alt = self.cb_abs_alt.isChecked()
        self.refresh_chart()
        self.save_config()

    def on_raw_alt_changed(self, state):
        if not hasattr(self, 'canvas_epoch_v') or self.canvas_epoch_v is None:
            return
        self.show_raw_alt = self.cb_raw_alt.isChecked()
        self.refresh_chart()
        self.save_config()

    def on_extrema_mode_changed(self, state):
        if not hasattr(self, 'canvas_epoch_v') or self.canvas_epoch_v is None:
            return
        self.show_extrema = self.cb_show_extrema.isChecked()
        self.refresh_chart()
        self.save_config()

    def on_show_sats_changed(self, state):
        if not hasattr(self, 'canvas_epoch_v') or self.canvas_epoch_v is None:
            return
        self.refresh_chart()

    def on_xaxis_changed(self, text):
        self.x_axis_mode = text

        # 容错校验：若多个分段起止时间跨度过大（如超过4小时且实际数据稀疏），自动切换回历元数以防折线图被压缩成极细线条
        if self.x_axis_mode == "时间轴" and self.segments:
            active_segs = [s for s in self.segments if s.get('active') and s.get('epochs')]
            if len(active_segs) > 1:
                try:
                    segs_time = []
                    for s in active_segs:
                        if s.get('file_id') == "COM_REALTIME" and s.get('epochs'):
                            t_start = s['epochs'][0]['utc_time_sec']
                            t_end = s['epochs'][-1]['utc_time_sec']
                        else:
                            t_start = time_str_to_seconds(s['start_time'])
                            t_end = time_str_to_seconds(s['end_time'])
                        segs_time.append((t_start, t_end))

                    min_sec = min(t[0] for t in segs_time)
                    max_sec = max(t[1] for t in segs_time)
                    total_span = max_sec - min_sec
                    sum_duration = sum(t[1] - t[0] for t in segs_time)

                    if total_span > 14400 and (sum_duration / total_span) < 0.15:
                        QMessageBox.warning(self, "时间轴对齐警告",
                                            "当前启用的多个分析分段起止时间相差过大（超过 4 小时）且互不连续。\n\n"
                                            "若强行使用时间轴对齐，曲线将被极度压缩。已自动帮您切换回「历元数」模式。")
                        self.cmb_xaxis.blockSignals(True)
                        self.cmb_xaxis.setCurrentText("历元数")
                        self.cmb_xaxis.blockSignals(False)
                        self.x_axis_mode = "历元数"
                except Exception:
                    pass

        self.refresh_chart()
        self.save_config()

    def on_settings_clicked(self):
        if not hasattr(self, 'app_config'):
            self.app_config = {}
        dialog = SettingsDialog(self.app_config, self)
        if dialog.exec():
            new_settings = dialog.get_settings()
            self.app_config.update(new_settings)
            self.apply_new_settings()
            self.save_config()

    def apply_new_settings(self):
        thresh = self.app_config.get('downsample_threshold', 100000)
        self.canvas_scatter.downsample_threshold = thresh
        self.canvas_status.downsample_threshold = thresh
        self.canvas_epoch_h.downsample_threshold = thresh
        self.canvas_epoch_v.downsample_threshold = thresh
        self.canvas_trajectory.downsample_threshold = thresh

        dpi = self.app_config.get('export_dpi', 150)
        self.canvas_scatter.export_dpi = dpi
        self.canvas_status.export_dpi = dpi
        self.canvas_epoch_h.export_dpi = dpi
        self.canvas_epoch_v.export_dpi = dpi
        self.canvas_trajectory.export_dpi = dpi

        self.recompute_all()

    # --- 配置存储与坐标历史管理逻辑 ---
    def load_config(self):
        CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "vcom_config.json")
        self.coordinate_history = []
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    self.app_config = json.load(f)

                # 读取配置
                self.time_zone = self.app_config.get("time_zone", "UTC")
                self.show_absolute_alt = self.app_config.get("show_absolute_alt", False)
                self.show_raw_alt = self.app_config.get("show_raw_alt", False)
                self.show_extrema = self.app_config.get("show_extrema", True)
                self.x_axis_mode = self.app_config.get("x_axis_mode", "历元数")
                self.truth_mode = self.app_config.get("truth_mode", "auto")

                # 坐标历史
                self.coordinate_history = self.app_config.get("coordinate_history", [])

                # 应用配置到UI
                if self.time_zone == 'Beijing':
                    self.cmb_timezone.setCurrentText("北京时间 (UTC+8)")
                else:
                    self.cmb_timezone.setCurrentText("UTC 时间")

                self.cb_abs_alt.setChecked(self.show_absolute_alt)
                self.cb_raw_alt.setChecked(self.show_raw_alt)
                self.cb_show_extrema.setChecked(self.show_extrema)
                self.cmb_xaxis.setCurrentText(self.x_axis_mode)

                # 填充坐标下拉框
                self.update_history_combo()

                if self.truth_mode == 'manual':
                    # 如果上次是手动模式，填入最新历史坐标或默认值
                    self.set_ref_mode_manual()
                    if self.coordinate_history:
                        latest = self.coordinate_history[0]
                        self.txt_ref_name.setText(latest.get('name', ''))
                        self.txt_lat.setText(f"{latest['lat']:.8f}")
                        self.txt_lon.setText(f"{latest['lon']:.8f}")
                        self.txt_alt.setText(f"{latest['alt']:.3f}")
                        self.truth = latest.copy()
                else:
                    self.set_ref_mode_auto()

                # 闰秒
                leap = self.app_config.get("leap_seconds", 18)
                self.txt_leap.setText(str(leap))

                view_visibility = self.app_config.get("view_visibility", {})
                self.apply_view_visibility_config(view_visibility)

            except Exception as e:
                print(f"读取配置出错: {e}")
        else:
            # 默认配置
            self.set_ref_mode_auto()
            self.update_history_combo()

    def apply_view_visibility_config(self, view_visibility):
        view_map = [
            ('serial_ctrl', 'group_serial_ctrl', 'action_toggle_serial_ctrl', self.adjust_serial_view_layout),
            ('dashboard', 'dashboard_tab', 'action_toggle_dashboard', self.adjust_serial_view_layout),
            ('console', 'group_console', 'action_toggle_console', self.adjust_serial_view_layout),
            ('cno', 'group_cno', 'action_toggle_cno', self.adjust_serial_view_layout),
            ('ref', 'group_ref', 'action_toggle_ref', self.adjust_sidebar_visibility),
            ('file', 'group_file', 'action_toggle_file', self.adjust_sidebar_visibility),
        ]
        for key, widget_name, action_name, _ in view_map:
            if key not in view_visibility or not hasattr(self, widget_name):
                continue
            visible = bool(view_visibility[key])
            getattr(self, widget_name).setVisible(visible)
            if hasattr(self, action_name):
                getattr(self, action_name).setChecked(visible)

        self.adjust_serial_view_layout()
        self.adjust_sidebar_visibility()

    def collect_view_visibility_config(self):
        return {
            "serial_ctrl": self._view_action_checked('action_toggle_serial_ctrl', 'group_serial_ctrl') if hasattr(self, 'group_serial_ctrl') else True,
            "dashboard": self._view_action_checked('action_toggle_dashboard', 'dashboard_tab') if hasattr(self, 'dashboard_tab') else True,
            "console": self._view_action_checked('action_toggle_console', 'group_console') if hasattr(self, 'group_console') else True,
            "cno": self._view_action_checked('action_toggle_cno', 'group_cno') if hasattr(self, 'group_cno') else True,
            "ref": self._view_action_checked('action_toggle_ref', 'group_ref') if hasattr(self, 'group_ref') else True,
            "file": self._view_action_checked('action_toggle_file', 'group_file') if hasattr(self, 'group_file') else True,
        }

    def save_config(self):
        CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "vcom_config.json")
        self.app_config.update({
            "time_zone": self.time_zone,
            "show_absolute_alt": self.show_absolute_alt,
            "show_raw_alt": self.show_raw_alt,
            "show_extrema": self.show_extrema,
            "x_axis_mode": self.x_axis_mode,
            "truth_mode": self.truth_mode,
            "leap_seconds": self.get_leap_seconds(),
            "coordinate_history": self.coordinate_history,
            "view_visibility": self.collect_view_visibility_config()
        })
        try:
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(self.app_config, f, indent=4, ensure_ascii=False)
        except Exception as e:
            print(f"保存配置出错: {e}")

    def add_coordinate_to_history(self, lat, lon, alt, name=""):
        # 查找是否已存在相同坐标（微小误差阈值内判为相同）
        exist_idx = -1
        for i, item in enumerate(self.coordinate_history):
            if abs(item['lat'] - lat) < 1e-8 and abs(item['lon'] - lon) < 1e-8 and abs(item['alt'] - alt) < 1e-3:
                exist_idx = i
                break

        coord = {"lat": lat, "lon": lon, "alt": alt, "name": name}
        if exist_idx != -1:
            if not name and 'name' in self.coordinate_history[exist_idx] and self.coordinate_history[exist_idx]['name']:
                coord['name'] = self.coordinate_history[exist_idx]['name']
            self.coordinate_history.pop(exist_idx)
        self.coordinate_history.insert(0, coord)

        # 限制个数为 10
        self.coordinate_history = self.coordinate_history[:10]

        # 更新下拉框并保存
        self.update_history_combo()
        self.save_config()

    def update_history_combo(self):
        # 阻止信号触发选择改变回调，防止死循环
        self.cmb_history.blockSignals(True)
        self.cmb_history.clear()
        self.cmb_history.addItem("--- 选择历史输入坐标 ---")
        for item in self.coordinate_history:
            name_str = f"[{item['name']}] " if item.get('name') else ""
            self.cmb_history.addItem(f"{name_str}纬:{item['lat']:.8f}, 经:{item['lon']:.8f}, 高:{item['alt']:.3f}")
        self.cmb_history.blockSignals(False)

    def on_history_coordinate_selected(self, index):
        if index <= 0 or not self.coordinate_history:
            return
        # 下拉框第0个元素是提示，真实坐标索引为 index - 1
        coord = self.coordinate_history[index - 1]
        self.txt_ref_name.setText(coord.get('name', ''))
        self.txt_lat.setText(f"{coord['lat']:.8f}")
        self.txt_lon.setText(f"{coord['lon']:.8f}")
        self.txt_alt.setText(f"{coord['alt']:.3f}")

        self.truth['lat'] = coord['lat']
        self.truth['lon'] = coord['lon']
        self.truth['alt'] = coord['alt']
        self.recompute_all()

        # 联动将所选坐标提到历史最前
        self.add_coordinate_to_history(coord['lat'], coord['lon'], coord['alt'], coord.get('name', ''))

    def get_display_time(self, utc_time_str):
        """将内部存储的 UTC 时间字符串转为当前应显示的时间字符串"""
        if self.time_zone == 'Beijing':
            try:
                parts = utc_time_str.split(':')
                h = int(parts[0])
                m = int(parts[1])
                s_part = parts[2]
                if '.' in s_part:
                    s_val = float(s_part)
                else:
                    s_val = int(s_part)
                h = (h + 8) % 24
                if isinstance(s_val, float):
                    return f"{h:02d}:{m:02d}:{s_val:05.2f}"
                else:
                    return f"{h:02d}:{m:02d}:{s_val:02d}"
            except Exception:
                return utc_time_str
        return utc_time_str

    def get_utc_time(self, display_time_str):
        """将当前编辑框显示的时间字符串转为内部存储的 UTC 时间字符串"""
        if self.time_zone == 'Beijing':
            try:
                parts = display_time_str.split(':')
                h = int(parts[0])
                m = int(parts[1])
                s_part = parts[2]
                if '.' in s_part:
                    s_val = float(s_part)
                else:
                    s_val = int(s_part)
                h = (h - 8 + 24) % 24
                if isinstance(s_val, float):
                    return f"{h:02d}:{m:02d}:{s_val:05.2f}"
                else:
                    return f"{h:02d}:{m:02d}:{s_val:02d}"
            except Exception:
                return display_time_str
        return display_time_str

    def on_seg_source_changed(self, seg_id, src_type):
        for s in self.segments:
            if s['id'] == seg_id:
                s['source_type'] = src_type
                break
        self.recompute_all()

    def on_seg_delete_clicked(self, seg_id):
        self.segments = [s for s in self.segments if s['id'] != seg_id]

        for i in range(self.list_segments.count()):
            item = self.list_segments.item(i)
            widget = self.list_segments.itemWidget(item)
            if isinstance(widget, SegmentListItemWidget) and widget.seg_id == seg_id:
                self.list_segments.takeItem(i)
                break
        self.recompute_all()

    # 7. 全局重算与渲染联动
    def check_xaxis_mode_availability(self):
        if not self.segments:
            model = self.cmb_xaxis.model()
            if model:
                item = model.item(1)
                if item:
                    item.setEnabled(True)
            return True

        active_segs = [s for s in self.segments if s.get('active') and s.get('epochs')]
        if len(active_segs) <= 1:
            model = self.cmb_xaxis.model()
            if model:
                item = model.item(1)
                if item:
                    item.setEnabled(True)
            return True

        try:
            segs_time = []
            for s in active_segs:
                if s.get('file_id') == "COM_REALTIME" and s.get('epochs'):
                    t_start = s['epochs'][0]['utc_time_sec']
                    t_end = s['epochs'][-1]['utc_time_sec']
                else:
                    t_start = time_str_to_seconds(s['start_time'])
                    t_end = time_str_to_seconds(s['end_time'])
                segs_time.append((t_start, t_end))

            segs_time.sort(key=lambda x: x[0])

            # 检查是否有严重的时间不同步：空隙大于 1.5 小时 (5400 秒) 或者总跨度超过 4 小时 (14400 秒)
            has_gap = False
            for i in range(len(segs_time) - 1):
                gap = segs_time[i+1][0] - segs_time[i][1]
                if gap > 5400:
                    has_gap = True
                    break

            min_sec = segs_time[0][0]
            max_sec = max(st[1] for st in segs_time)
            if max_sec - min_sec > 14400:
                has_gap = True

            model = self.cmb_xaxis.model()
            if model:
                item = model.item(1)
                if item:
                    if has_gap:
                        item.setEnabled(False) # 禁用下拉列表中的“时间轴”
                        if self.cmb_xaxis.currentText() == "时间轴":
                            QMessageBox.warning(self, "X轴对齐限制",
                                                "检测到启用的多个分段之间录制时间不一致且存在大跨度空隙（或超过4小时）。\n"
                                                "如果使用时间轴，图表曲线将被极度压缩。已自动为您切换为「历元数」模式。")
                            self.cmb_xaxis.blockSignals(True)
                            self.cmb_xaxis.setCurrentText("历元数")
                            self.cmb_xaxis.blockSignals(False)
                            self.x_axis_mode = "历元数"
                    else:
                        item.setEnabled(True)
        except Exception:
            pass

        return True

    def _resolve_segment_time_range(self, epochs, start_time, end_time):
        if not epochs:
            return None

        start_sec = time_str_to_seconds(start_time)
        end_sec = time_str_to_seconds(end_time)
        first_time = epochs[0]['utc_time_sec']
        last_time = epochs[-1]['utc_time_sec']
        base_day = math.floor(first_time / 86400) * 86400

        start_abs = base_day + start_sec
        end_abs = base_day + end_sec
        if start_sec > end_sec:
            end_abs += 86400

        while end_abs < first_time:
            start_abs += 86400
            end_abs += 86400

        while start_abs > last_time and start_abs - 86400 <= last_time:
            start_abs -= 86400
            end_abs -= 86400

        return start_abs, end_abs

    def _find_epoch_range(self, epochs, start_sec, end_sec):
        left, right = 0, len(epochs)
        while left < right:
            mid = (left + right) // 2
            if epochs[mid]['utc_time_sec'] < start_sec:
                left = mid + 1
            else:
                right = mid
        start_idx = left

        left, right = start_idx, len(epochs)
        while left < right:
            mid = (left + right) // 2
            if epochs[mid]['utc_time_sec'] <= end_sec:
                left = mid + 1
            else:
                right = mid
        return epochs[start_idx:left]

    def _unwrap_stream_time(self, raw_time_sec, state):
        if raw_time_sec < 0:
            return -1

        current_time = raw_time_sec + state['day_offset']
        last_time = state.get('last_time')
        if last_time is not None:
            if current_time < last_time - 43200:
                state['day_offset'] += 86400
                current_time = raw_time_sec + state['day_offset']
            elif current_time > last_time + 43200:
                state['day_offset'] -= 86400
                current_time = raw_time_sec + state['day_offset']

        state['last_time'] = current_time
        return current_time

    def _build_stream_time_state(self, file_epochs):
        day_offset = 0
        if file_epochs:
            first_time = file_epochs[0]['utc_time_sec']
            day_offset = first_time - (first_time % 86400)
        return {
            'day_offset': day_offset,
            'last_time': None
        }

    def recompute_all(self):
        if not self.parsed_epochs:
            return

        for seg in self.segments:
            if seg.get('file_id') == "COM_REALTIME":
                # 根据当前选定的数据源类型进行动态过滤，保留最近 2000 个
                if seg['source_type'] == 'GGA':
                    seg['epochs'] = [ep for ep in self.realtime_raw_epochs if ep['type'] in ['GGA', 'POSOL', 'BK_PNT_NAV']][-2000:]
                else:
                    seg['epochs'] = [ep for ep in self.realtime_raw_epochs if ep['type'] == seg['source_type']][-2000:]
            else:
                # 使用基于 file_id 的字典查询代替遍历数十万级别的全集列表
                file_epochs = self.file_epochs_map.get(seg.get('file_id'), [])

                # 使用二分查找 O(log N) 获取时间区间内的历元
                time_range = self._resolve_segment_time_range(file_epochs, seg['start_time'], seg['end_time'])
                range_epochs = self._find_epoch_range(file_epochs, *time_range) if time_range else []
                seg['epochs'] = [ep for ep in range_epochs if ep['type'] == seg['source_type']]

            # 防御性过滤：确保所有参与精度统计和绘图的历元都包含必需的定位坐标及质量字段，防止 KeyError
            seg['epochs'] = [p for p in seg['epochs'] if isinstance(p, dict) and 'lat' in p and 'lon' in p and 'alt' in p and 'quality' in p]

            dynamic_truth_array = None
            if self.truth_mode == 'auto':
                if seg['epochs']:
                    seg_truth = {
                        'lat': sum(p['lat'] for p in seg['epochs']) / len(seg['epochs']),
                        'lon': sum(p['lon'] for p in seg['epochs']) / len(seg['epochs']),
                        'alt': sum(p['alt'] for p in seg['epochs']) / len(seg['epochs'])
                    }
                else:
                    seg_truth = self.truth
            elif self.truth_mode == 'dynamic':
                seg_truth = self.truth
                if self.truth.get('epochs') and seg['epochs']:
                    dynamic_truth_array = interpolate_dynamic_truth(seg['epochs'], self.truth['epochs'])
                if not dynamic_truth_array:
                    QMessageBox.warning(self, "动态真值错误", f"分段 [{seg['name']}] 与动态真值时间无交集，无法完成指标计算。")
                    seg['epochs'] = []
                    seg['metrics'] = None
                    continue
            else:
                seg_truth = self.truth

            metrics, filtered_epochs = calculate_metrics(
                seg['epochs'], seg_truth,
                filter_outliers=self.app_config.get('filter_outliers', False),
                outlier_thresh=self.app_config.get('outlier_threshold', 50.0),
                dynamic_truth_array=dynamic_truth_array
            )
            seg['epochs'] = filtered_epochs
            seg['metrics'] = metrics

        # 联动检查X轴时间可用性并自动禁用
        self.check_xaxis_mode_availability()
        self.update_metrics_table()
        self.refresh_chart()

    def update_metrics_table(self):
        self.table_metrics.setRowCount(0)
        for seg in self.segments:
            m = seg['metrics']
            if not m:
                continue

            row = self.table_metrics.rowCount()
            self.table_metrics.insertRow(row)

            item_id = QTableWidgetItem(seg['name'])
            item_id.setForeground(QColor(seg['color']))

            align_rate = 100.0
            if m.get('original_count', 0) > 0:
                align_rate = (m['count'] / m['original_count']) * 100.0

            self.table_metrics.setItem(row, 0, item_id)
            self.table_metrics.setItem(row, 1, QTableWidgetItem(f"{align_rate:.1f}%"))
            self.table_metrics.setItem(row, 2, QTableWidgetItem(str(m['count'])))
            self.table_metrics.setItem(row, 3, QTableWidgetItem(f"{m['rtk_fix_rate']:.1f}%"))
            self.table_metrics.setItem(row, 4, QTableWidgetItem(f"{m['cep50']:.3f}"))
            self.table_metrics.setItem(row, 5, QTableWidgetItem(f"{m['cep68']:.3f}"))
            self.table_metrics.setItem(row, 6, QTableWidgetItem(f"{m['cep95']:.3f}"))
            self.table_metrics.setItem(row, 7, QTableWidgetItem(f"{m['rms_h']:.3f}"))
            self.table_metrics.setItem(row, 8, QTableWidgetItem(f"{m['rms_v']:.3f}"))
            self.table_metrics.setItem(row, 9, QTableWidgetItem(f"{m['max_h']:.3f}"))
            self.table_metrics.setItem(row, 10, QTableWidgetItem(f"{m['max_v']:.3f}"))

    def on_tab_changed(self, index):
        self.refresh_chart()

    def refresh_chart(self):
        index = self.tab_widget.currentIndex()
        if index == 0:
            self.canvas_scatter.render_data('scatter', self.segments, self.truth, self.time_zone)
        elif index == 1:
            self.canvas_epoch_h.render_data('epoch_h', self.segments, self.truth, self.time_zone, show_extrema=self.show_extrema, x_axis_mode=self.x_axis_mode, show_sats=self.cb_show_sats.isChecked())
        elif index == 2:
            self.canvas_epoch_v.render_data('epoch_v', self.segments, self.truth, self.time_zone, self.show_absolute_alt, show_extrema=self.show_extrema, x_axis_mode=self.x_axis_mode, show_sats=self.cb_show_sats.isChecked(), show_raw_alt=self.show_raw_alt)
        elif index == 3:
            self.canvas_status.render_data('status', self.segments, self.truth, self.time_zone)
        elif index == 4:
            self.canvas_trajectory.render_data('trajectory', self.segments, self.truth)

    # 8. 导出数据逻辑
    def on_export_raw_clicked(self):
        if not self.segments:
            return

        active_segs = [s for s in self.segments if s['active']]
        if not active_segs:
            QMessageBox.warning(self, "提示", "请在列表中勾选要导出的活跃分段。")
            return

        target_seg = active_segs[0]
        export_mode = self.app_config.get('export_dir_mode', '同源文件目录')
        default_dir = ""
        if export_mode == "记忆上次目录":
            default_dir = self.app_config.get('last_export_dir', "")
        if not default_dir and target_seg.get('file_id'):
            default_dir = os.path.dirname(target_seg['file_id'])

        default_path = os.path.join(default_dir, f"sliced_{target_seg['name']}.log")
        save_path, _ = QFileDialog.getSaveFileName(self, "导出截取原始数据", default_path, "GNSS Logs (*.log *.txt)")
        if not save_path:
            return

        if export_mode == "记忆上次目录":
            self.app_config['last_export_dir'] = os.path.dirname(save_path)
            self.save_config()

        leap_secs = self.get_leap_seconds()
        file_epochs = self.file_epochs_map.get(target_seg.get('file_id'), [])
        time_range = self._resolve_segment_time_range(file_epochs, target_seg['start_time'], target_seg['end_time'])
        if not time_range:
            start_sec = time_str_to_seconds(target_seg['start_time'])
            end_sec = time_str_to_seconds(target_seg['end_time'])
            time_range = (start_sec, end_sec + (86400 if start_sec > end_sec else 0))

        try:
            file_size = os.path.getsize(target_seg['file_id'])
            processed = 0
            stream_state = self._build_stream_time_state(file_epochs)

            from PySide6.QtWidgets import QProgressDialog, QApplication
            progress = QProgressDialog("正在导出原始数据...", "取消", 0, 100, self)
            progress.setWindowTitle("导出进度")
            progress.setWindowModality(Qt.WindowModal)
            progress.show()

            with open(target_seg['file_id'], 'r', encoding='utf-8', errors='ignore') as f_in, \
                 open(save_path, 'w', encoding='utf-8') as f_out:
                for line in f_in:
                    if progress.wasCanceled():
                        break

                    processed += len(line.encode('utf-8', errors='ignore'))
                    if processed % (1024 * 50) < 100:  # ~50KB update
                        pct = int((processed / file_size) * 100) if file_size > 0 else 0
                        progress.setValue(pct)
                        QApplication.processEvents()

                    if not line.startswith('$'):
                        continue

                    parts = line.strip().split(',')
                    stype = parts[0]
                    utc_time_sec = -1

                    if stype.endswith('GGA') and len(parts) > 1:
                        utc_time_sec = time_str_to_seconds(parts[1])
                    elif (stype == '$POGOS' or stype == '$PODRS') and len(parts) > 2:
                        try:
                            tow = float(parts[2])
                            utc_time_sec = time_str_to_seconds(gps_tow_to_utc_time(tow, leap_secs))
                        except ValueError:
                            pass

                    utc_time_sec = self._unwrap_stream_time(utc_time_sec, stream_state)
                    if utc_time_sec != -1 and time_range[0] <= utc_time_sec <= time_range[1]:
                        f_out.write(line)

            progress.setValue(100)

            QMessageBox.information(self, "导出成功", f"文件已成功保存到:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")

    def on_export_gga_clicked(self):
        active_segs = [s for s in self.segments if s['active']]
        if not active_segs:
            QMessageBox.warning(self, "提示", "请在列表中勾选要转换的活跃分段。")
            return

        target_seg = active_segs[0]
        export_mode = self.app_config.get('export_dir_mode', '同源文件目录')
        default_dir = ""
        if export_mode == "记忆上次目录":
            default_dir = self.app_config.get('last_export_dir', "")
        if not default_dir and target_seg.get('file_id'):
            default_dir = os.path.dirname(target_seg['file_id'])

        default_path = os.path.join(default_dir, f"converted_{target_seg['name']}.nmea")
        save_path, _ = QFileDialog.getSaveFileName(self, "格式转换导出为标准 NMEA GGA", default_path, "NMEA Files (*.nmea *.gga)")
        if not save_path:
            return

        if export_mode == "记忆上次目录":
            self.app_config['last_export_dir'] = os.path.dirname(save_path)
            self.save_config()

        leap_secs = self.get_leap_seconds()
        file_epochs = self.file_epochs_map.get(target_seg.get('file_id'), [])
        time_range = self._resolve_segment_time_range(file_epochs, target_seg['start_time'], target_seg['end_time'])
        if not time_range:
            start_sec = time_str_to_seconds(target_seg['start_time'])
            end_sec = time_str_to_seconds(target_seg['end_time'])
            time_range = (start_sec, end_sec + (86400 if start_sec > end_sec else 0))

        try:
            file_size = os.path.getsize(target_seg['file_id'])
            processed = 0
            stream_state = self._build_stream_time_state(file_epochs)

            from PySide6.QtWidgets import QProgressDialog, QApplication
            progress = QProgressDialog("正在转换导出 GGA...", "取消", 0, 100, self)
            progress.setWindowTitle("转换进度")
            progress.setWindowModality(Qt.WindowModal)
            progress.show()

            with open(target_seg['file_id'], 'r', encoding='utf-8', errors='ignore') as f_in, \
                 open(save_path, 'w', encoding='utf-8') as f_out:
                for line in f_in:
                    if progress.wasCanceled():
                        break

                    processed += len(line.encode('utf-8', errors='ignore'))
                    if processed % (1024 * 50) < 100:
                        pct = int((processed / file_size) * 100) if file_size > 0 else 0
                        progress.setValue(pct)
                        QApplication.processEvents()

                    if line.startswith('$POGOS'):
                        epoch = parse_log_line(line, leap_secs)
                        if epoch:
                            utc_time_sec = self._unwrap_stream_time(epoch['utc_time_sec'], stream_state)
                        else:
                            utc_time_sec = -1
                        if epoch and time_range[0] <= utc_time_sec <= time_range[1]:
                            gga = convert_pogos_to_gga(line, 'GN', leap_secs)
                            if gga:
                                f_out.write(gga + '\n')

            progress.setValue(100)

            QMessageBox.information(self, "导出成功", f"格式转换已成功保存到:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"转换保存失败: {str(e)}")

    def on_export_kml_clicked(self):
        active_segs = [s for s in self.segments if s.get('active', True) and s.get('epochs')]
        if not active_segs and not (self.truth_mode == 'dynamic' and self.truth.get('epochs')):
            QMessageBox.warning(self, "提示", "没有可导出的有效轨迹数据（请确保已加载测试分段或动态真值）。")
            return

        default_dir = self.app_config.get('last_export_dir', "")
        if not default_dir and active_segs and active_segs[0].get('file_id'):
            default_dir = os.path.dirname(active_segs[0]['file_id'])

        default_path = os.path.join(default_dir, "trajectory_export.kml")
        save_path, _ = QFileDialog.getSaveFileName(self, "导出 KML 轨迹", default_path, "KML Files (*.kml)")
        if not save_path:
            return

        self.app_config['last_export_dir'] = os.path.dirname(save_path)
        self.save_config()

        kml_header = '''<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>GNSS Trajectory Export</name>
'''
        kml_footer = '''  </Document>
</kml>
'''

        try:
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(kml_header)

                # 导出动态真值
                if self.truth_mode == 'dynamic' and self.truth.get('epochs'):
                    f.write('    <Folder>\n      <name>Dynamic Truth</name>\n')
                    # 1. 写入真值连续轨迹线 (LineString)
                    f.write('      <Placemark>\n')
                    f.write('        <name>Dynamic Truth (真值轨迹线)</name>\n')
                    f.write('        <Style>\n          <LineStyle>\n            <color>ff000000</color>\n            <width>3</width>\n          </LineStyle>\n        </Style>\n')
                    f.write('        <LineString>\n')
                    f.write('          <altitudeMode>absolute</altitudeMode>\n')
                    f.write('          <coordinates>\n')
                    for ep in self.truth['epochs']:
                        if ep['lat'] == 0.0 or ep['lon'] == 0.0 or math.isnan(ep['lat']) or math.isnan(ep['lon']) or math.isnan(ep.get('alt', 0.0)):
                            continue
                        f.write(f'            {ep["lon"]},{ep["lat"]},{ep.get("alt", 0.0)}\n')
                    f.write('          </coordinates>\n')
                    f.write('        </LineString>\n')
                    f.write('      </Placemark>\n')

                    # 2. 写入稀疏化的位置标志点 (Placemark)
                    truth_count = len(self.truth['epochs'])
                    step_t = max(1, truth_count // 1000)
                    for i in range(0, truth_count, step_t):
                        ep = self.truth['epochs'][i]
                        if ep['lat'] == 0.0 or ep['lon'] == 0.0 or math.isnan(ep['lat']) or math.isnan(ep['lon']) or math.isnan(ep.get('alt', 0.0)):
                            continue
                        is_key = (i == 0 or i >= truth_count - step_t)
                        label_scale = 1.0 if is_key else 0.0
                        pt_name = "真值起点" if i == 0 else ("真值终点" if i >= truth_count - step_t else f"真值点 {i}")

                        f.write('      <Placemark>\n')
                        f.write(f'        <name>{pt_name}</name>\n')
                        f.write(f'        <Style>\n          <LabelStyle>\n            <scale>{label_scale}</scale>\n          </LabelStyle>\n')
                        f.write('          <IconStyle>\n            <color>ff000000</color>\n            <scale>0.4</scale>\n          </IconStyle>\n        </Style>\n')
                        time_str = ep.get('time_str', '')
                        if time_str and len(time_str) >= 6:
                            clean_time = time_str.replace(':', '')
                            if len(clean_time) >= 6:
                                iso_time = f"2024-01-01T{clean_time[:2]}:{clean_time[2:4]}:{clean_time[4:6]}Z"
                                f.write(f'        <TimeStamp><when>{iso_time}</when></TimeStamp>\n')
                        f.write('        <Point>\n')
                        f.write(f'          <coordinates>{ep["lon"]},{ep["lat"]},{ep.get("alt", 0.0)}</coordinates>\n')
                        f.write('        </Point>\n')
                        f.write('      </Placemark>\n')
                    f.write('    </Folder>\n')

                # 导出各个测试分段
                for seg in active_segs:
                    color_hex = seg['color'].lstrip('#')
                    if len(color_hex) == 6:
                        # KML 颜色为 aabbggrr，Matplotlib 为 #rrggbb
                        r, g, b = color_hex[0:2], color_hex[2:4], color_hex[4:6]
                        kml_color = f"ff{b}{g}{r}"
                    else:
                        kml_color = "ffff0000"

                    f.write(f'    <Folder>\n      <name>{seg["name"]}</name>\n')

                    # 1. 写入测试分段轨迹线 (LineString)
                    f.write('      <Placemark>\n')
                    f.write(f'        <name>{seg["name"]} (轨迹线)</name>\n')
                    f.write(f'        <Style>\n          <LineStyle>\n            <color>{kml_color}</color>\n            <width>3</width>\n          </LineStyle>\n        </Style>\n')
                    f.write('        <LineString>\n')
                    f.write('          <altitudeMode>absolute</altitudeMode>\n')
                    f.write('          <coordinates>\n')
                    for ep in seg['epochs']:
                        if ep['lat'] == 0.0 or ep['lon'] == 0.0 or math.isnan(ep['lat']) or math.isnan(ep['lon']) or math.isnan(ep.get('alt', 0.0)):
                            continue
                        f.write(f'            {ep["lon"]},{ep["lat"]},{ep.get("alt", 0.0)}\n')
                    f.write('          </coordinates>\n')
                    f.write('        </LineString>\n')
                    f.write('      </Placemark>\n')

                    # 2. 写入稀疏化的位置标志点 (Placemark)
                    seg_count = len(seg['epochs'])
                    step_s = max(1, seg_count // 1000)
                    for i in range(0, seg_count, step_s):
                        ep = seg['epochs'][i]
                        if ep['lat'] == 0.0 or ep['lon'] == 0.0 or math.isnan(ep['lat']) or math.isnan(ep['lon']) or math.isnan(ep.get('alt', 0.0)):
                            continue
                        is_key = (i == 0 or i >= seg_count - step_s)
                        label_scale = 1.0 if is_key else 0.0
                        pt_name = f"{seg['name']}_起点" if i == 0 else (f"{seg['name']}_终点" if i >= seg_count - step_s else f"点 {i}")

                        f.write('      <Placemark>\n')
                        f.write(f'        <name>{pt_name}</name>\n')
                        f.write(f'        <Style>\n          <LabelStyle>\n            <scale>{label_scale}</scale>\n          </LabelStyle>\n')
                        f.write(f'          <IconStyle>\n            <color>{kml_color}</color>\n            <scale>0.4</scale>\n          </IconStyle>\n        </Style>\n')
                        time_str = ep.get('time_str', '')
                        if time_str and len(time_str) >= 6:
                            clean_time = time_str.replace(':', '')
                            if len(clean_time) >= 6:
                                iso_time = f"2024-01-01T{clean_time[:2]}:{clean_time[2:4]}:{clean_time[4:6]}Z"
                                f.write(f'        <TimeStamp><when>{iso_time}</when></TimeStamp>\n')
                        f.write('        <Point>\n')
                        f.write(f'          <coordinates>{ep["lon"]},{ep["lat"]},{ep.get("alt", 0.0)}</coordinates>\n')
                        f.write('        </Point>\n')
                        f.write('      </Placemark>\n')
                    f.write('    </Folder>\n')

                f.write(kml_footer)
            QMessageBox.information(self, "导出成功", f"KML 轨迹已成功导出至:\n{save_path}")
        except PermissionError:
            QMessageBox.critical(self, "错误", "目标文件被占用！请确认是否在 Google Earth 中打开了该文件，关闭后再试。")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"KML 导出失败: {str(e)}")

    def on_export_report_clicked(self):
        try:
            import docx
            from docx.shared import Inches
            from datetime import datetime
        except ImportError:
            QMessageBox.warning(self, "提示", "未检测到 python-docx 模块。请在终端运行: pip install python-docx")
            return

        if not self.segments:
            QMessageBox.warning(self, "提示", "没有测试数据可供导出。")
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "导出 Word 报告", "测试报告.docx", "Word Documents (*.docx)")
        if not save_path:
            return

        try:
            from PySide6.QtWidgets import QProgressDialog, QApplication
            progress = QProgressDialog("正在生成 Word 报告...", "取消", 0, 8, self)
            progress.setWindowTitle("导出报告")
            progress.setWindowModality(Qt.WindowModal)
            progress.setMinimumDuration(0)
            progress.setValue(0)

            doc = docx.Document()
            doc.add_heading('GNSS 定位精度测试报告', 0)

            doc.add_heading('1. 测试概览', level=1)
            doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            progress.setValue(1)
            QApplication.processEvents()
            if progress.wasCanceled():
                return

            # 添加表格数据 (精度指标)
            doc.add_heading('2. 精度指标对比', level=1)
            table = doc.add_table(rows=1, cols=self.table_metrics.columnCount())
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for j in range(self.table_metrics.columnCount()):
                hdr_cells[j].text = self.table_metrics.horizontalHeaderItem(j).text()

            for i in range(self.table_metrics.rowCount()):
                row_cells = table.add_row().cells
                for j in range(self.table_metrics.columnCount()):
                    item = self.table_metrics.item(i, j)
                    row_cells[j].text = item.text() if item else ""
            progress.setValue(2)
            QApplication.processEvents()
            if progress.wasCanceled():
                return

            # 导出图表截图
            doc.add_heading('3. 图表分析', level=1)
            import io

            def add_plot_to_doc(canvas, title, progress_value):
                doc.add_heading(title, level=2)
                buf = io.BytesIO()
                canvas.figure.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                buf.seek(0)
                doc.add_picture(buf, width=Inches(6.0))
                buf.close()
                progress.setValue(progress_value)
                QApplication.processEvents()

            # 强制渲染所有图表，防止用户未点击的选项卡截图为空白
            progress.setLabelText("正在渲染图表...")
            self.canvas_scatter.render_data('scatter', self.segments, self.truth, self.time_zone)
            self.canvas_trajectory.render_data('trajectory', self.segments, self.truth)
            self.canvas_status.render_data('status', self.segments, self.truth, self.time_zone)
            self.canvas_epoch_h.render_data('epoch_h', self.segments, self.truth, self.time_zone, show_extrema=self.show_extrema, x_axis_mode=self.x_axis_mode, show_sats=self.cb_show_sats.isChecked())
            self.canvas_epoch_v.render_data('epoch_v', self.segments, self.truth, self.time_zone, self.show_absolute_alt, show_extrema=self.show_extrema, x_axis_mode=self.x_axis_mode, show_sats=self.cb_show_sats.isChecked(), show_raw_alt=self.show_raw_alt)
            progress.setValue(3)
            QApplication.processEvents()
            if progress.wasCanceled():
                return

            progress.setLabelText("正在写入图表...")
            add_plot_to_doc(self.canvas_trajectory, '3.1 绝对二维轨迹投影图', 4)
            if progress.wasCanceled():
                return
            add_plot_to_doc(self.canvas_scatter, '3.2 定位偏差分布图 (靶心图)', 5)
            if progress.wasCanceled():
                return
            add_plot_to_doc(self.canvas_status, '3.3 定位解状态分布', 6)
            if progress.wasCanceled():
                return
            add_plot_to_doc(self.canvas_epoch_h, '3.4 水平位置误差分布', 7)
            if progress.wasCanceled():
                return
            add_plot_to_doc(self.canvas_epoch_v, '3.5 高程位置误差分布', 8)
            if progress.wasCanceled():
                return

            progress.setLabelText("正在保存报告...")
            doc.save(save_path)
            progress.close()
            QMessageBox.information(self, "导出成功", f"Word 报告已成功导出至:\n{save_path}")
        except PermissionError:
            QMessageBox.critical(self, "错误", "目标 Word 文件正被另一程序(如 Microsoft Word)打开，请关闭后重试。")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"报告生成失败: {str(e)}")

    def refresh_serial_ports(self):
        self.cmb_port.clear()
        ports = QSerialPortInfo.availablePorts()
        for p in ports:
            display_name = f"{p.portName()} ({p.description()})" if p.description() else p.portName()
            self.cmb_port.addItem(display_name, p.portName())

        # 设置下拉菜单最小宽度为 320px 保证长串口名字完整可见
        self.cmb_port.view().setMinimumWidth(320)

        if self.cmb_port.count() == 0:
            self.lbl_status.setText("未检测到可用串口")
        else:
            self.lbl_status.setText(f"已扫描到 {self.cmb_port.count()} 个可用串口")

    def toggle_serial_connection(self):
        if self.serial_port.isOpen():
            self.serial_port.close()
            if hasattr(self, 'realtime_timeout_timer'):
                self.realtime_timeout_timer.stop()
            if hasattr(self, 'cno_refresh_timer'):
                self.cno_refresh_timer.stop()
            self.btn_serial_connect.setText("打开串口")
            self.btn_serial_connect.setStyleSheet("""
                QPushButton {
                    background-color: rgba(16, 185, 129, 0.15);
                    border: 1px solid #10B981;
                    color: #10B981;
                    font-weight: bold;
                    font-size: 12px;
                }
                QPushButton:hover {
                    background-color: rgba(16, 185, 129, 0.25);
                    color: #FFFFFF;
                    font-size: 12px;
                }
                QPushButton:pressed {
                    background-color: rgba(16, 185, 129, 0.35);
                    font-size: 12px;
                }
            """)
            self.lbl_status.setText("串口已关闭")
            if self.record_file:
                try:
                    self.record_file.close()
                except Exception:
                    pass
                self.record_file = None
                self.cb_record.setChecked(False)
                self.record_filepath = None
        else:
            port_name = self.cmb_port.currentData()
            if not port_name:
                port_name = self.cmb_port.currentText()
            if not port_name:
                QMessageBox.warning(self, "警告", "没有选择有效的串口！")
                return
            baud_rate = int(self.cmb_baud.currentText())
            self.serial_port.setPortName(port_name)
            self.serial_port.setBaudRate(baud_rate)
            if self.serial_port.open(QSerialPort.ReadWrite):
                self.btn_serial_connect.setText("关闭串口")
                self.btn_serial_connect.setStyleSheet("""
                    QPushButton {
                        background-color: rgba(239, 68, 68, 0.15);
                        border: 1px solid #EF4444;
                        color: #EF4444;
                        font-weight: bold;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: rgba(239, 68, 68, 0.25);
                        color: #FFFFFF;
                        font-size: 12px;
                    }
                    QPushButton:pressed {
                        background-color: rgba(239, 68, 68, 0.35);
                        font-size: 12px;
                    }
                """)
                self.lbl_status.setText(f"已连接串口 {port_name}，波特率 {baud_rate}...")
                self.add_realtime_segment_item()
                self.reset_live_status_ui()
                if hasattr(self, 'cno_refresh_timer'):
                    self.cno_refresh_timer.start(1000)
            else:
                QMessageBox.critical(self, "连接失败", f"无法打开串口 {port_name}，可能已被占用或未连接。")

    def on_record_state_changed(self, state):
        if state == 2:  # Checked (Qt.Checked is 2)
            filepath, _ = QFileDialog.getSaveFileName(self, "选择保存的原始数据日志文件", "", "GNSS Logs (*.log *.txt *.nmea *.dat)")
            if not filepath:
                self.cb_record.setChecked(False)
                return
            try:
                self.record_file = open(filepath, 'wb')
                self.record_filepath = filepath
                self.record_error_reported = False
                self.safe_append_console(f"[录制启动] 正在将原始流实时保存到: {filepath}\n")
            except Exception as e:
                QMessageBox.critical(self, "启动录制失败", f"无法写入目标文件: {str(e)}")
                self.cb_record.setChecked(False)
        else:
            if self.record_file:
                try:
                    self.record_file.close()
                except Exception:
                    pass
                self.record_file = None
                self.safe_append_console(f"[录制停止] 文件已成功保存: {self.record_filepath}\n")
                self.record_filepath = None
                self.record_error_reported = False

    def stop_recording_with_error(self, err_msg):
        if self.record_error_reported:
            return
        self.record_error_reported = True

        failed_path = self.record_filepath
        if self.record_file:
            try:
                self.record_file.close()
            except Exception:
                pass
        self.record_file = None
        self.record_filepath = None

        if hasattr(self, 'cb_record'):
            self.cb_record.blockSignals(True)
            self.cb_record.setChecked(False)
            self.cb_record.blockSignals(False)

        self.lbl_status.setText("录制失败，已停止保存原始流")
        self.safe_append_console(f"[录制失败] 已停止保存原始流: {failed_path or 'unknown'}\n错误: {err_msg}\n")
        QMessageBox.warning(self, "录制失败", f"写入原始数据日志失败，录制已停止。\n\n{err_msg}")

    def safe_append_console(self, text, scroll=True):
        scrollbar = self.txt_console.verticalScrollBar()
        old_val = scrollbar.value()
        self._is_programmatic_scroll = True
        try:
            self.txt_console.append(text)
            if not self.cb_scroll.isChecked():
                scrollbar.setValue(old_val)
        finally:
            self._is_programmatic_scroll = False
        if scroll:
            self.request_console_scroll_to_bottom()

    def safe_clear_console(self):
        self._is_programmatic_scroll = True
        try:
            self.txt_console.clear()
        finally:
            self._is_programmatic_scroll = False

    def scroll_console_to_bottom(self):
        if self.cb_scroll.isChecked():
            self._is_programmatic_scroll = True
            self.txt_console.moveCursor(QTextCursor.End)
            self._is_programmatic_scroll = False

    def request_console_scroll_to_bottom(self):
        if self.cb_scroll.isChecked() and hasattr(self, 'console_scroll_timer'):
            self.console_scroll_timer.start(50)

    def on_console_scrollbar_value_changed(self, value):
        if self._is_programmatic_scroll:
            return

        scrollbar = self.txt_console.verticalScrollBar()
        max_val = scrollbar.maximum()

        if max_val - value > 10:
            if self.cb_scroll.isChecked():
                self.cb_scroll.blockSignals(True)
                self.cb_scroll.setChecked(False)
                self.cb_scroll.blockSignals(False)
        elif max_val > 0 and max_val - value <= 10:
            if not self.cb_scroll.isChecked():
                self.cb_scroll.blockSignals(True)
                self.cb_scroll.setChecked(True)
                self.cb_scroll.blockSignals(False)

    def on_scroll_checkbox_changed(self, state):
        if state == 2 or state == Qt.Checked:
            self.scroll_console_to_bottom()

    def update_send_ln_state(self):
        from PySide6.QtWidgets import QGraphicsOpacityEffect
        is_hex = self.cb_send_hex.isChecked()
        self.cb_send_ln.setEnabled(not is_hex)
        if is_hex:
            effect = QGraphicsOpacityEffect(self.cb_send_ln)
            effect.setOpacity(0.4)
            self.cb_send_ln.setGraphicsEffect(effect)
        else:
            self.cb_send_ln.setGraphicsEffect(None)

    def send_custom_data(self):
        if not self.serial_port.isOpen():
            QMessageBox.warning(self, "警告", "请先打开串口连接！")
            return

        text = self.txt_send_input.text()
        if not text:
            return

        if self.cb_send_hex.isChecked():
            hex_cleaned = "".join(text.split())
            try:
                data = bytes.fromhex(hex_cleaned)
            except ValueError:
                QMessageBox.warning(self, "错误", "非法的十六进制数据格式！请确保输入只包含 0-9, a-f, A-F 等字符（可含空格）。")
                return
        else:
            if self.cb_send_ln.isChecked():
                text += "\r\n"
            data = text.encode('utf-8', errors='ignore')

        try:
            self.serial_port.write(data)

            if self.cb_send_hex.isChecked():
                echo_str = f">> [HEX] {' '.join(f'{b:02X}' for b in data)}\n"
            else:
                echo_str = f">> {text}"
                if not echo_str.endswith('\n'):
                    echo_str += '\n'

            self.safe_append_console(echo_str)
            self.txt_send_input.clear()
        except Exception as e:
            QMessageBox.critical(self, "错误", f"数据发送失败: {e}")

    def on_mode_tab_changed(self, index):
        if index == 1:  # 日志回放 Tab
            if self.serial_port.isOpen():
                self.toggle_serial_connection()
        else:  # 串口连接 Tab
            if self.is_replaying:
                self.pause_replay()

    def on_replay_browse(self):
        filepath, _ = QFileDialog.getOpenFileName(
            self, "选择回放日志文件", "", "GNSS Logs (*.log *.txt *.nmea *.dat)"
        )
        if filepath:
            self.load_replay_file(filepath)

    def clear_replay_data(self):
        self.stop_replay()
        self.replay_blocks = []
        self.replay_filepath = None
        self.replay_snapshots.clear()
        self.txt_replay_file.clear()
        self.btn_replay_play.setEnabled(False)
        self.btn_replay_stop.setEnabled(False)
        self.btn_replay_clear.setEnabled(False)
        self.slider_replay.setEnabled(False)
        self.slider_replay.setValue(0)
        self.lbl_replay_time.setText("00:00:00 / 00:00:00")

        # 清除所有解析数据缓存并重绘为空白
        self.gsv_satellites.clear()
        self.used_satellites.clear()
        self.sat_metadata.clear()
        self.has_received_gsa = False
        self.realtime_raw_epochs.clear()
        self.parsed_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') != "COM_REALTIME"]

        self.reset_live_status_ui()
        self.recompute_all()
        self.update_cno_chart()

    def capture_replay_snapshot(self, index):
        if index < 0:
            return
        self.replay_snapshots[index] = {
            'serial_buffer': copy.deepcopy(self.serial_buffer),
            'gsv_satellites': copy.deepcopy(self.gsv_satellites),
            'used_satellites': set(self.used_satellites),
            'sat_metadata': copy.deepcopy(self.sat_metadata),
            'has_received_gsa': self.has_received_gsa,
            'realtime_raw_epochs': copy.deepcopy(list(self.realtime_raw_epochs)),
            'parsed_realtime_epochs': copy.deepcopy([ep for ep in self.parsed_epochs if ep.get('file_id') == "COM_REALTIME"]),
            'latest_quality': self.latest_quality,
            'latest_num_sats': self.latest_num_sats,
            'latest_hdop': self.latest_hdop,
            'latest_pdop': self.latest_pdop
        }

    def restore_replay_snapshot(self, index):
        snapshot = self.replay_snapshots.get(index)
        if not snapshot:
            return False

        self.serial_buffer = copy.deepcopy(snapshot['serial_buffer'])
        self.gsv_satellites = copy.deepcopy(snapshot['gsv_satellites'])
        self.used_satellites = set(snapshot['used_satellites'])
        self.sat_metadata = copy.deepcopy(snapshot['sat_metadata'])
        self.has_received_gsa = snapshot['has_received_gsa']
        self.realtime_raw_epochs.clear()
        self.realtime_raw_epochs.extend(snapshot['realtime_raw_epochs'])
        self.parsed_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') != "COM_REALTIME"]
        self.parsed_epochs.extend(copy.deepcopy(snapshot['parsed_realtime_epochs']))
        self.latest_quality = snapshot['latest_quality']
        self.latest_num_sats = snapshot['latest_num_sats']
        self.latest_hdop = snapshot['latest_hdop']
        self.latest_pdop = snapshot['latest_pdop']
        self._last_cno_snapshot = None
        return True

    def read_replay_block(self, index):
        if not self.replay_filepath or index < 0 or index >= len(self.replay_blocks):
            return b''

        with open(self.replay_filepath, 'rb') as f:
            return self._read_replay_block_from_file(f, index)

    def _read_replay_block_from_file(self, file_obj, index):
        _, offset, length = self.replay_blocks[index]
        if length <= 0:
            return b''

        file_obj.seek(offset)
        return file_obj.read(length)

    def load_replay_file(self, filepath):
        import os
        from PySide6.QtWidgets import QProgressDialog

        if not os.path.exists(filepath):
            QMessageBox.critical(self, "错误", "目标日志文件不存在！")
            return

        file_size = os.path.getsize(filepath)
        if file_size == 0:
            QMessageBox.warning(self, "警告", "选定的日志文件大小为 0 字节，无法回放！")
            return

        # 预分析特征检出
        nmea_count = 0
        bk_count = 0
        try:
            with open(filepath, 'rb') as f:
                header_data = f.read(50000)
                nmea_count = header_data.count(b'$')
                bk_count = 0
                idx = 0
                while idx < len(header_data) - 1:
                    if header_data[idx] == 0x42 and header_data[idx+1] == 0x4B:
                        bk_count += 1
                        idx += 8
                    else:
                        idx += 1
        except Exception as e:
            QMessageBox.critical(self, "错误", f"读取日志特征失败: {e}")
            return

        if nmea_count == 0 and bk_count == 0:
            QMessageBox.warning(self, "错误", "未能在文件中检测到任何符合规范的 NMEA 或博通二进制协议数据特征！请检查选择的文件。")
            return

        self.stop_replay()
        self.reset_live_status_ui()
        self.serial_buffer = BKStreamParser()
        self.gsv_satellites.clear()
        self.used_satellites.clear()
        self.sat_metadata.clear()
        self.has_received_gsa = False
        self.realtime_raw_epochs.clear()
        self.parsed_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') != "COM_REALTIME"]
        self.recompute_all()
        self.update_cno_chart()
        self.replay_blocks = []
        self.replay_filepath = filepath
        self.replay_snapshots.clear()
        self.replay_index = 0

        progress = QProgressDialog("正在分析日志并切分时间块...", "取消", 0, 100, self)
        progress.setWindowTitle("载入中")
        progress.setWindowModality(Qt.WindowModal)
        progress.setMinimumDuration(500)

        def extract_time_str(line):
            if not line.startswith('$'):
                return None
            parts = line.split(',')
            if len(parts) > 1:
                stype = parts[0]
                if 'GGA' in stype or 'RMC' in stype:
                    t_field = parts[1]
                    if len(t_field) >= 6:
                        return f"{t_field[0:2]}:{t_field[2:4]}:{t_field[4:6]}"
            return None

        try:
            is_pure_binary = (nmea_count == 0 and bk_count > 0)
            blocks = []

            if is_pure_binary:
                with open(filepath, 'rb') as f:
                    chunk_idx = 0
                    while True:
                        if progress.wasCanceled():
                            break
                        block_offset = f.tell()
                        chunk = f.read(2048)
                        if not chunk:
                            break
                        t_str = f"Offset:{chunk_idx * 2}KB"
                        blocks.append((t_str, block_offset, len(chunk)))
                        chunk_idx += 1
                        progress.setValue(min(int((f.tell() / file_size) * 100), 99))
            else:
                current_time = "00:00:00"
                current_block_start = None
                current_block_len = 0

                with open(filepath, 'rb') as f:
                    line_idx = 0
                    while True:
                        if progress.wasCanceled():
                            break
                        line_offset = f.tell()
                        line_bytes = f.readline()
                        if not line_bytes:
                            break

                        if line_idx % 2000 == 0:
                            progress.setValue(min(int((f.tell() / file_size) * 100), 99))
                        line_idx += 1

                        line_str = line_bytes.decode('utf-8', errors='replace')
                        t = extract_time_str(line_str)
                        if t is not None and t != current_time:
                            if current_block_start is not None and current_block_len > 0:
                                blocks.append((current_time, current_block_start, current_block_len))
                            current_time = t
                            current_block_start = line_offset
                            current_block_len = len(line_bytes)
                        else:
                            if current_block_start is None:
                                current_block_start = line_offset
                            current_block_len += len(line_bytes)

                    if current_block_start is not None and current_block_len > 0 and not progress.wasCanceled():
                        blocks.append((current_time, current_block_start, current_block_len))

            progress.setValue(100)

            if progress.wasCanceled():
                self.replay_blocks = []
                self.replay_filepath = None
                self.replay_snapshots.clear()
                self.txt_replay_file.clear()
                self.btn_replay_play.setEnabled(False)
                self.btn_replay_stop.setEnabled(False)
                self.btn_replay_clear.setEnabled(False)
                self.slider_replay.setEnabled(False)
                self.lbl_replay_time.setText("00:00:00 / 00:00:00")
                return

            self.replay_blocks = blocks
            if not self.replay_blocks:
                self.replay_filepath = None
                self.replay_snapshots.clear()
                QMessageBox.warning(self, "警告", "日志文件中未切分出可回放的数据块。")
                return

            self.txt_replay_file.setText(filepath)
            self.btn_replay_play.setEnabled(True)
            self.btn_replay_stop.setEnabled(True)
            self.btn_replay_clear.setEnabled(True)
            self.slider_replay.setEnabled(True)
            self.slider_replay.setRange(0, len(self.replay_blocks) - 1)
            self.slider_replay.setValue(0)

            # 导入后立即解析并重绘第1个分包以显示初始数据
            if self.replay_blocks:
                block_bytes = self.read_replay_block(0)
                self.parse_raw_chunk(block_bytes)
                self.recompute_all()
                self.update_cno_chart()

            self.update_replay_time_display()

        except Exception as e:
            self.replay_blocks = []
            self.replay_filepath = None
            self.replay_snapshots.clear()
            QMessageBox.critical(self, "错误", f"日志预分包失败: {e}")

    def update_replay_time_display(self):
        if not self.replay_blocks:
            self.lbl_replay_time.setText("00:00:00 / 00:00:00")
            return
        curr_time = self.replay_blocks[self.replay_index][0]
        total_time = self.replay_blocks[-1][0]
        self.lbl_replay_time.setText(f"{curr_time} / {total_time}")

    def reset_replay_live_state(self):
        self.serial_buffer = BKStreamParser()
        self.gsv_satellites.clear()
        self.used_satellites.clear()
        self.sat_metadata.clear()
        self.has_received_gsa = False
        self.realtime_raw_epochs.clear()
        self.parsed_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') != "COM_REALTIME"]
        self._last_cno_snapshot = None
        if hasattr(self, 'canvas_cno'):
            self.canvas_cno.render_cno(self.gsv_satellites, self.used_satellites, self.has_received_gsa, self.sat_metadata)

    def parse_replay_blocks_until(self, target_index):
        if not self.replay_blocks or not self.replay_filepath:
            return False

        target_index = max(0, min(target_index, len(self.replay_blocks) - 1))
        snapshot_candidates = [idx for idx in self.replay_snapshots if idx <= target_index]
        snapshot_index = max(snapshot_candidates) if snapshot_candidates else None
        restored_snapshot = snapshot_index is not None and self.restore_replay_snapshot(snapshot_index)
        if restored_snapshot:
            start_index = snapshot_index + 1
        else:
            self.reset_replay_live_state()
            start_index = 0

        has_new_epoch = restored_snapshot
        progress = None
        blocks_to_parse = target_index - start_index + 1
        if blocks_to_parse >= 200:
            from PySide6.QtWidgets import QProgressDialog, QApplication
            progress = QProgressDialog("正在定位回放进度...", "取消", 0, max(1, blocks_to_parse), self)
            progress.setWindowTitle("回放定位")
            progress.setWindowModality(Qt.WindowModal)
            progress.setMinimumDuration(300)
            progress.setValue(0)

        cancelled = False
        completed_index = snapshot_index if snapshot_index is not None else -1
        self._disable_console_append = True
        try:
            with open(self.replay_filepath, 'rb') as f:
                for step_idx, idx in enumerate(range(start_index, target_index + 1)):
                    if progress and progress.wasCanceled():
                        cancelled = True
                        break
                    block_bytes = self._read_replay_block_from_file(f, idx)
                    if self.parse_raw_chunk(block_bytes):
                        has_new_epoch = True
                    completed_index = idx
                    if idx % self.replay_snapshot_interval == 0 or idx == target_index:
                        self.capture_replay_snapshot(idx)
                    if progress and (step_idx % 100 == 0 or idx == target_index):
                        progress.setValue(step_idx + 1)
                        QApplication.processEvents()
        finally:
            self._disable_console_append = False
            if progress:
                progress.close()
        if cancelled:
            self.replay_index = max(0, completed_index)
            self.slider_replay.blockSignals(True)
            self.slider_replay.setValue(self.replay_index)
            self.slider_replay.blockSignals(False)

        return has_new_epoch

    def toggle_replay_playback(self):
        if not self.replay_blocks:
            return
        if self.is_replaying:
            self.pause_replay()
        else:
            self.start_replay()

    def start_replay(self):
        if not self.replay_blocks or self.is_replaying:
            return
        if self.replay_index >= len(self.replay_blocks) - 1:
            self.replay_index = 0
            self.slider_replay.setValue(0)

        self.is_replaying = True
        self.btn_replay_play.setText("暂停")
        self.btn_replay_play.setStyleSheet("""
            QPushButton {
                background-color: rgba(245, 158, 11, 0.15);
                border: 1px solid #F59E0B;
                color: #F59E0B;
                font-weight: bold;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: rgba(245, 158, 11, 0.25);
                color: #FFFFFF;
            }
        """)
        import time
        self.replay_start_time = time.time()
        self.replay_start_index = self.replay_index
        self.replay_timer.start(50)

    def pause_replay(self):
        if not self.is_replaying:
            return
        self.is_replaying = False
        self.replay_timer.stop()
        self.btn_replay_play.setText("播放")
        self.btn_replay_play.setStyleSheet("""
            QPushButton {
                background-color: rgba(56, 189, 248, 0.15);
                border: 1px solid #38BDF8;
                color: #38BDF8;
                font-weight: bold;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: rgba(56, 189, 248, 0.25);
                color: #FFFFFF;
            }
        """)

    def stop_replay(self):
        self.pause_replay()
        self.replay_index = 0
        if self.replay_blocks:
            self.slider_replay.setValue(0)
            self.update_replay_time_display()

    def on_replay_speed_changed(self, text):
        if self.is_replaying:
            import time
            self.replay_start_time = time.time()
            self.replay_start_index = self.replay_index

    def get_replay_speed_multiplier(self):
        txt = self.cmb_replay_speed.currentText()
        try:
            return float(txt.replace('x', ''))
        except ValueError:
            return 1.0

    def replay_tick(self):
        if not self.is_replaying or not self.replay_blocks or not self.replay_filepath:
            self.replay_timer.stop()
            return

        import time
        multiplier = self.get_replay_speed_multiplier()
        elapsed = (time.time() - self.replay_start_time) * multiplier
        target_index = self.replay_start_index + int(elapsed)

        target_index = min(target_index, len(self.replay_blocks) - 1)

        if target_index > self.replay_index:
            blocks_to_process = target_index - self.replay_index
            blocks_to_process = min(blocks_to_process, 50)

            has_new_epoch = False
            if multiplier >= 5.0:
                self._disable_console_append = True
            try:
                with open(self.replay_filepath, 'rb') as f:
                    for i in range(blocks_to_process):
                        curr_idx = self.replay_index + 1
                        if curr_idx >= len(self.replay_blocks):
                            break
                        self.replay_index = curr_idx
                        block_bytes = self._read_replay_block_from_file(f, self.replay_index)

                        has_epoch = self.parse_raw_chunk(block_bytes)
                        if has_epoch:
                            has_new_epoch = True
                        if self.replay_index % self.replay_snapshot_interval == 0:
                            self.capture_replay_snapshot(self.replay_index)
            finally:
                if multiplier >= 5.0:
                    self._disable_console_append = False

            self.slider_replay.blockSignals(True)
            self.slider_replay.setValue(self.replay_index)
            self.slider_replay.blockSignals(False)
            self.update_replay_time_display()

            cur_time = time.time()
            if cur_time - self.last_recompute_time >= 0.5:
                if has_new_epoch:
                    self.recompute_all()
                self.update_cno_chart()
                self.last_recompute_time = cur_time

            if self.replay_index >= len(self.replay_blocks) - 1:
                self.stop_replay()

    def on_slider_pressed(self):
        self.is_slider_dragging = True
        if self.is_replaying:
            self.replay_timer.stop()

    def on_slider_released(self):
        self.is_slider_dragging = False
        self.replay_index = self.slider_replay.value()
        self.reset_replay_live_state()
        has_new_epoch = self.parse_replay_blocks_until(self.replay_index)

        self.update_replay_time_display()
        if has_new_epoch:
            self.recompute_all()
        self.update_cno_chart()

        if self.is_replaying:
            import time
            self.replay_start_time = time.time()
            self.replay_start_index = self.replay_index
            self.replay_timer.start(50)

    def on_slider_value_changed(self, value):
        if self.is_slider_dragging:
            self.replay_index = value
            self.update_replay_time_display()

    def clear_serial_console(self):
        self.safe_clear_console()

    def send_serial_command(self, cmd_type):
        if not self.serial_port.isOpen():
            QMessageBox.warning(self, "警告", "请先打开串口连接！")
            return

        if cmd_type == 'cold':
            # 明文冷启动
            self.serial_port.write(b"$POLCFGRESET,1\r\n")
            self.safe_append_console("[下发指令] 冷启动复位\n")
            self.reset_live_status_ui()

        elif cmd_type == 'hot':
            # 明文热启动
            self.serial_port.write(b"$POLCFGRESET,0\r\n")
            self.safe_append_console("[下发指令] 热启动复位\n")
            self.reset_live_status_ui()

        elif cmd_type == 'version':
            self.waiting_for_version = True
            self.version_lines = []
            QTimer.singleShot(3000, self.reset_version_wait_flag)
            self.serial_port.write(b"$POLCFGPTVER\r\n")
            header_data = bytes([0x02, 0x07, 0x00, 0x00])
            crc = crc16_ccitt(header_data)
            bk_cmd = bytes([0x42, 0x4B, (crc >> 8) & 0xFF, crc & 0xFF]) + header_data
            self.serial_port.write(bk_cmd)
            self.safe_append_console("[下发指令] 查询版本号\n")

        elif cmd_type == 'save':
            self.serial_port.write(b"$POLCFGSAVE\r\n")
            header_data = bytes([0x02, 0x26, 0x00, 0x04])
            crc = crc16_ccitt(header_data)
            bk_cmd = bytes([0x42, 0x4B, (crc >> 8) & 0xFF, crc & 0xFF]) + header_data
            self.serial_port.write(bk_cmd)
            self.safe_append_console("[下发指令] 保存配置到 Flash\n")

    def add_realtime_segment_item(self):
        # 避免重复创建
        for s in self.segments:
            if s.get('file_id') == "COM_REALTIME":
                return

        color = None
        used_colors = {s['color'].upper() for s in self.segments}
        for candidate in self.default_colors:
            if candidate.upper() not in used_colors:
                color = candidate
                break
        if not color:
            color = self.default_colors[self.segment_counter % len(self.default_colors)]

        seg_id = self.segment_counter
        self.segment_counter += 1

        seg = {
            'id': seg_id,
            'name': "串口实时数据",
            'start_time': "00:00:00",
            'end_time': "23:59:59",
            'source_type': 'GGA',
            'color': color,
            'active': True,
            'metrics': None,
            'epochs': [],
            'file_id': "COM_REALTIME"
        }
        self.segments.append(seg)

        item_widget = SegmentListItemWidget(seg_id, "串口实时数据", "00:00:00", "23:59:59", "GGA", color, True, True, True)
        item_widget.active_toggled.connect(self.on_seg_active_toggled)
        item_widget.name_changed.connect(self.on_seg_name_changed)
        item_widget.color_changed.connect(self.on_seg_color_changed)
        item_widget.source_changed.connect(self.on_seg_source_changed)
        item_widget.delete_clicked.connect(self.on_seg_delete_clicked)

        list_item = QListWidgetItem(self.list_segments)
        list_item.setSizeHint(item_widget.sizeHint())
        self.list_segments.addItem(list_item)
        self.list_segments.setItemWidget(list_item, item_widget)

        # 激活全局变量
        if not self.parsed_epochs:
            self.parsed_epochs = []

    def handle_serial_read(self):
        if not self.serial_port.isOpen():
            return

        data = self.serial_port.readAll().data()
        if not data:
            return

        # 写入录制文件
        if self.record_file:
            try:
                self.record_file.write(data)
            except Exception as e:
                self.stop_recording_with_error(str(e))

        has_new_epoch = self.parse_raw_chunk(data)
        if has_new_epoch:
            import time
            cur_time = time.time()
            if cur_time - self.last_live_recompute_time >= 0.5:
                self.recompute_all()
                self.last_live_recompute_time = cur_time

    def parse_raw_chunk(self, data):
        self.serial_buffer.feed(data)
        has_new_epoch = False

        while True:
            res = self.serial_buffer.next_frame()
            if res is None:
                break

            frame_type, frame_data = res

            if frame_type == 'NMEA':
                line_str = frame_data.decode('gbk', errors='replace')

                # 如果没有被阻断，正常输出
                if not getattr(self, '_disable_console_append', False):
                    if self.cb_hex.isChecked():
                        self.safe_append_console(frame_data.hex(' ').upper() + '\n', scroll=False)
                    else:
                        self.safe_append_console(line_str, scroll=False)

                # 检查是否是版本查询返回
                if self.waiting_for_version:
                    is_version = False
                    lower_line = line_str.upper()
                    for kw in ['POLCFGPTVER', 'POLCFGVER', 'POLCFGGETVER', '$BKCHIP', '$POLRS', '$POSYS_BM', '$FWVER', '$HWVER']:
                        if kw in lower_line:
                            is_version = True
                            break
                    if is_version:
                        cleaned_line = line_str.strip()
                        if '*' in cleaned_line:
                            cleaned_line = cleaned_line.split('*')[0]
                        if cleaned_line and cleaned_line not in self.version_lines:
                            self.version_lines.append(cleaned_line)

                        if self.version_timer is not None:
                            self.version_timer.stop()
                        self.version_timer = QTimer(self)
                        self.version_timer.setSingleShot(True)
                        self.version_timer.timeout.connect(self.show_version_dialog)
                        self.version_timer.start(500)

                # 解析 NMEA 行
                epoch = parse_log_line(line_str, self.get_leap_seconds())
                if epoch:
                    if epoch['type'] == 'GSV':
                        prefix = epoch['prefix']
                        total_msg = epoch['total_msg']
                        msg_num = epoch['msg_num']
                        signal_id = epoch['signal_id']

                        # 第一包时清空当前星座该频段的历史数据，防止残留已消失的卫星
                        if msg_num == 1:
                            keys_to_remove = []
                            for k in list(self.gsv_satellites.keys()):
                                mapped_sys, _, _ = get_sat_info(prefix, k[1])
                                if k[0] == mapped_sys:
                                    if signal_id in self.gsv_satellites[k]:
                                        del self.gsv_satellites[k][signal_id]
                                    if not self.gsv_satellites[k]:
                                        keys_to_remove.append(k)
                            for k in keys_to_remove:
                                self.gsv_satellites.pop(k, None)

                        # 缓存可见卫星载噪比及仰角方位角元数据
                        for sat in epoch['sats']:
                            prn = sat['prn']
                            snr = sat['snr']
                            sys_prefix, real_prn, _ = get_sat_info(prefix, prn)
                            key = (sys_prefix, real_prn)
                            if key not in self.gsv_satellites:
                                self.gsv_satellites[key] = {}
                            self.gsv_satellites[key][signal_id] = snr

                            # 更新卫星的仰角和方位角信息
                            elev = sat.get('elevation')
                            azim = sat.get('azimuth')
                            if key not in self.sat_metadata:
                                self.sat_metadata[key] = {}
                            if elev is not None:
                                self.sat_metadata[key]['elevation'] = elev
                            if azim is not None:
                                self.sat_metadata[key]['azimuth'] = azim
                    else:
                        self.process_live_epoch(epoch)
                        has_new_epoch = True

            elif frame_type == 'BK':
                mtype = frame_data[4]
                stype = frame_data[5]

                # 如果没有被阻断，正常输出
                if not getattr(self, '_disable_console_append', False):
                    if self.cb_hex.isChecked():
                        self.safe_append_console(frame_data.hex(' ').upper() + '\n', scroll=False)
                    else:
                        payload_len = ((frame_data[6] & 0x0F) << 8) | frame_data[7]
                        self.safe_append_console(f"[BK 二进制帧] MTYPE={hex(mtype)} STYPE={hex(stype)} 长度={payload_len}\n", scroll=False)

                # 检查是否是二进制版本返回 (MTYPE=0x02, STYPE=0x07)
                if self.waiting_for_version and mtype == 0x02 and stype == 0x07:
                    payload = frame_data[8:]
                    try:
                        version_info = payload.decode('ascii', errors='replace').strip()
                        version_info = "".join(ch for ch in version_info if ch.isprintable())
                        if version_info:
                            cleaned_info = f"[BK 二进制版本] {version_info}"
                            if cleaned_info not in self.version_lines:
                                self.version_lines.append(cleaned_info)

                            if self.version_timer is not None:
                                self.version_timer.stop()
                            self.version_timer = QTimer(self)
                            self.version_timer.setSingleShot(True)
                            self.version_timer.timeout.connect(self.show_version_dialog)
                            self.version_timer.start(500)
                    except Exception:
                        pass

                # 解析 BK 帧
                epoch = parse_bk_frame(frame_data)
                if epoch:
                    if epoch['type'] == 'BK_PNT_NAV':
                        self.process_live_epoch(epoch)
                        has_new_epoch = True

        self.request_console_scroll_to_bottom()
        return has_new_epoch

    def show_version_dialog(self):
        self.waiting_for_version = False
        if not self.version_lines:
            return
        content = "\n".join(self.version_lines)
        self.version_lines = []

        msg = QMessageBox(self)
        msg.setWindowTitle("固件版本信息")
        msg.setText(f"查询到设备固件版本：\n\n{content}")
        msg.setTextInteractionFlags(Qt.TextSelectableByMouse)
        msg.setIcon(QMessageBox.Information)
        msg.exec()

    def reset_version_wait_flag(self):
        self.waiting_for_version = False

    def reset_live_status_ui(self):
        # 1. 重置最新的缓存数据
        self.latest_quality = 0
        self.latest_num_sats = 0
        self.latest_hdop = 1.0
        self.latest_pdop = 1.0

        # 2. 刷新面板坐标和状态展示为默认值
        self.lbl_pnt_utc.setText("--:--:--.--")
        self.lbl_pnt_lat.setText("0.00000000")
        self.lbl_pnt_lon.setText("0.00000000")
        self.lbl_pnt_alt.setText("0.000 米")
        self.lbl_pnt_num.setText("0 颗")
        self.lbl_pnt_pdop.setText("1.0")
        self.lbl_pnt_hdop.setText("1.0")

        # 定位状态卡片重置为 [0] 未定位
        self.lbl_pnt_quality.setText("[0] 未定位")
        self.lbl_pnt_quality.setStyleSheet(
            "background-color: #475569; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
        )

        # 惯导状态卡片重置为 [0] 未激活
        if hasattr(self, 'lbl_ins_status'):
            self.lbl_ins_status.setText("未激活")
            self.lbl_ins_status.setStyleSheet(
                "background-color: #334155; color: #94A3B8; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
            )

        # 3. 清空卫星载噪比缓存及重绘空图表
        self.gsv_satellites = {}
        self.used_satellites.clear()
        self.has_received_gsa = False
        self.sat_metadata.clear()
        self._last_cno_snapshot = None
        if hasattr(self, 'canvas_cno'):
            self.canvas_cno.render_cno(self.gsv_satellites, self.used_satellites, self.has_received_gsa, self.sat_metadata)

    def update_cno_chart(self):
        if hasattr(self, 'canvas_cno') and self.canvas_cno.isVisible():
            if getattr(self.canvas_cno, '_is_resizing', False):
                return
            snapshot = (
                tuple(sorted((key, tuple(sorted(value.items()))) for key, value in self.gsv_satellites.items())),
                tuple(sorted(self.used_satellites)),
                self.has_received_gsa,
                tuple(sorted((key, tuple(sorted(value.items()))) for key, value in self.sat_metadata.items()))
            )
            if snapshot == self._last_cno_snapshot:
                return
            self._last_cno_snapshot = snapshot
            self.canvas_cno.render_cno(self.gsv_satellites, self.used_satellites, self.has_received_gsa, self.sat_metadata)

    def toggle_cno_visibility(self, visible):
        if hasattr(self, 'group_cno'):
            self.group_cno.setVisible(visible)
            if hasattr(self, 'action_toggle_cno'):
                self.action_toggle_cno.setChecked(visible)
            self.adjust_serial_view_layout()
            self.save_config()

    def toggle_serial_ctrl_visibility(self, visible):
        if hasattr(self, 'group_serial_ctrl'):
            self.group_serial_ctrl.setVisible(visible)
            if hasattr(self, 'action_toggle_serial_ctrl'):
                self.action_toggle_serial_ctrl.setChecked(visible)
            self.adjust_serial_view_layout()
            self.save_config()

    def toggle_dashboard_visibility(self, visible):
        if hasattr(self, 'dashboard_tab'):
            self.dashboard_tab.setVisible(visible)
            if hasattr(self, 'action_toggle_dashboard'):
                self.action_toggle_dashboard.setChecked(visible)
            self.adjust_serial_view_layout()
            self.save_config()

    def toggle_console_visibility(self, visible):
        if hasattr(self, 'group_console'):
            self.group_console.setVisible(visible)
            if hasattr(self, 'action_toggle_console'):
                self.action_toggle_console.setChecked(visible)
            self.adjust_serial_view_layout()
            self.save_config()

    def toggle_ref_visibility(self, visible):
        if hasattr(self, 'group_ref'):
            self.group_ref.setVisible(visible)
            if hasattr(self, 'action_toggle_ref'):
                self.action_toggle_ref.setChecked(visible)
            self.adjust_sidebar_visibility()
            self.save_config()

    def toggle_file_visibility(self, visible):
        if hasattr(self, 'group_file'):
            self.group_file.setVisible(visible)
            if hasattr(self, 'action_toggle_file'):
                self.action_toggle_file.setChecked(visible)
            self.adjust_sidebar_visibility()
            self.save_config()

    def adjust_serial_left_panel_visibility(self):
        self.adjust_serial_view_layout()

    def _view_action_checked(self, action_name, widget_name):
        if hasattr(self, action_name):
            return getattr(self, action_name).isChecked()
        return hasattr(self, widget_name) and getattr(self, widget_name).isVisible()

    def adjust_serial_view_layout(self):
        if hasattr(self, 'serial_left_panel') and hasattr(self, 'group_serial_ctrl') and hasattr(self, 'dashboard_tab'):
            serial_ctrl_visible = self._view_action_checked('action_toggle_serial_ctrl', 'group_serial_ctrl')
            dashboard_visible = self._view_action_checked('action_toggle_dashboard', 'dashboard_tab')
            console_visible = self._view_action_checked('action_toggle_console', 'group_console')
            cno_visible = self._view_action_checked('action_toggle_cno', 'group_cno')

            left_visible = serial_ctrl_visible or dashboard_visible
            upper_visible = left_visible or console_visible

            self.serial_left_panel.setVisible(left_visible)
            if hasattr(self, 'serial_upper_splitter'):
                self.serial_upper_splitter.setVisible(upper_visible)

            if upper_visible and left_visible and console_visible and hasattr(self, 'serial_upper_splitter'):
                sizes = self.serial_upper_splitter.sizes()
                if len(sizes) >= 2 and sizes[0] < 50:
                    total_w = sum(sizes)
                    w_left = 320
                    w_right = max(100, total_w - w_left) if total_w > w_left else 920
                    self.serial_upper_splitter.setSizes([w_left, w_right])

            if hasattr(self, 'serial_vertical_splitter'):
                if upper_visible and cno_visible:
                    sizes = self.serial_vertical_splitter.sizes()
                    total_h = sum(sizes) if sizes else self.serial_vertical_splitter.height()
                    if total_h <= 0:
                        total_h = 900
                    if len(sizes) < 2 or sizes[0] < 80 or sizes[1] < 80:
                        h_top = max(260, int(total_h * 0.65))
                        h_bottom = max(180, total_h - h_top)
                        self.serial_vertical_splitter.setSizes([h_top, h_bottom])
                elif upper_visible:
                    self.serial_vertical_splitter.setSizes([1, 0])
                elif cno_visible:
                    self.serial_vertical_splitter.setSizes([0, 1])

    def adjust_sidebar_visibility(self):
        if hasattr(self, 'sidebar_widget') and hasattr(self, 'group_ref') and hasattr(self, 'group_file'):
            ref_visible = self._view_action_checked('action_toggle_ref', 'group_ref')
            file_visible = self._view_action_checked('action_toggle_file', 'group_file')
            sidebar_visible = ref_visible or file_visible

            self.sidebar_widget.setVisible(sidebar_visible)
            if sidebar_visible and hasattr(self, 'main_splitter'):
                sizes = self.main_splitter.sizes()
                if len(sizes) >= 2 and sizes[1] < 80:
                    total_w = sum(sizes)
                    sidebar_w = 340
                    if total_w <= sidebar_w:
                        total_w = self.width()
                    left_w = max(400, total_w - sidebar_w - self.main_splitter.handleWidth())
                    self.main_splitter.setSizes([left_w, sidebar_w])


    def process_live_epoch(self, epoch):
        # 刷新实时超时定时器，2.5秒内没有新数据则自动清空UI
        if hasattr(self, 'realtime_timeout_timer') and self.serial_port.isOpen():
            self.realtime_timeout_timer.start(2500)

        self.add_realtime_segment_item()

        realtime_seg = None
        for s in self.segments:
            if s.get('file_id') == "COM_REALTIME":
                realtime_seg = s
                break
        if not realtime_seg:
            return

        # 1. 如果是主定位语句，更新系统状态缓存，并清空当前 Epoch 在用卫星集合
        if epoch['type'] in ['GGA', 'POSOL', 'BK_PNT_NAV']:
            self.used_satellites.clear()
            self.latest_quality = epoch.get('quality', 0)
            if 'num_sats' in epoch:
                self.latest_num_sats = epoch['num_sats']
            if 'hdop' in epoch:
                self.latest_hdop = epoch['hdop']
            if 'pdop' in epoch:
                self.latest_pdop = epoch['pdop']
            elif 'hdop' in epoch:
                self.latest_pdop = epoch['hdop']

        # 1.5. 如果是 GSA 语句，收集当前在用卫星 PRN
        elif epoch['type'] == 'GSA':
            self.has_received_gsa = True
            sats_used = epoch.get('sats_used', [])
            sentence_type = epoch.get('sentence_type', '')
            talker = sentence_type[1:3] if len(sentence_type) >= 3 else ''

            # 识别基本星座
            gsa_prefix = None
            if talker == 'GP':
                gsa_prefix = 'GPS'
            elif talker in ['BD', 'GB']:
                gsa_prefix = 'BD'
            elif talker == 'GL':
                gsa_prefix = 'GL'
            elif talker == 'GA':
                gsa_prefix = 'GA'

            # 若是多星座合一 NMEA 帧，且存在 system ID 扩展字段
            raw_line = epoch.get('raw_line', '')
            parts = [p.strip() for p in raw_line.split(',')]
            if talker == 'GN' and len(parts) > 18:
                sys_id = parts[18].split('*')[0].strip()
                if sys_id == '1':
                    gsa_prefix = 'GPS'
                elif sys_id == '2':
                    gsa_prefix = 'GL'
                elif sys_id == '3':
                    gsa_prefix = 'GA'
                elif sys_id == '4':
                    gsa_prefix = 'BD'

            for prn in sats_used:
                # 备用方案：若无法获取前缀，根据常规 NMEA 规范 PRN 区间猜测星座
                prn_prefix = gsa_prefix
                if prn_prefix is None:
                    if 1 <= prn <= 32 or 193 <= prn <= 202:
                        prn_prefix = 'GPS'
                    elif 65 <= prn <= 99:
                        prn_prefix = 'GL'
                    elif 141 <= prn <= 172 or 1 <= prn <= 63:
                        prn_prefix = 'BD'
                    else:
                        prn_prefix = 'GPS'

                if prn_prefix:
                    sys_prefix, real_prn, _ = get_sat_info(prn_prefix, prn)
                    self.used_satellites.add((sys_prefix, real_prn))

        # 2. 如果是 GOS 或 DRS，强制继承最新的 GGA/POSOL 系统状态，确保统一性
        elif epoch['type'] in ['POGOS', 'PODRS']:
            epoch['quality'] = self.latest_quality
            epoch['num_sats'] = self.latest_num_sats
            epoch['hdop'] = self.latest_hdop
            epoch['pdop'] = self.latest_pdop

        # 3. 追加到全局实时原始缓冲队列中（限制 6000 帧）
        if epoch['type'] in ['GGA', 'POSOL', 'BK_PNT_NAV', 'POGOS', 'PODRS']:
            self.realtime_raw_epochs.append(epoch)

            # 为了与其他绘图/导出功能兼容，我们需要维护 realtime_seg['epochs']
            # 我们直接把当前过滤出来的 epoch 赋值刷新给 realtime_seg['epochs']，以便外部通过 realtime_seg['epochs'] 读取当前选定源的实时数据
            if realtime_seg['source_type'] == 'GGA':
                realtime_seg['epochs'] = [ep for ep in self.realtime_raw_epochs if ep['type'] in ['GGA', 'POSOL', 'BK_PNT_NAV']][-2000:]
            else:
                realtime_seg['epochs'] = [ep for ep in self.realtime_raw_epochs if ep['type'] == realtime_seg['source_type']][-2000:]

            # 同时更新全局 parsed_epochs 中对应的 COM_REALTIME 帧（上限 2000 点），供全局索引
            # 先给当前帧打上 COM_REALTIME 标记
            epoch['file_id'] = "COM_REALTIME"
            self.parsed_epochs.append(epoch)

            com_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') == "COM_REALTIME"]
            if len(com_epochs) > 2000:
                self.parsed_epochs = [ep for ep in self.parsed_epochs if ep.get('file_id') != "COM_REALTIME"] + com_epochs[-2000:]

            # 建立 file_epochs_map，确保二分查找等其他重计算正常获取
            self.file_epochs_map["COM_REALTIME"] = realtime_seg['epochs']

        # 4. 更新实时解析仪表盘
        self.update_live_dashboard(epoch)

    def update_live_dashboard(self, epoch):
        # 查找当前实时分段选择的数据源类型
        realtime_seg = None
        for s in self.segments:
            if s.get('file_id') == "COM_REALTIME":
                realtime_seg = s
                break

        current_src_type = 'GGA'
        if realtime_seg:
            current_src_type = realtime_seg.get('source_type', 'GGA')

        is_matched_epoch = False
        if current_src_type == 'GGA' and epoch['type'] in ['GGA', 'POSOL', 'BK_PNT_NAV']:
            is_matched_epoch = True
        elif current_src_type == epoch['type']:
            is_matched_epoch = True

        if is_matched_epoch:
            qual = self.latest_quality
            if qual == 0:
                # 未定位状态：重置面板（仅保留时间戳指示数据链连通）
                time_str = epoch.get('time_str', '--:--:--.--')
                self.reset_live_status_ui()
                self.lbl_pnt_utc.setText(time_str)
            else:
                # 已定位状态：更新面板坐标和时间（使用选定数据源的真实值）
                self.lbl_pnt_utc.setText(epoch.get('time_str', '--:--:--.--'))
                self.lbl_pnt_lat.setText(f"{epoch.get('lat', 0.0):.8f}")
                self.lbl_pnt_lon.setText(f"{epoch.get('lon', 0.0):.8f}")
                self.lbl_pnt_alt.setText(f"{epoch.get('alt', 0.0):.3f} 米")

                # 以下系统级指标统一展示系统级最新的缓存数据
                self.lbl_pnt_num.setText(f"{self.latest_num_sats} 颗")
                self.lbl_pnt_pdop.setText(f"{self.latest_pdop:.1f}")
                self.lbl_pnt_hdop.setText(f"{self.latest_hdop:.1f}")

                qual_str = f"[{qual}] 未定位"
                badge_style = "background-color: #475569; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
                if qual == 4:
                    qual_str = f"[{qual}] RTK 固定 (FIXED)"
                    badge_style = "background-color: #10B981; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
                elif qual == 5:
                    qual_str = f"[{qual}] RTK 浮点 (FLOAT)"
                    badge_style = "background-color: #F59E0B; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
                elif qual == 2:
                    qual_str = f"[{qual}] 差分 (DGPS)"
                    badge_style = "background-color: #3B82F6; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
                elif qual == 1:
                    qual_str = f"[{qual}] 单点 (SINGLE)"
                    badge_style = "background-color: #6366F1; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
                elif qual == 6:
                    qual_str = f"[{qual}] 惯导推算 (DR)"
                    badge_style = "background-color: #8B5CF6; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
                elif qual == 7: # BK 二进制下的 RTK Fix
                    qual_str = f"[{qual}] RTK 固定 (FIXED)"
                    badge_style = "background-color: #10B981; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"

                self.lbl_pnt_quality.setText(qual_str)
                self.lbl_pnt_quality.setStyleSheet(badge_style)

        elif epoch['type'] == 'POINS':
            ins_stat = epoch.get('ins_status', 0)
            ins_str = f"[{ins_stat}] 未激活"
            ins_style = "background-color: #475569; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
            if ins_stat == 5:
                ins_str = f"[{ins_stat}] 已收敛"
                ins_style = "background-color: #10B981; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
            elif ins_stat == 4:
                ins_str = f"[{ins_stat}] 未收敛"
                ins_style = "background-color: #F59E0B; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"
            elif ins_stat in [1, 2, 3]:
                ins_str = f"[{ins_stat}] 初始化/对准中"
                ins_style = "background-color: #3B82F6; color: #FFFFFF; font-weight: bold; border-radius: 4px; padding: 2px 6px; font-size: 12px;"

            self.lbl_ins_status.setText(ins_str)
            self.lbl_ins_status.setStyleSheet(ins_style)

            mot = epoch.get('motion_status', 0)
            mot_map = {0: "0: 未知", 1: "1: 静止", 2: "2: 运动", 3: "3: 直线运动", 4: "4: 曲线运动"}
            self.lbl_ins_motion.setText(mot_map.get(mot, "0: 未知"))

            self.lbl_ins_roll.setText(f"{epoch.get('roll', 0.0):.3f} 度")
            self.lbl_ins_pitch.setText(f"{epoch.get('pitch', 0.0):.3f} 度")
            self.lbl_ins_yaw.setText(f"{epoch.get('yaw', 0.0):.3f} 度")
            self.lbl_ins_speed.setText(f"{epoch.get('velocity_forward', 0.0):.2f} m/s")
            self.lbl_ins_mileage.setText(f"{epoch.get('drive_mileage', 0.0):.1f} 米")
            self.lbl_ins_tow.setText(f"{epoch.get('gps_tow', 0.0):.3f}")

    def closeEvent(self, event):
        if hasattr(self, 'serial_port') and self.serial_port.isOpen():
            self.serial_port.close()
        if hasattr(self, 'record_file') and self.record_file:
            try:
                self.record_file.close()
            except Exception:
                pass
        super().closeEvent(event)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
